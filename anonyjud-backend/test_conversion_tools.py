#!/usr/bin/env python3
"""
Script de test pour vérifier les outils de conversion Word → PDF
Teste LibreOffice et unoconv sur Railway
"""

import subprocess
import sys
import logging
import tempfile
import os
from docx import Document
from docx.shared import Inches

# Configuration des logs
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def create_test_docx():
    """Crée un document Word de test"""
    try:
        doc = Document()
        doc.add_heading('Test de Conversion', 0)
        doc.add_paragraph('Ceci est un test de conversion Word → PDF.')
        doc.add_paragraph('Ce document contient :')
        
        # Ajouter une liste
        p = doc.add_paragraph()
        p.add_run('• ').bold = True
        p.add_run('Du texte en gras')
        p = doc.add_paragraph()
        p.add_run('• ').bold = True
        p.add_run('Du texte en italique').italic = True
        p = doc.add_paragraph()
        p.add_run('• ').bold = True
        p.add_run('Des polices différentes')
        
        doc.add_heading('Tableau de test', level=1)
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        
        cells = table.rows[0].cells
        cells[0].text = 'Colonne 1'
        cells[1].text = 'Colonne 2'
        
        cells = table.rows[1].cells
        cells[0].text = 'Données 1'
        cells[1].text = 'Données 2'
        
        # Sauvegarder temporairement
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
            doc.save(temp_file.name)
            logger.info(f"✅ Document Word de test créé: {temp_file.name}")
            return temp_file.name
            
    except Exception as e:
        logger.error(f"❌ Erreur création document test: {str(e)}")
        return None

def test_libreoffice():
    """Test de LibreOffice"""
    logger.info("🔄 Test de LibreOffice...")
    
    # Créer un document de test
    test_docx = create_test_docx()
    if not test_docx:
        return False
    
    try:
        # Commandes LibreOffice à tester
        libreoffice_commands = [
            "soffice",
            "libreoffice",
            "/usr/bin/soffice",
            "/usr/bin/libreoffice"
        ]
        
        # Variables d'environnement
        env = os.environ.copy()
        env.update({
            'HOME': '/tmp',
            'XDG_CONFIG_HOME': '/tmp/.config',
            'XDG_DATA_HOME': '/tmp/.local/share',
            'XDG_CACHE_HOME': '/tmp/.cache'
        })
        
        for cmd in libreoffice_commands:
            try:
                logger.info(f"📄 Test LibreOffice avec: {cmd}")
                
                output_dir = os.path.dirname(test_docx)
                
                result = subprocess.run([
                    cmd,
                    "--headless",
                    "--invisible", 
                    "--nodefault",
                    "--nolockcheck",
                    "--nologo",
                    "--norestore",
                    "--convert-to", "pdf",
                    "--outdir", output_dir,
                    test_docx
                ], capture_output=True, text=True, timeout=60, env=env)
                
                if result.returncode == 0:
                    expected_pdf = test_docx.replace('.docx', '.pdf')
                    if os.path.exists(expected_pdf):
                        file_size = os.path.getsize(expected_pdf)
                        logger.info(f"✅ LibreOffice réussi avec {cmd}")
                        logger.info(f"📄 PDF généré: {expected_pdf} ({file_size} bytes)")
                        
                        # Nettoyer
                        os.unlink(expected_pdf)
                        return True
                    else:
                        logger.warning(f"⚠️ PDF non généré: {expected_pdf}")
                else:
                    logger.debug(f"❌ Échec {cmd}: {result.stderr}")
                    
            except Exception as e:
                logger.debug(f"❌ Erreur {cmd}: {str(e)}")
                continue
        
        logger.error("❌ Aucune commande LibreOffice n'a fonctionné")
        return False
        
    finally:
        if os.path.exists(test_docx):
            os.unlink(test_docx)

def test_unoconv():
    """Test de unoconv"""
    logger.info("🔄 Test de unoconv...")
    
    # Créer un document de test
    test_docx = create_test_docx()
    if not test_docx:
        return False
    
    try:
        # Commandes unoconv à tester
        unoconv_commands = [
            "unoconv",
            "/usr/bin/unoconv"
        ]
        
        for cmd in unoconv_commands:
            try:
                logger.info(f"📄 Test unoconv avec: {cmd}")
                
                output_dir = os.path.dirname(test_docx)
                output_file = os.path.join(output_dir, "test_unoconv.pdf")
                
                result = subprocess.run([
                    cmd,
                    "-f", "pdf",
                    "-o", output_file,
                    test_docx
                ], capture_output=True, text=True, timeout=60)
                
                if result.returncode == 0:
                    if os.path.exists(output_file):
                        file_size = os.path.getsize(output_file)
                        logger.info(f"✅ unoconv réussi avec {cmd}")
                        logger.info(f"📄 PDF généré: {output_file} ({file_size} bytes)")
                        
                        # Nettoyer
                        os.unlink(output_file)
                        return True
                    else:
                        logger.warning(f"⚠️ PDF non généré: {output_file}")
                else:
                    logger.debug(f"❌ Échec {cmd}: {result.stderr}")
                    
            except Exception as e:
                logger.debug(f"❌ Erreur {cmd}: {str(e)}")
                continue
        
        logger.error("❌ Aucune commande unoconv n'a fonctionné")
        return False
        
    finally:
        if os.path.exists(test_docx):
            os.unlink(test_docx)

def test_system_info():
    """Affiche les informations système"""
    logger.info("🔍 Informations système:")
    
    # Version Python
    logger.info(f"🐍 Python: {sys.version}")
    
    # Système d'exploitation
    try:
        with open('/etc/os-release') as f:
            os_info = f.read()
            logger.info(f"🖥️ OS: {os_info.split('PRETTY_NAME=')[1].split('\"')[1]}")
    except:
        logger.info(f"🖥️ OS: {sys.platform}")
    
    # Vérifier les commandes disponibles
    commands_to_check = [
        "soffice", "libreoffice", "unoconv", 
        "/usr/bin/soffice", "/usr/bin/libreoffice", "/usr/bin/unoconv"
    ]
    
    logger.info("📋 Commandes disponibles:")
    for cmd in commands_to_check:
        try:
            result = subprocess.run([cmd, "--version"], 
                                  capture_output=True, text=True, timeout=5)
            if result.returncode == 0:
                version = result.stdout.strip().split('\n')[0]
                logger.info(f"   ✅ {cmd}: {version}")
            else:
                logger.info(f"   ❌ {cmd}: non fonctionnel")
        except:
            logger.info(f"   ❌ {cmd}: non trouvé")

def main():
    """Fonction principale"""
    logger.info("🚀 Début des tests de conversion Word → PDF")
    
    # Informations système
    test_system_info()
    
    # Test LibreOffice
    libreoffice_ok = test_libreoffice()
    
    # Test unoconv
    unoconv_ok = test_unoconv()
    
    # Résumé
    logger.info("📊 Résumé des tests:")
    logger.info(f"   - LibreOffice: {'✅' if libreoffice_ok else '❌'}")
    logger.info(f"   - unoconv: {'✅' if unoconv_ok else '❌'}")
    
    if libreoffice_ok or unoconv_ok:
        logger.info("🎉 Au moins un outil de conversion fonctionne!")
        return True
    else:
        logger.error("❌ Aucun outil de conversion ne fonctionne")
        return False

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1) 