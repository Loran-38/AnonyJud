from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Body
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import Dict, List, Any, Optional
import json
import os
import tempfile
from pathlib import Path
import fitz  # PyMuPDF pour les PDF
from docx import Document  # python-docx pour les fichiers Word
import io
from odf import text as odf_text, teletype
from odf.opendocument import load
import re # Added for regex in deanonymize_docx_file
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import inch
from reportlab.lib.utils import ImageReader

from .anonymizer import anonymize_text
from .deanonymizer import deanonymize_text
from .models import TextAnonymizationRequest, TextDeanonymizationRequest
from .pdf_utils import safe_extract_text_from_pdf, validate_pdf_content

app = FastAPI()

# Configuration CORS pour permettre les requêtes depuis le frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # En production, spécifier l'origine exacte
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.get("/")
def read_root():
    return {"message": "AnonyJud API is running"}

@app.post("/anonymize/text")
def anonymize_text_endpoint(request: TextAnonymizationRequest):
    """
    Anonymise un texte en utilisant les tiers fournis.
    """
    try:
        anonymized, mapping = anonymize_text(request.text, request.tiers)
        return {"anonymized_text": anonymized, "mapping": mapping}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/deanonymize/text")
def deanonymize_text_endpoint(request: TextDeanonymizationRequest):
    """
    Dé-anonymise un texte en utilisant le mapping fourni.
    Si aucun mapping n'est fourni, tente de générer le mapping à partir des tiers.
    """
    try:
        # Si le mapping est vide, générer le mapping à partir des tiers
        if not request.has_mapping or not request.mapping or len(request.mapping) == 0:
            if request.tiers and len(request.tiers) > 0:
                mapping = generate_mapping_from_tiers(request.tiers)
            else:
                # Fallback: essayer de détecter automatiquement
                mapping = detect_anonymized_patterns(request.anonymized_text)
                
                if not mapping:
                    return {"deanonymized_text": request.anonymized_text, "mapping": {}, "message": "Aucun pattern d'anonymisation détecté dans le texte"}
        else:
            # Utiliser le mapping fourni
            mapping = request.mapping
        
        deanonymized = deanonymize_text(request.anonymized_text, mapping)
        
        return {"deanonymized_text": deanonymized, "mapping": mapping}
        
    except Exception as e:
        print(f"❌ Erreur dans deanonymize_text_endpoint: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/anonymize/file")
async def anonymize_file(
    file: UploadFile = File(...),
    tiers_json: str = Form(...)
):
    """
    Anonymise un fichier Word, PDF ou ODT en utilisant les tiers fournis.
    """
    try:
        # Convertir la chaîne JSON en liste de tiers
        tiers = json.loads(tiers_json)
        
        # Vérifier le type de fichier
        filename = file.filename or ""
        file_extension = os.path.splitext(filename)[1].lower()
        
        if file_extension == ".pdf":
            # Traitement des fichiers PDF
            content = await file.read()
            pdf_text, mapping = extract_and_anonymize_pdf(content, tiers)
            return {"text": pdf_text, "mapping": mapping}
            
        elif file_extension in [".doc", ".docx"]:
            # Traitement des fichiers Word
            content = await file.read()
            doc_text, mapping = extract_and_anonymize_docx(content, tiers)
            return {"text": doc_text, "mapping": mapping}
            
        elif file_extension == ".odt":
            # Traitement des fichiers ODT (OpenDocument Text)
            content = await file.read()
            odt_text, mapping = extract_and_anonymize_odt(content, tiers)
            return {"text": odt_text, "mapping": mapping}
            
        else:
            raise HTTPException(status_code=400, detail="Format de fichier non supporté. Utilisez PDF, DOCX ou ODT.")
            
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

def detect_anonymized_patterns(text: str) -> Dict[str, str]:
    """
    Détecte automatiquement les patterns d'anonymisation dans un texte.
    Utilise les mêmes patterns que l'anonymisation (NOM1, PRENOM1, etc.)
    Retourne un mapping des patterns détectés.
    """
    import re
    
    print(f"🔍 DETECT_ANONYMIZED_PATTERNS - Début de la détection")
    print(f"📝 Texte à analyser (premiers 500 chars): {text[:500]}...")
    
    mapping = {}
    
    # Patterns pour détecter les champs anonymisés (même format que l'anonymisation)
    patterns = [
        (r'\bNOM(\d+)\b', 'NOM'),
        (r'\bPRENOM(\d+)\b', 'PRENOM'),
        (r'\bADRESSE(\d+)\b', 'ADRESSE'),
        (r'\bNUMERO(\d+)\b', 'NUMERO'),
        (r'\bVOIE(\d+)\b', 'VOIE'),
        (r'\bCODEPOSTAL(\d+)\b', 'CODEPOSTAL'),
        (r'\bVILLE(\d+)\b', 'VILLE'),
        (r'\bTEL(\d+)\b', 'TEL'),
        (r'\bPORTABLE(\d+)\b', 'PORTABLE'),
        (r'\bEMAIL(\d+)\b', 'EMAIL'),
        (r'\bSOCIETE(\d+)\b', 'SOCIETE'),
        (r'\bPERSO(\d+)\b', 'PERSO'),
        # Patterns génériques pour les champs personnalisés
        (r'\b([A-Z]+)(\d+)\b', 'CUSTOM'),
    ]
    
    # Chercher tous les patterns dans le texte
    for pattern, field_type in patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        print(f"🔍 Pattern '{pattern}' -> {len(matches)} correspondances trouvées")
        
        for match in matches:
            if field_type == 'CUSTOM':
                # Pour les patterns personnalisés, match est un tuple (prefix, number)
                prefix, number = match
                tag = f"{prefix.upper()}{number}"
            else:
                # Pour les patterns standards, match est juste le numéro
                number = match if isinstance(match, str) else match
                tag = f"{field_type}{number}"
            
            # Créer un mapping de base (tag -> tag pour l'instant)
            # En réalité, nous n'avons pas les valeurs originales, donc on garde les tags
            mapping[tag] = tag
            print(f"✅ Pattern détecté: '{tag}'")
    
    print(f"📊 Total des patterns détectés: {len(mapping)}")
    print(f"🗂️ Mapping généré: {mapping}")
    print(f"🏁 DETECT_ANONYMIZED_PATTERNS - Fin de la détection")
    
    return mapping

def generate_mapping_from_tiers(tiers: List[Dict[str, Any]]) -> Dict[str, str]:
    """
    Génère le mapping d'anonymisation à partir des tiers.
    Utilise la même logique que l'anonymisation pour créer les balises.
    """
    import re
    
    print(f"🔧 GENERATE_MAPPING_FROM_TIERS - Début de la génération")
    print(f"📊 Nombre de tiers reçus: {len(tiers)}")
    
    mapping = {}
    
    for tier_index, tier in enumerate(tiers):
        # Utiliser le numéro fixe du tiers ou fallback sur l'index + 1
        tier_number = tier.get("numero", tier_index + 1)
        print(f"🔍 Traitement du tiers {tier_number}: {tier}")
        
        # Traiter le nom
        if tier.get("nom"):
            nom = tier["nom"].strip()
            if nom and len(nom) > 1:
                tag = f"NOM{tier_number}"
                mapping[tag] = nom
                print(f"✅ Ajouté: {tag} -> {nom}")
        
        # Traiter le prénom
        if tier.get("prenom"):
            prenom = tier["prenom"].strip()
            if prenom and len(prenom) > 1:
                tag = f"PRENOM{tier_number}"
                mapping[tag] = prenom
                print(f"✅ Ajouté: {tag} -> {prenom}")
        
        # Traiter les composants de l'adresse
        if tier.get("adresse_numero"):
            numero = tier["adresse_numero"].strip()
            if numero and len(numero) > 0:
                tag = f"NUMERO{tier_number}"
                mapping[tag] = numero
                print(f"✅ Ajouté: {tag} -> {numero}")
        
        if tier.get("adresse_voie"):
            voie = tier["adresse_voie"].strip()
            if voie and len(voie) > 2:
                tag = f"VOIE{tier_number}"
                mapping[tag] = voie
                print(f"✅ Ajouté: {tag} -> {voie}")
        
        if tier.get("adresse_code_postal"):
            code_postal = tier["adresse_code_postal"].strip()
            if code_postal and len(code_postal) > 0:
                tag = f"CODEPOSTAL{tier_number}"
                mapping[tag] = code_postal
                print(f"✅ Ajouté: {tag} -> {code_postal}")
        
        if tier.get("adresse_ville"):
            ville = tier["adresse_ville"].strip()
            if ville and len(ville) > 1:
                tag = f"VILLE{tier_number}"
                mapping[tag] = ville
                print(f"✅ Ajouté: {tag} -> {ville}")
        
        # Traiter la ville (format simple, pour compatibilité)
        if tier.get("ville") and not tier.get("adresse_ville"):
            ville = tier["ville"].strip()
            if ville and len(ville) > 1:
                tag = f"VILLE{tier_number}"
                mapping[tag] = ville
                print(f"✅ Ajouté: {tag} -> {ville}")
        
        # Traiter l'adresse complète (compatibilité)
        if tier.get("adresse"):
            adresse = tier["adresse"].strip()
            if adresse and len(adresse) > 5:
                tag = f"ADRESSE{tier_number}"
                mapping[tag] = adresse
                print(f"✅ Ajouté: {tag} -> {adresse}")
        
        # Traiter le téléphone
        if tier.get("telephone"):
            tel = tier["telephone"].strip()
            if tel and len(tel) > 5:
                tag = f"TEL{tier_number}"
                mapping[tag] = tel
                print(f"✅ Ajouté: {tag} -> {tel}")
        
        # Traiter le portable
        if tier.get("portable"):
            portable = tier["portable"].strip()
            if portable and len(portable) > 5:
                tag = f"PORTABLE{tier_number}"
                mapping[tag] = portable
                print(f"✅ Ajouté: {tag} -> {portable}")
        
        # Traiter l'email
        if tier.get("email"):
            email = tier["email"].strip()
            if email and '@' in email:
                tag = f"EMAIL{tier_number}"
                mapping[tag] = email
                print(f"✅ Ajouté: {tag} -> {email}")
        
        # Traiter la société
        if tier.get("societe"):
            societe = tier["societe"].strip()
            if societe and len(societe) > 1:
                tag = f"SOCIETE{tier_number}"
                mapping[tag] = societe
                print(f"✅ Ajouté: {tag} -> {societe}")
        
        # Traiter les champs personnalisés
        if tier.get("customFields") and isinstance(tier["customFields"], list):
            for custom_field in tier["customFields"]:
                if isinstance(custom_field, dict):
                    champ_value = custom_field.get("value")
                    champ_label = custom_field.get("label")
                    
                    if champ_value and isinstance(champ_value, str):
                        champ_value = champ_value.strip()
                        if champ_value and len(champ_value) > 1:
                            label_base = "PERSO"
                            if champ_label and isinstance(champ_label, str):
                                label_perso = champ_label.strip()
                                if label_perso:
                                    label_base = re.sub(r'[^A-Za-z]', '', label_perso.upper())
                                    if not label_base:
                                        label_base = "PERSO"
                            
                            tag = f"{label_base}{tier_number}"
                            mapping[tag] = champ_value
                            print(f"✅ Ajouté: {tag} -> {champ_value}")
        
        # Traiter le champ personnalisé (ancien format)
        if tier.get("champPerso"):
            champ_perso = tier["champPerso"]
            if champ_perso and isinstance(champ_perso, str):
                champ_perso = champ_perso.strip()
                if champ_perso and len(champ_perso) > 1:
                    label_base = "PERSO"
                    if tier.get("labelChampPerso") and isinstance(tier["labelChampPerso"], str):
                        label_perso = tier["labelChampPerso"].strip()
                        if label_perso:
                            label_base = re.sub(r'[^A-Za-z]', '', label_perso.upper())
                            if not label_base:
                                label_base = "PERSO"
                    
                    tag = f"{label_base}{tier_number}"
                    mapping[tag] = champ_perso
                    print(f"✅ Ajouté: {tag} -> {champ_perso}")
    
    print(f"🏁 GENERATE_MAPPING_FROM_TIERS - Mapping généré avec {len(mapping)} éléments")
    print(f"🗂️ Mapping final: {mapping}")
    return mapping

@app.post("/deanonymize/file")
async def deanonymize_file(
    file: UploadFile = File(...),
    mapping_json: str = Form(...),
    tiers_json: str = Form("[]"),
    has_mapping: str = Form("true")
):
    """
    Dé-anonymise un fichier Word, PDF ou ODT en utilisant le mapping fourni.
    Si le mapping est vide, essaie de détecter automatiquement les patterns.
    """
    try:
        print(f"🚀 DEANONYMIZE_FILE ENDPOINT - Début du traitement")
        print(f"📁 Fichier reçu: {file.filename}")
        print(f"🗂️ Mapping JSON brut: {mapping_json}")
        print(f"🗂️ Tiers JSON brut: {tiers_json}")
        print(f"🔄 A mapping: {has_mapping}")
        
        # Convertir la chaîne JSON en mapping
        mapping = json.loads(mapping_json)
        tiers = json.loads(tiers_json)
        print(f"🗂️ Mapping parsé: {mapping}")
        print(f"📊 Nombre de balises dans le mapping: {len(mapping)}")
        print(f"👥 Nombre de tiers: {len(tiers)}")
        
        # Vérifier le type de fichier
        filename = file.filename or ""
        file_extension = os.path.splitext(filename)[1].lower()
        print(f"📄 Extension du fichier: {file_extension}")
        
        # Lire le contenu du fichier
        content = await file.read()
        print(f"📦 Taille du fichier: {len(content)} bytes")
        
        # Si le mapping est vide, générer le mapping à partir des tiers
        if has_mapping.lower() == "false" or not mapping or len(mapping) == 0:
            print(f"⚠️ Mapping vide détecté, génération à partir des tiers...")
            if tiers and len(tiers) > 0:
                mapping = generate_mapping_from_tiers(tiers)
                print(f"🔧 Mapping généré à partir des tiers: {mapping}")
            else:
                print(f"❌ Aucun tiers disponible pour générer le mapping")
                # Fallback: essayer de détecter automatiquement
                print(f"🔍 Tentative de détection automatique...")
                # Extraire d'abord le texte pour détecter les patterns
                if file_extension == ".pdf":
        with fitz.open(stream=content, filetype="pdf") as pdf:
            text = ""
            for page in pdf:
                text += page.get_text()
                elif file_extension in [".doc", ".docx"]:
                    doc = Document(io.BytesIO(content))
                    text = ""
                    for para in doc.paragraphs:
                        text += para.text + "\n"
                elif file_extension == ".odt":
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".odt") as temp_file:
                        temp_file.write(content)
                        temp_path = temp_file.name
                    try:
                        doc = load(temp_path)
                        text = ""
                        for paragraph in doc.getElementsByType(odf_text.P):
                            text += teletype.extractText(paragraph) + "\n"
                    finally:
                        if os.path.exists(temp_path):
                            os.unlink(temp_path)
                else:
                    raise HTTPException(status_code=400, detail="Format de fichier non supporté. Utilisez PDF, DOCX ou ODT.")
                
                # Détecter les patterns anonymisés automatiquement
                mapping = detect_anonymized_patterns(text)
                print(f"🔍 Patterns détectés automatiquement: {mapping}")
                
                if not mapping:
                    print(f"❌ Aucun pattern d'anonymisation détecté")
                    return {"text": text, "mapping": {}, "message": "Aucun pattern d'anonymisation détecté dans le fichier"}
        
        print(f"🔄 Début de la désanonymisation avec mapping: {mapping}")
        
        # Procéder à la dé-anonymisation
        if file_extension == ".pdf":
            print(f"📄 Traitement PDF...")
            pdf_text = extract_and_deanonymize_pdf(content, mapping)
            print(f"✅ PDF désanonymisé avec succès")
            return {"text": pdf_text, "mapping": mapping}
            
        elif file_extension in [".doc", ".docx"]:
            print(f"📄 Traitement DOCX...")
            doc_text = extract_and_deanonymize_docx(content, mapping)
            print(f"✅ DOCX désanonymisé avec succès")
            return {"text": doc_text, "mapping": mapping}
            
        elif file_extension == ".odt":
            print(f"📄 Traitement ODT...")
            odt_text = extract_and_deanonymize_odt(content, mapping)
            print(f"✅ ODT désanonymisé avec succès")
            return {"text": odt_text, "mapping": mapping}
            
        else:
            raise HTTPException(status_code=400, detail="Format de fichier non supporté. Utilisez PDF, DOCX ou ODT.")
            
    except Exception as e:
        print(f"❌ Erreur dans deanonymize_file endpoint: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/anonymize/file/download")
async def anonymize_file_download(
    file: UploadFile = File(...),
    tiers_json: str = Form(...)
):
    """
    Anonymise un fichier Word ou ODT et retourne le fichier modifié pour téléchargement.
    """
    try:
        print(f"🚀 ANONYMIZE_FILE_DOWNLOAD - Début du traitement")
        print(f"📁 Fichier reçu: {file.filename}")
        
        # Convertir la chaîne JSON en liste de tiers
        tiers = json.loads(tiers_json)
        print(f"👥 Nombre de tiers: {len(tiers)}")
        
        # Vérifier le type de fichier
        filename = file.filename or ""
        file_extension = os.path.splitext(filename)[1].lower()
        print(f"📄 Extension du fichier: {file_extension}")
        
        if file_extension == ".pdf":
            print(f"📄 Traitement fichier PDF...")
            # Traitement des fichiers PDF - Utilisation de la méthode sécurisée par défaut
            content = await file.read()
            anonymized_file, mapping = anonymize_pdf_secure_with_graphics(content, tiers)
            
            # Créer un nom de fichier pour le téléchargement
            base_name = os.path.splitext(filename)[0]
            anonymized_filename = f"{base_name}_ANONYM_SECURE.pdf"
            
            print(f"✅ Fichier PDF anonymisé: {anonymized_filename}")
            
            # Retourner le fichier modifié
            return StreamingResponse(
                io.BytesIO(anonymized_file),
                media_type="application/pdf",
                headers={"Content-Disposition": f"attachment; filename={anonymized_filename}"}
            )
        elif file_extension in [".doc", ".docx"]:
            print(f"📄 Traitement fichier Word...")
            # Traitement des fichiers Word
            content = await file.read()
            anonymized_file, mapping = anonymize_docx_file(content, tiers)
            
            # Créer un nom de fichier pour le téléchargement
            base_name = os.path.splitext(filename)[0]
            anonymized_filename = f"{base_name}_ANONYM.docx"
            
            print(f"✅ Fichier Word anonymisé: {anonymized_filename}")
            
            # Retourner le fichier modifié
            return StreamingResponse(
                io.BytesIO(anonymized_file),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f"attachment; filename={anonymized_filename}"}
            )
        elif file_extension == ".odt":
            print(f"📄 Traitement fichier ODT...")
            # Traitement des fichiers ODT
            content = await file.read()
            anonymized_file, mapping = anonymize_odt_file(content, tiers)
            
            # Créer un nom de fichier pour le téléchargement
            base_name = os.path.splitext(filename)[0]
            anonymized_filename = f"{base_name}_ANONYM.odt"
            
            print(f"✅ Fichier ODT anonymisé: {anonymized_filename}")
            
            # Retourner le fichier modifié
            return StreamingResponse(
                io.BytesIO(anonymized_file),
                media_type="application/vnd.oasis.opendocument.text",
                headers={"Content-Disposition": f"attachment; filename={anonymized_filename}"}
            )
        else:
            print(f"❌ Format de fichier non supporté: {file_extension}")
            raise HTTPException(status_code=400, detail="Seuls les fichiers PDF (.pdf), Word (.docx) et ODT (.odt) sont supportés pour le téléchargement.")
            
    except Exception as e:
        print(f"❌ Erreur dans anonymize_file_download: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/deanonymize/file/download")
async def deanonymize_file_download(
    file: UploadFile = File(...),
    mapping_json: str = Form(...)
):
    """
    Dé-anonymise un fichier Word ou ODT et retourne le fichier modifié pour téléchargement.
    """
    try:
        print(f"🚀 DEANONYMIZE_FILE_DOWNLOAD - Début du traitement")
        print(f"📁 Fichier reçu: {file.filename}")
        
        # Convertir la chaîne JSON en mapping
        mapping = json.loads(mapping_json)
        print(f"🗂️ Mapping reçu: {mapping}")
        print(f"📊 Nombre de balises dans le mapping: {len(mapping)}")
        
        # Vérifier le type de fichier
        filename = file.filename or ""
        file_extension = os.path.splitext(filename)[1].lower()
        print(f"📄 Extension du fichier: {file_extension}")
        
        if file_extension == ".pdf":
            print(f"📄 Traitement fichier PDF...")
            # Traitement des fichiers PDF - Utilisation de la méthode sécurisée
            content = await file.read()
            deanonymized_file = deanonymize_pdf_secure_with_graphics(content, mapping)
            
            # Créer un nom de fichier pour le téléchargement
            base_name = os.path.splitext(filename)[0]
            # Retirer les suffixes d'anonymisation si présents
            for suffix in ["_ANONYM_SECURE", "_ANONYM"]:
                if base_name.endswith(suffix):
                    base_name = base_name[:-len(suffix)]
                    break
            deanonymized_filename = f"{base_name}_DESANONYM_SECURE.pdf"
            
            print(f"✅ Fichier PDF dé-anonymisé: {deanonymized_filename}")
            
            # Retourner le fichier modifié
            return StreamingResponse(
                io.BytesIO(deanonymized_file),
                media_type="application/pdf",
                headers={"Content-Disposition": f"attachment; filename={deanonymized_filename}"}
            )
        elif file_extension in [".doc", ".docx"]:
            print(f"📄 Traitement fichier Word...")
            # Traitement des fichiers Word
            content = await file.read()
            deanonymized_file = deanonymize_docx_file(content, mapping)
            
            # Créer un nom de fichier pour le téléchargement
            base_name = os.path.splitext(filename)[0]
            # Retirer "_ANONYM" du nom si présent
            if base_name.endswith("_ANONYM"):
                base_name = base_name[:-7]
            deanonymized_filename = f"{base_name}_DESANONYM.docx"
            
            print(f"✅ Fichier Word dé-anonymisé: {deanonymized_filename}")
            
            # Retourner le fichier modifié
            return StreamingResponse(
                io.BytesIO(deanonymized_file),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f"attachment; filename={deanonymized_filename}"}
            )
        elif file_extension == ".odt":
            print(f"📄 Traitement fichier ODT...")
            # Traitement des fichiers ODT
            content = await file.read()
            deanonymized_file = deanonymize_odt_file(content, mapping)
            
            # Créer un nom de fichier pour le téléchargement
            base_name = os.path.splitext(filename)[0]
            # Retirer "_ANONYM" du nom si présent
            if base_name.endswith("_ANONYM"):
                base_name = base_name[:-7]
            deanonymized_filename = f"{base_name}_DESANONYM.odt"
            
            print(f"✅ Fichier ODT dé-anonymisé: {deanonymized_filename}")
            
            # Retourner le fichier modifié
            return StreamingResponse(
                io.BytesIO(deanonymized_file),
                media_type="application/vnd.oasis.opendocument.text",
                headers={"Content-Disposition": f"attachment; filename={deanonymized_filename}"}
            )
        else:
            print(f"❌ Format de fichier non supporté: {file_extension}")
            raise HTTPException(status_code=400, detail="Seuls les fichiers PDF (.pdf), Word (.docx) et ODT (.odt) sont supportés pour le téléchargement.")
            
    except Exception as e:
        print(f"❌ Erreur dans deanonymize_file_download: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

def extract_and_anonymize_pdf(content: bytes, tiers: List[Dict[str, Any]]):
    """
    Extrait le texte d'un PDF et l'anonymise de manière sécurisée.
    Utilise safe_extract_text_from_pdf pour gérer les erreurs get_text().
    """
    try:
        print(f"🔍 EXTRACT_AND_ANONYMIZE_PDF - Début du traitement sécurisé")
        print(f"👥 Nombre de tiers: {len(tiers)}")
        
        # Extraction sécurisée du texte
        text, extraction_success = safe_extract_text_from_pdf(content)
        
        if not extraction_success:
            print(f"⚠️ Extraction partielle ou échec - tentative avec méthode de fallback")
            # Fallback: essayer l'ancienne méthode pour compatibilité
            try:
                with fitz.open(stream=content, filetype="pdf") as pdf:
                    text = ""
                    for page_num in range(pdf.page_count):
                        try:
                            page = pdf[page_num]
                            if page is not None:
                                try:
                                    page_text = page.get_text()
                                    text += page_text
                                except Exception as page_error:
                                    print(f"⚠️ Erreur get_text page {page_num + 1}: {str(page_error)}")
                                    # Essayer avec get_text("text")
                                    try:
                                        page_text = page.get_text("text")
                                        text += page_text
                                    except Exception as page_error2:
                                        print(f"⚠️ Erreur get_text('text') page {page_num + 1}: {str(page_error2)}")
                                        # Passer à la page suivante sans arrêter le processus
                                        continue
                        except Exception as e:
                            print(f"⚠️ Erreur traitement page {page_num + 1}: {str(e)}")
                            continue
                            
                print(f"📝 Fallback: {len(text)} caractères extraits")
            except Exception as fallback_error:
                print(f"❌ Fallback aussi en échec: {str(fallback_error)}")
                if not text:  # Si aucun texte n'a été extrait
                    raise Exception(f"Impossible d'extraire le texte du PDF: {str(fallback_error)}")
        
        if not text or len(text.strip()) == 0:
            print(f"⚠️ Aucun texte extrait du PDF")
            return "", {}
        
        print(f"📝 Texte extrait: {len(text)} caractères")
        print(f"📝 Aperçu (premiers 200 chars): {text[:200]}...")
        
        # Anonymiser le texte extrait
        anonymized, mapping = anonymize_text(text, tiers)
        
        print(f"🔒 Anonymisation terminée: {len(mapping)} remplacements")
        return anonymized, mapping
        
    except Exception as e:
        print(f"❌ Erreur dans extract_and_anonymize_pdf: {str(e)}")
        raise Exception(f"Erreur lors du traitement du PDF: {str(e)}")

def extract_and_anonymize_docx(content: bytes, tiers: List[Dict[str, Any]]):
    """
    Extrait le texte d'un document Word et l'anonymise.
    """
    try:
        # Ouvrir le document Word depuis les bytes
        doc = Document(io.BytesIO(content))
        text = ""
        
        # Extraire le texte de chaque paragraphe
        for para in doc.paragraphs:
            text += para.text + "\n"
            
        # Anonymiser le texte extrait
        anonymized, mapping = anonymize_text(text, tiers)
        return anonymized, mapping
        
    except Exception as e:
        raise Exception(f"Erreur lors du traitement du document Word: {str(e)}")

def extract_and_anonymize_odt(content: bytes, tiers: List[Dict[str, Any]]):
    """
    Extrait le texte d'un document OpenDocument Text (ODT) et l'anonymise.
    """
    try:
        # Créer un fichier temporaire pour stocker le contenu ODT
        with tempfile.NamedTemporaryFile(delete=False, suffix=".odt") as temp_file:
            temp_file.write(content)
            temp_path = temp_file.name
        
        try:
            # Charger le document ODT
            doc = load(temp_path)
            
            # Extraire le texte du document
            text = ""
            
            # Parcourir tous les éléments de texte dans le document
            for paragraph in doc.getElementsByType(odf_text.P):
                text += teletype.extractText(paragraph) + "\n"
            
            # Anonymiser le texte extrait
            anonymized, mapping = anonymize_text(text, tiers)
            
            return anonymized, mapping
            
        finally:
            # Supprimer le fichier temporaire
            if os.path.exists(temp_path):
                os.unlink(temp_path)
        
    except Exception as e:
        raise Exception(f"Erreur lors du traitement du document ODT: {str(e)}") 

def anonymize_docx_file(content: bytes, tiers: List[Dict[str, Any]]):
    """
    Anonymise directement un fichier Word en modifiant son contenu.
    Préserve le formatage (police, style, majuscules/minuscules, mise en page).
    Retourne le fichier modifié et le mapping d'anonymisation.
    """
    try:
        print(f"DEBUG: Début anonymize_docx_file avec {len(tiers)} tiers")
        
        # Ouvrir le document Word depuis les bytes
        doc = Document(io.BytesIO(content))
        
        # Collecter tout le texte du document
        full_text = ""
        for para in doc.paragraphs:
            full_text += para.text + "\n"
        
        print(f"DEBUG: Texte extrait (premiers 200 chars): {full_text[:200]}...")
        
        # Anonymiser le texte complet pour obtenir le mapping
        anonymized_text, mapping = anonymize_text(full_text, tiers)
        
        print(f"DEBUG: Mapping généré: {mapping}")
        print(f"DEBUG: Texte anonymisé (premiers 200 chars): {anonymized_text[:200]}...")
        
        # Fonction pour anonymiser un run en préservant son formatage
        def anonymize_run_preserving_format(run, mapping):
            """Anonymise le texte d'un run en préservant son formatage"""
            if not run.text:
                return
                
            original_text = run.text
            modified_text = original_text
            
            # Appliquer chaque remplacement du mapping (inverser pour que les tags remplacent les originaux)
            for tag, original_value in mapping.items():
                if original_value in modified_text:
                    # Préserver la casse du texte original
                    if original_value.isupper():
                        modified_text = modified_text.replace(original_value, tag.upper())
                    elif original_value.islower():
                        modified_text = modified_text.replace(original_value, tag.lower())
                    elif original_value.istitle():
                        modified_text = modified_text.replace(original_value, tag.title())
                    else:
                        modified_text = modified_text.replace(original_value, tag)
            
            # Remplacer le texte du run seulement si modifié (le formatage est automatiquement préservé)
            if modified_text != original_text:
                print(f"DEBUG: Run modifié: '{original_text}' -> '{modified_text}'")
                run.text = modified_text
        
        # Appliquer les anonymisations directement dans le document en préservant le formatage
        paragraphs_processed = 0
        for para in doc.paragraphs:
            if para.text.strip():  # Seulement pour les paragraphes non vides
                # Traiter chaque run individuellement pour préserver le formatage
                for run in para.runs:
                    anonymize_run_preserving_format(run, mapping)
                paragraphs_processed += 1
        
        # Traiter également les tableaux
        cells_processed = 0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            # Traiter chaque run individuellement pour préserver le formatage
                            for run in para.runs:
                                anonymize_run_preserving_format(run, mapping)
                            cells_processed += 1
        
        print(f"DEBUG: Traitement terminé - {paragraphs_processed} paragraphes, {cells_processed} cellules")
        
        # Sauvegarder le document modifié en bytes
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        print(f"DEBUG: Fichier anonymisé généré avec succès")
        return output.getvalue(), mapping
        
    except Exception as e:
        print(f"DEBUG: Erreur dans anonymize_docx_file: {str(e)}")
        raise Exception(f"Erreur lors de l'anonymisation du fichier Word: {str(e)}")

def deanonymize_docx_file(content: bytes, mapping: Dict[str, str]):
    """
    Dé-anonymise directement un fichier Word en utilisant le mapping fourni.
    Préserve le formatage (police, style, majuscules/minuscules, mise en page).
    Retourne le fichier modifié.
    """
    try:
        print(f"🔍 DEANONYMIZE_DOCX_FILE - Début du processus")
        print(f"🗂️ Mapping reçu: {mapping}")
        print(f"📊 Nombre de balises dans le mapping: {len(mapping)}")
        
        # Ouvrir le document Word depuis les bytes
        doc = Document(io.BytesIO(content))
        
        # Extraire tout le texte du document pour analyse
        full_text = ""
        for para in doc.paragraphs:
            full_text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        full_text += para.text + "\n"
        
        print(f"📝 Texte extrait du document (premiers 300 chars): {full_text[:300]}...")
        
        # Analyser quelles balises sont présentes dans le texte
        found_tags = []
        for tag in mapping.keys():
            if tag in full_text:
                found_tags.append(tag)
                print(f"✅ Balise '{tag}' trouvée dans le document")
            else:
                print(f"❌ Balise '{tag}' NON trouvée dans le document")
        
        print(f"📋 Résumé: {len(found_tags)}/{len(mapping)} balises trouvées: {found_tags}")
        
        # Le mapping est déjà dans le bon sens (balise -> valeur_originale)
        # Pas besoin d'inverser car generate_mapping_from_tiers() crée déjà le mapping correct
        print(f"🔄 Mapping reçu (balise -> valeur): {mapping}")
        
        # Vérifier si le mapping est dans le bon sens
        sample_key = list(mapping.keys())[0] if mapping else ""
        if sample_key and not sample_key.isupper():
            # Le mapping semble être dans le mauvais sens (valeur -> balise), l'inverser
            reverse_mapping = {v: k for k, v in mapping.items()}
            print(f"🔄 Mapping inversé car dans le mauvais sens: {reverse_mapping}")
        else:
            # Le mapping est dans le bon sens (balise -> valeur)
            reverse_mapping = mapping
            print(f"🔄 Mapping utilisé tel quel: {reverse_mapping}")
        
        # Trier les balises par longueur décroissante pour éviter les remplacements partiels
        sorted_tags = sorted(reverse_mapping.keys(), key=len, reverse=True)
        print(f"🔢 Balises triées par longueur: {sorted_tags}")
        
        # Fonction pour dé-anonymiser un run en préservant son formatage
        def deanonymize_run_preserving_format(run, mapping, sorted_tags):
            """Dé-anonymise le texte d'un run en préservant son formatage"""
            if not run.text:
                return False
                
            original_text = run.text
            modified_text = original_text
            has_changes = False
            
            # Appliquer chaque remplacement du mapping (balise -> valeur originale)
            for balise in sorted_tags:
                valeur_originale = mapping[balise]
                
                # Chercher la balise dans toutes les variantes de casse possibles
                balise_variants = [
                    balise,           # PRENOM1
                    balise.lower(),   # prenom1
                    balise.title(),   # Prenom1
                    balise.capitalize() # Prenom1 (même résultat que title pour ce cas)
                ]
                
                for variant in balise_variants:
                    if variant in modified_text:
                        count_before = modified_text.count(variant)
                        
                        # Utiliser une expression régulière pour remplacer la balise exacte (pas de remplacement partiel)
                        pattern = re.compile(r'\b' + re.escape(variant) + r'\b')
                        modified_text_new = pattern.sub(valeur_originale, modified_text)
                        
                        count_after = modified_text_new.count(variant)
                        actual_replacements = count_before - count_after
                        
                        if actual_replacements > 0:
                            print(f"✅ Run: {actual_replacements} occurrence(s) de '{variant}' remplacée(s) par '{valeur_originale}'")
                            modified_text = modified_text_new
                            has_changes = True
                            break  # Sortir de la boucle des variants une fois qu'on a trouvé et remplacé
                        else:
                            print(f"⚠️ Run: Balise '{variant}' présente mais aucun remplacement de mot entier effectué")
                            print(f"❌ DEBUG: La balise '{variant}' pourrait être une sous-chaîne d'une autre balise (ex: NOM dans PRENOM)")
                            # ✅ NE PLUS FAIRE DE FALLBACK avec replace() simple car cela cause le problème PRENOM
                            # Le fallback replace() sans limites de mots causait: PRENOM1 -> PREHuissoud1
                            print(f"🚫 Pas de remplacement fallback pour éviter les remplacements partiels dans d'autres balises")
                            # break supprimé car on continue à chercher les autres variants
            
            # Remplacer le texte du run seulement si modifié (le formatage est automatiquement préservé)
            if has_changes:
                print(f"📝 Run modifié: '{original_text}' -> '{modified_text}'")
                run.text = modified_text
            
            return has_changes
        
        paragraphs_modified = 0
        runs_modified = 0
        
        # Appliquer les dé-anonymisations dans les paragraphes
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():  # Seulement pour les paragraphes non vides
                para_has_changes = False
                # Traiter chaque run individuellement pour préserver le formatage
                for run in para.runs:
                    if deanonymize_run_preserving_format(run, reverse_mapping, sorted_tags):
                        runs_modified += 1
                        para_has_changes = True
                
                if para_has_changes:
                    paragraphs_modified += 1
        
        # Traiter également les tableaux
        cells_modified = 0
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, para in enumerate(cell.paragraphs):
                        if para.text.strip():
                            cell_has_changes = False
                            # Traiter chaque run individuellement pour préserver le formatage
                            for run in para.runs:
                                if deanonymize_run_preserving_format(run, reverse_mapping, sorted_tags):
                                    runs_modified += 1
                                    cell_has_changes = True
                            
                            if cell_has_changes:
                                cells_modified += 1
        
        print(f"📈 Résultats: {paragraphs_modified} paragraphes, {cells_modified} cellules, {runs_modified} runs modifiés")
        
        # Sauvegarder le document modifié en bytes
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        print(f"🏁 DEANONYMIZE_DOCX_FILE - Fichier modifié généré avec succès")
        return output.getvalue()
        
    except Exception as e:
        print(f"❌ Erreur dans deanonymize_docx_file: {str(e)}")
        raise Exception(f"Erreur lors de la dé-anonymisation du fichier Word: {str(e)}")

def extract_and_deanonymize_pdf(content: bytes, mapping: Dict[str, str]):
    """
    Extrait le texte d'un PDF et le dé-anonymise de manière sécurisée.
    Utilise safe_extract_text_from_pdf pour gérer les erreurs get_text().
    """
    try:
        print(f"🔍 EXTRACT_AND_DEANONYMIZE_PDF - Début du traitement sécurisé")
        print(f"🗂️ Mapping reçu: {mapping}")
        print(f"📊 Nombre de balises: {len(mapping)}")
        
        # Extraction sécurisée du texte
        text, extraction_success = safe_extract_text_from_pdf(content)
        
        if not extraction_success:
            print(f"⚠️ Extraction partielle ou échec - tentative avec méthode de fallback")
            # Fallback: essayer l'ancienne méthode pour compatibilité
            try:
                with fitz.open(stream=content, filetype="pdf") as pdf:
                    text = ""
                    for page_num in range(pdf.page_count):
                        try:
                            page = pdf[page_num]
                            if page is not None:
                                try:
                                    page_text = page.get_text()
                                    text += page_text
                                except Exception as page_error:
                                    print(f"⚠️ Erreur get_text page {page_num + 1}: {str(page_error)}")
                                    # Essayer avec get_text("text")
                                    try:
                                        page_text = page.get_text("text")
                                        text += page_text
                                    except Exception as page_error2:
                                        print(f"⚠️ Erreur get_text('text') page {page_num + 1}: {str(page_error2)}")
                                        # Passer à la page suivante sans arrêter le processus
                                        continue
                        except Exception as e:
                            print(f"⚠️ Erreur traitement page {page_num + 1}: {str(e)}")
                            continue
                            
                print(f"📝 Fallback: {len(text)} caractères extraits")
            except Exception as fallback_error:
                print(f"❌ Fallback aussi en échec: {str(fallback_error)}")
                if not text:  # Si aucun texte n'a été extrait
                    raise Exception(f"Impossible d'extraire le texte du PDF: {str(fallback_error)}")
        
        if not text or len(text.strip()) == 0:
            print(f"⚠️ Aucun texte extrait du PDF")
            return ""
        
        print(f"📝 Texte extrait: {len(text)} caractères")
        print(f"📝 Aperçu (premiers 200 chars): {text[:200]}...")
        
        # Dé-anonymiser le texte extrait
        deanonymized = deanonymize_text(text, mapping)
        
        print(f"🔓 Dé-anonymisation terminée")
        return deanonymized
        
    except Exception as e:
        print(f"❌ Erreur dans extract_and_deanonymize_pdf: {str(e)}")
        raise Exception(f"Erreur lors du traitement du PDF: {str(e)}")

def extract_and_deanonymize_docx(content: bytes, mapping: Dict[str, str]):
    """
    Extrait le texte d'un document Word et le dé-anonymise.
    """
    try:
        print(f"🔍 EXTRACT_AND_DEANONYMIZE_DOCX - Début du processus")
        print(f"🗂️ Mapping reçu: {mapping}")
        
        # Ouvrir le document Word depuis les bytes
        doc = Document(io.BytesIO(content))
        text = ""
        
        # Extraire le texte de chaque paragraphe
        paragraph_count = 0
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
                paragraph_count += 1
        
        # Extraire le texte des tableaux aussi
        table_count = 0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            text += para.text + "\n"
                            table_count += 1
        
        print(f"📝 Texte extrait: {paragraph_count} paragraphes, {table_count} cellules de tableau")
        print(f"📝 Texte complet (premiers 300 chars): {text[:300]}...")
        print(f"📊 Longueur totale du texte: {len(text)} caractères")
            
        # Dé-anonymiser le texte extrait
        print(f"🔄 Appel de deanonymize_text...")
        deanonymized = deanonymize_text(text, mapping)
        print(f"✅ Désanonymisation terminée")
        
        return deanonymized
        
    except Exception as e:
        print(f"❌ Erreur dans extract_and_deanonymize_docx: {str(e)}")
        raise Exception(f"Erreur lors du traitement du document Word: {str(e)}")

def extract_and_deanonymize_odt(content: bytes, mapping: Dict[str, str]):
    """
    Extrait le texte d'un document OpenDocument Text (ODT) et le dé-anonymise.
    """
    try:
        # Créer un fichier temporaire pour stocker le contenu ODT
        with tempfile.NamedTemporaryFile(delete=False, suffix=".odt") as temp_file:
            temp_file.write(content)
            temp_path = temp_file.name
        
        try:
            # Charger le document ODT
            doc = load(temp_path)
            
            # Extraire le texte du document
            text = ""
            
            # Parcourir tous les éléments de texte dans le document
            for paragraph in doc.getElementsByType(odf_text.P):
                text += teletype.extractText(paragraph) + "\n"
            
            # Dé-anonymiser le texte extrait
            deanonymized = deanonymize_text(text, mapping)
            
            return deanonymized
            
        finally:
            # Supprimer le fichier temporaire
            if os.path.exists(temp_path):
                os.unlink(temp_path)
        
    except Exception as e:
        raise Exception(f"Erreur lors du traitement du document ODT: {str(e)}") 

def anonymize_odt_file(content: bytes, tiers: List[Dict[str, Any]]):
    """
    Anonymise directement un fichier ODT en modifiant son contenu.
    Préserve le formatage et retourne le fichier modifié et le mapping d'anonymisation.
    """
    try:
        print(f"🚀 DEBUG: Début anonymize_odt_file avec {len(tiers)} tiers")
        
        # Créer un fichier temporaire pour stocker le contenu ODT
        with tempfile.NamedTemporaryFile(delete=False, suffix=".odt") as temp_file:
            temp_file.write(content)
            temp_path = temp_file.name
        
        try:
            # Charger le document ODT
            doc = load(temp_path)
            
            # Collecter tout le texte du document
            full_text = ""
            for paragraph in doc.getElementsByType(odf_text.P):
                full_text += teletype.extractText(paragraph) + "\n"
            
            print(f"📝 DEBUG: Texte extrait (premiers 200 chars): {full_text[:200]}...")
            
            # Anonymiser le texte complet pour obtenir le mapping
            anonymized_text, mapping = anonymize_text(full_text, tiers)
            
            print(f"🗂️ DEBUG: Mapping généré: {mapping}")
            print(f"📝 DEBUG: Texte anonymisé (premiers 200 chars): {anonymized_text[:200]}...")
            
            # Appliquer l'anonymisation au document ODT
            paragraphs_processed = 0
            for paragraph in doc.getElementsByType(odf_text.P):
                paragraph_text = teletype.extractText(paragraph)
                if paragraph_text.strip():
                    # Anonymiser le texte du paragraphe
                    anonymized_paragraph = anonymize_text(paragraph_text, tiers)[0]
                    
                    # Remplacer le contenu du paragraphe de manière sécurisée
                    try:
                        # Vider le paragraphe
                        paragraph.childNodes = []
                        # Ajouter le texte anonymisé
                        paragraph.addText(anonymized_paragraph)
                        paragraphs_processed += 1
                    except Exception as e:
                        print(f"⚠️ DEBUG: Erreur lors du traitement du paragraphe: {str(e)}")
                        # Continuer avec le paragraphe suivant
                        continue
            
            print(f"📊 DEBUG: Traitement terminé - {paragraphs_processed} paragraphes")
            
            # Sauvegarder le document modifié
            output_path = temp_path + "_anonym"
            doc.save(output_path)
            
            # Lire le fichier anonymisé
            with open(output_path, 'rb') as f:
                anonymized_file_content = f.read()
            
            # Nettoyer les fichiers temporaires
            if os.path.exists(output_path):
                os.unlink(output_path)
            
            print(f"✅ DEBUG: Fichier ODT anonymisé généré avec succès")
            return anonymized_file_content, mapping
            
        finally:
            # Supprimer le fichier temporaire
            if os.path.exists(temp_path):
                os.unlink(temp_path)
        
    except Exception as e:
        print(f"❌ DEBUG: Erreur dans anonymize_odt_file: {str(e)}")
        raise Exception(f"Erreur lors de l'anonymisation du fichier ODT: {str(e)}")

def deanonymize_odt_file(content: bytes, mapping: Dict[str, str]):
    """
    Dé-anonymise directement un fichier ODT en utilisant le mapping fourni.
    Préserve le formatage et retourne le fichier modifié.
    """
    try:
        print(f"🚀 DEBUG: Début deanonymize_odt_file")
        print(f"🗂️ DEBUG: Mapping reçu: {mapping}")
        print(f"📊 DEBUG: Nombre de balises dans le mapping: {len(mapping)}")
        
        # Créer un fichier temporaire pour stocker le contenu ODT
        with tempfile.NamedTemporaryFile(delete=False, suffix=".odt") as temp_file:
            temp_file.write(content)
            temp_path = temp_file.name
        
        try:
            # Charger le document ODT
            doc = load(temp_path)
            
            # Extraire tout le texte du document pour analyse
            full_text = ""
            for paragraph in doc.getElementsByType(odf_text.P):
                full_text += teletype.extractText(paragraph) + "\n"
            
            print(f"📝 DEBUG: Texte extrait (premiers 300 chars): {full_text[:300]}...")
            
            # Analyser quelles balises sont présentes dans le texte
            found_tags = []
            for tag in mapping.keys():
                if tag in full_text:
                    found_tags.append(tag)
                    print(f"✅ DEBUG: Balise '{tag}' trouvée dans le document")
                else:
                    print(f"❌ DEBUG: Balise '{tag}' NON trouvée dans le document")
            
            print(f"📋 DEBUG: Résumé: {len(found_tags)}/{len(mapping)} balises trouvées: {found_tags}")
            
            # Appliquer la dé-anonymisation au document ODT
            paragraphs_processed = 0
            for paragraph in doc.getElementsByType(odf_text.P):
                paragraph_text = teletype.extractText(paragraph)
                if paragraph_text.strip():
                    # Dé-anonymiser le texte du paragraphe
                    deanonymized_paragraph = deanonymize_text(paragraph_text, mapping)
                    
                    # Si le texte a changé, le remplacer
                    if deanonymized_paragraph != paragraph_text:
                        try:
                            # Vider le paragraphe
                            paragraph.childNodes = []
                            # Ajouter le texte dé-anonymisé
                            paragraph.addText(deanonymized_paragraph)
                            paragraphs_processed += 1
                        except Exception as e:
                            print(f"⚠️ DEBUG: Erreur lors du traitement du paragraphe: {str(e)}")
                            # Continuer avec le paragraphe suivant
                            continue
            
            print(f"📊 DEBUG: Traitement terminé - {paragraphs_processed} paragraphes modifiés")
            
            # Sauvegarder le document modifié
            output_path = temp_path + "_deanonym"
            doc.save(output_path)
            
            # Lire le fichier dé-anonymisé
            with open(output_path, 'rb') as f:
                deanonymized_file_content = f.read()
            
            # Nettoyer les fichiers temporaires
            if os.path.exists(output_path):
                os.unlink(output_path)
            
            print(f"✅ DEBUG: Fichier ODT dé-anonymisé généré avec succès")
            return deanonymized_file_content
            
        finally:
            # Supprimer le fichier temporaire
            if os.path.exists(temp_path):
                os.unlink(temp_path)
        
    except Exception as e:
        print(f"❌ DEBUG: Erreur dans deanonymize_odt_file: {str(e)}")
        raise Exception(f"Erreur lors de la dé-anonymisation du fichier ODT: {str(e)}") 

def create_pdf_from_text(text: str, filename: str) -> bytes:
    """
    Crée un nouveau PDF à partir du texte fourni en utilisant reportlab.
    Préserve les sauts de ligne et la mise en forme basique.
    """
    try:
        print(f"🚀 CREATE_PDF_FROM_TEXT - Début de la génération PDF")
        print(f"📄 Nom du fichier: {filename}")
        print(f"📝 Longueur du texte: {len(text)} caractères")
        
        # Créer un buffer en mémoire pour le PDF
        buffer = io.BytesIO()
        
        # Créer le document PDF
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=72,  # 1 inch
            leftMargin=72,   # 1 inch
            topMargin=72,    # 1 inch
            bottomMargin=72  # 1 inch
        )
        
        # Définir les styles
        styles = getSampleStyleSheet()
        
        # Style pour le titre
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=16,
            spaceAfter=30,
            alignment=1  # Centré
        )
        
        # Style pour le texte normal
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            leading=14,  # Espacement entre les lignes
            spaceAfter=12,
            alignment=0  # Justifié à gauche
        )
        
        # Construire le contenu du PDF
        story = []
        
        # Ajouter le titre
        title = f"Document traité - {filename}"
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 20))
        
        # Diviser le texte en paragraphes
        paragraphs = text.split('\n')
        
        paragraph_count = 0
        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text:  # Ignorer les lignes vides
                # Échapper les caractères spéciaux pour reportlab
                para_text = para_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                
                # Ajouter le paragraphe
                story.append(Paragraph(para_text, normal_style))
                paragraph_count += 1
            else:
                # Ajouter un espacement pour les lignes vides
                story.append(Spacer(1, 6))
        
        print(f"📊 Nombre de paragraphes traités: {paragraph_count}")
        
        # Générer le PDF
        doc.build(story)
        
        # Récupérer les bytes du PDF
        buffer.seek(0)
        pdf_bytes = buffer.getvalue()
        buffer.close()
        
        print(f"✅ PDF généré avec succès, taille: {len(pdf_bytes)} bytes")
        return pdf_bytes
        
    except Exception as e:
        print(f"❌ Erreur lors de la génération du PDF: {str(e)}")
        raise Exception(f"Erreur lors de la génération du PDF: {str(e)}")

def anonymize_pdf_file(content: bytes, tiers: List[Dict[str, Any]]):
    """
    Anonymise un fichier PDF en extrayant le texte, l'anonymisant, 
    puis générant un nouveau PDF avec le texte anonymisé.
    """
    try:
        print(f"🚀 ANONYMIZE_PDF_FILE - Début du traitement")
        print(f"👥 Nombre de tiers: {len(tiers)}")
        
        # Extraire le texte du PDF original de manière sécurisée
        text, extraction_success = safe_extract_text_from_pdf(content)
        
        if not extraction_success:
            print(f"⚠️ Extraction partielle ou échec - tentative avec méthode de fallback")
            # Fallback: essayer l'ancienne méthode pour compatibilité
            try:
                with fitz.open(stream=content, filetype="pdf") as pdf:
                    text = ""
                    page_count = 0
                    for page_num in range(pdf.page_count):
                        try:
                            page = pdf[page_num]
                            if page is not None:
                                try:
                                    page_text = page.get_text()
                                    text += page_text
                                    page_count += 1
                                except Exception as page_error:
                                    print(f"⚠️ Erreur get_text page {page_num + 1}: {str(page_error)}")
                                    # Essayer avec get_text("text")
                                    try:
                                        page_text = page.get_text("text")
                                        text += page_text
                                        page_count += 1
                                    except Exception as page_error2:
                                        print(f"⚠️ Erreur get_text('text') page {page_num + 1}: {str(page_error2)}")
                                        # Passer à la page suivante sans arrêter le processus
                                        continue
                        except Exception as e:
                            print(f"⚠️ Erreur traitement page {page_num + 1}: {str(e)}")
                            continue
                            
                print(f"📝 Fallback: {len(text)} caractères extraits de {page_count} pages")
            except Exception as fallback_error:
                print(f"❌ Fallback aussi en échec: {str(fallback_error)}")
                if not text:  # Si aucun texte n'a été extrait
                    raise Exception(f"Impossible d'extraire le texte du PDF: {str(fallback_error)}")
        else:
            # Compter les pages pour les logs
            try:
                with fitz.open(stream=content, filetype="pdf") as pdf:
                    page_count = pdf.page_count
            except:
                page_count = "inconnu"
        
        if not text or len(text.strip()) == 0:
            print(f"⚠️ Aucun texte extrait du PDF")
            return b"", {}
        
        print(f"📄 Texte extrait de {page_count} pages")
        print(f"📝 Longueur du texte extrait: {len(text)} caractères")
        
        # Anonymiser le texte
        anonymized_text, mapping = anonymize_text(text, tiers)
        
        print(f"🔒 Texte anonymisé, {len(mapping)} remplacements")
        
        # Générer le nouveau PDF avec le texte anonymisé
        pdf_bytes = create_pdf_from_text(anonymized_text, "document_anonymise.pdf")
        
        print(f"✅ PDF anonymisé généré avec succès")
        return pdf_bytes, mapping
        
    except Exception as e:
        print(f"❌ Erreur dans anonymize_pdf_file: {str(e)}")
        raise Exception(f"Erreur lors de l'anonymisation du fichier PDF: {str(e)}")

def deanonymize_pdf_file(content: bytes, mapping: Dict[str, str]):
    """
    Dé-anonymise un fichier PDF en extrayant le texte, le dé-anonymisant,
    puis générant un nouveau PDF avec le texte dé-anonymisé.
    """
    try:
        print(f"🚀 DEANONYMIZE_PDF_FILE - Début du traitement")
        print(f"🗂️ Mapping reçu: {mapping}")
        print(f"📊 Nombre de balises dans le mapping: {len(mapping)}")
        
        # Extraire le texte du PDF anonymisé
        with fitz.open(stream=content, filetype="pdf") as pdf:
            text = ""
            page_count = 0
            for page in pdf:
                text += page.get_text()
                page_count += 1
        
        print(f"📄 Texte extrait de {page_count} pages")
        print(f"📝 Longueur du texte extrait: {len(text)} caractères")
        
        # Dé-anonymiser le texte
        deanonymized_text = deanonymize_text(text, mapping)
        
        print(f"🔓 Texte dé-anonymisé")
        
        # Générer le nouveau PDF avec le texte dé-anonymisé
        pdf_bytes = create_pdf_from_text(deanonymized_text, "document_desanonymise.pdf")
        
        print(f"✅ PDF dé-anonymisé généré avec succès")
        return pdf_bytes
        
    except Exception as e:
        print(f"❌ Erreur dans deanonymize_pdf_file: {str(e)}")
        raise Exception(f"Erreur lors de la dé-anonymisation du fichier PDF: {str(e)}") 

def extract_pdf_elements(doc):
    """
    Extrait tous les éléments du PDF : texte, images, graphiques avec leurs positions.
    """
    pdf_elements = []
    
    for page_num in range(doc.page_count):
        page = doc[page_num]
        page_elements = {
            "page_number": page_num,
            "page_size": page.rect,
            "text_elements": [],
            "images": [],
            "drawings": []
        }
        
        # Extraire le texte avec positions exactes
        text_dict = page.get_text("dict")
        for block in text_dict.get("blocks", []):
            if "lines" in block:  # Bloc de texte
                for line in block["lines"]:
                    for span in line["spans"]:
                        text_element = {
                            "text": span["text"],
                            "bbox": span["bbox"],
                            "font": span["font"],
                            "size": span["size"],
                            "flags": span["flags"],
                            "color": span.get("color", 0)
                        }
                        page_elements["text_elements"].append(text_element)
        
        # Extraire les images
        image_list = page.get_images()
        for img_index, img in enumerate(image_list):
            try:
                xref = img[0]
                pix = fitz.Pixmap(doc, xref)
                if pix.n < 5:  # GRAY or RGB
                    img_data = pix.tobytes("png")
                else:  # CMYK: convert first
                    pix1 = fitz.Pixmap(fitz.csRGB, pix)
                    img_data = pix1.tobytes("png")
                    pix1 = None
                
                # Obtenir la position de l'image sur la page
                img_rect = page.get_image_bbox(img)
                
                image_element = {
                    "data": img_data,
                    "bbox": img_rect,
                    "xref": xref
                }
                page_elements["images"].append(image_element)
                pix = None
            except Exception as e:
                print(f"⚠️ Erreur lors de l'extraction d'image: {e}")
        
        # Extraire les dessins/graphiques vectoriels
        try:
            drawings = page.get_drawings()
            for drawing in drawings:
                page_elements["drawings"].append(drawing)
        except Exception as e:
            print(f"⚠️ Erreur lors de l'extraction des dessins: {e}")
        
        pdf_elements.append(page_elements)
    
    return pdf_elements

def anonymize_pdf_secure_with_graphics(pdf_content: bytes, tiers: List[Any]) -> tuple[bytes, Dict[str, str]]:
    """
    Anonymise un PDF de manière sécurisée en remplaçant RÉELLEMENT le texte
    tout en préservant images, graphiques et mise en page exacte.
    
    Méthode sécurisée :
    1. Extrait tous les éléments (texte, images, graphiques)
    2. Remplace le texte de manière irréversible
    3. Reconstitue le PDF avec reportlab en préservant la mise en page
    """
    try:
        print(f"🔒 ANONYMIZE_PDF_SECURE_WITH_GRAPHICS - Début du traitement sécurisé")
        
        # Ouvrir le PDF original
        doc = fitz.open(stream=pdf_content, filetype="pdf")
        mapping = {}
        replacement_counter = 1
        
        print(f"📄 PDF ouvert: {doc.page_count} pages")
        
        # Créer les remplacements basés sur les tiers
        replacements = {}
        for tiers_data in tiers:
            numero = tiers_data.get('numero', replacement_counter)
            print(f"🔍 DEBUG - Traitement tiers numéro {numero}: {tiers_data}")
            
            if tiers_data.get('nom'):
                original_nom = tiers_data['nom'].strip()
                anonymized_nom = f"NOM{numero}"  # ✅ CORRIGÉ: MAJUSCULES comme le reste du système
                replacements[original_nom] = anonymized_nom
                mapping[anonymized_nom] = original_nom
                print(f"📝 DEBUG - Nom: '{original_nom}' -> '{anonymized_nom}' (CORRIGÉ: majuscules!)")
                
            if tiers_data.get('prenom'):
                original_prenom = tiers_data['prenom'].strip()
                anonymized_prenom = f"PRENOM{numero}"  # ✅ CORRIGÉ: MAJUSCULES comme le reste du système
                replacements[original_prenom] = anonymized_prenom
                mapping[anonymized_prenom] = original_prenom
                print(f"📝 DEBUG - Prénom: '{original_prenom}' -> '{anonymized_prenom}' (CORRIGÉ: majuscules!)")
                
            if tiers_data.get('adresse'):
                original_adresse = tiers_data['adresse'].strip()
                anonymized_adresse = f"ADRESSE{numero}"  # ✅ CORRIGÉ: MAJUSCULES comme le reste du système
                replacements[original_adresse] = anonymized_adresse
                mapping[anonymized_adresse] = original_adresse
                print(f"📝 DEBUG - Adresse: '{original_adresse}' -> '{anonymized_adresse}' (CORRIGÉ: majuscules!)")
            
            # Ajouter support pour les autres champs pour être cohérent avec anonymize_text()
            if tiers_data.get('telephone'):
                original_tel = tiers_data['telephone'].strip()
                anonymized_tel = f"TEL{numero}"
                replacements[original_tel] = anonymized_tel
                mapping[anonymized_tel] = original_tel
                print(f"📝 DEBUG - Téléphone: '{original_tel}' -> '{anonymized_tel}'")
            
            if tiers_data.get('portable'):
                original_portable = tiers_data['portable'].strip()
                anonymized_portable = f"PORTABLE{numero}"
                replacements[original_portable] = anonymized_portable
                mapping[anonymized_portable] = original_portable
                print(f"📝 DEBUG - Portable: '{original_portable}' -> '{anonymized_portable}'")
            
            if tiers_data.get('email'):
                original_email = tiers_data['email'].strip()
                anonymized_email = f"EMAIL{numero}"
                replacements[original_email] = anonymized_email
                mapping[anonymized_email] = original_email
                print(f"📝 DEBUG - Email: '{original_email}' -> '{anonymized_email}'")
            
            if tiers_data.get('societe'):
                original_societe = tiers_data['societe'].strip()
                anonymized_societe = f"SOCIETE{numero}"
                replacements[original_societe] = anonymized_societe
                mapping[anonymized_societe] = original_societe
                print(f"📝 DEBUG - Société: '{original_societe}' -> '{anonymized_societe}'")
                
            replacement_counter += 1
        
        print(f"🔄 {len(replacements)} remplacements à effectuer")
        
        # Extraire tous les éléments du PDF
        pdf_elements = extract_pdf_elements(doc)
        doc.close()  # Fermer le document original
        
        # Anonymiser le texte dans les éléments extraits
        for page_data in pdf_elements:
            for text_element in page_data["text_elements"]:
                original_text = text_element["text"]
                anonymized_text = original_text
                
                # Appliquer tous les remplacements
                for original, anonymized in replacements.items():
                    if original.lower() in anonymized_text.lower():
                        anonymized_text = anonymized_text.replace(original, anonymized)
                        print(f"🔄 Remplacement sécurisé: '{original}' → '{anonymized}'")
                
                # Remplacer DÉFINITIVEMENT le texte
                text_element["text"] = anonymized_text
        
        # Reconstituer le PDF avec reportlab
        buffer = io.BytesIO()
        
        # Utiliser reportlab pour créer le nouveau PDF
        from reportlab.pdfgen import canvas as rl_canvas
        
        # Créer le document avec la taille de la première page
        first_page = pdf_elements[0] if pdf_elements else None
        if first_page:
            page_size = (first_page["page_size"].width, first_page["page_size"].height)
        else:
            page_size = A4
            
        c = rl_canvas.Canvas(buffer, pagesize=page_size)
        
        # Reconstituer chaque page
        for page_data in pdf_elements:
            page_size = (page_data["page_size"].width, page_data["page_size"].height)
            c.setPageSize(page_size)
            
            # Ajouter les images d'abord (arrière-plan)
            for image_element in page_data["images"]:
                try:
                    img_data = image_element["data"]
                    bbox = image_element["bbox"]
                    
                    # Créer un ImageReader à partir des données PNG
                    img_reader = ImageReader(io.BytesIO(img_data))
                    
                    # Dessiner l'image à sa position exacte
                    c.drawImage(
                        img_reader,
                        bbox.x0, page_size[1] - bbox.y1,  # Conversion coordonnées
                        width=bbox.width,
                        height=bbox.height
                    )
                except Exception as e:
                    print(f"⚠️ Erreur lors de l'ajout d'image: {e}")
            
            # Ajouter les dessins vectoriels
            for drawing in page_data["drawings"]:
                try:
                    # Simplification : dessiner les formes basiques
                    if drawing.get("type") == "l":  # Ligne
                        points = drawing.get("items", [])
                        if len(points) >= 2:
                            c.line(points[0][1][0], page_size[1] - points[0][1][1],
                                  points[1][1][0], page_size[1] - points[1][1][1])
                    elif drawing.get("type") == "re":  # Rectangle
                        rect = drawing.get("rect")
                        if rect:
                            c.rect(rect[0], page_size[1] - rect[3], 
                                  rect[2] - rect[0], rect[3] - rect[1])
                except Exception as e:
                    print(f"⚠️ Erreur lors de l'ajout de dessin: {e}")
            
            # Ajouter le texte anonymisé avec la mise en forme exacte
            for text_element in page_data["text_elements"]:
                try:
                    text = text_element["text"].strip()
                    if not text:
                        continue
                        
                    bbox = text_element["bbox"]
                    font_size = text_element["size"]
                    
                    # Convertir les coordonnées (PDF vs reportlab)
                    x = bbox[0]
                    y = page_size[1] - bbox[3]  # Inverser Y
                    
                    # Appliquer la police et la taille
                    c.setFont("Helvetica", font_size)
                    
                    # Dessiner le texte à la position exacte
                    c.drawString(x, y, text)
                    
                except Exception as e:
                    print(f"⚠️ Erreur lors de l'ajout de texte: {e}")
            
            # Passer à la page suivante
            c.showPage()
        
        # Finaliser le PDF
        c.save()
        pdf_bytes = buffer.getvalue()
        
        print(f"✅ PDF anonymisé sécurisé avec graphiques préservés généré")
        print(f"🗂️ Mapping créé avec {len(mapping)} entrées")
        
        return pdf_bytes, mapping
        
    except Exception as e:
        print(f"❌ Erreur dans anonymize_pdf_secure_with_graphics: {str(e)}")
        raise Exception(f"Erreur lors de l'anonymisation sécurisée du PDF: {str(e)}")

def deanonymize_pdf_secure_with_graphics(pdf_content: bytes, mapping: Dict[str, str]) -> bytes:
    """
    Dé-anonymise un PDF en restaurant le texte original de manière sécurisée
    tout en préservant les images et graphiques.
    """
    try:
        print(f"🔒 DEANONYMIZE_PDF_SECURE_WITH_GRAPHICS - Début du traitement")
        print(f"📊 DEBUG - Mapping reçu: {mapping}")
        print(f"📊 DEBUG - Nombre de balises dans le mapping: {len(mapping)}")
        
        # Analyser le mapping pour identifier le problème de casse
        for tag, original in mapping.items():
            if tag.islower():
                print(f"❌ DEBUG - PROBLÈME DE CASSE DÉTECTÉ: balise '{tag}' en minuscules (devrait être en majuscules)")
            else:
                print(f"✅ DEBUG - Balise correcte: '{tag}' en majuscules")
        
        # Créer les remplacements inverses
        reverse_replacements = {anonymized: original for anonymized, original in mapping.items()}
        print(f"🔄 DEBUG - Remplacements inverses créés: {reverse_replacements}")
        
        # Même processus que l'anonymisation mais avec les remplacements inversés
        doc = fitz.open(stream=pdf_content, filetype="pdf")
        
        # Extraire tous les éléments
        pdf_elements = extract_pdf_elements(doc)
        doc.close()
        
        # Dé-anonymiser le texte
        for page_data in pdf_elements:
            for text_element in page_data["text_elements"]:
                anonymized_text = text_element["text"]
                restored_text = anonymized_text
                
                # ✅ CORRECTION: Appliquer les remplacements inverses avec limites de mots
                # Trier par longueur décroissante pour éviter les remplacements partiels
                sorted_anonymized = sorted(reverse_replacements.keys(), key=len, reverse=True)
                
                for anonymized in sorted_anonymized:
                    original = reverse_replacements[anonymized]
                    if anonymized in restored_text:
                        # Utiliser une expression régulière avec limites de mots pour éviter 
                        # le problème PRENOM1 -> PREHuissoud1 quand on remplace NOM1
                        pattern = re.compile(r'\b' + re.escape(anonymized) + r'\b')
                        if pattern.search(restored_text):
                            restored_text = pattern.sub(original, restored_text)
                            print(f"✅ Restauration avec limites de mots: '{anonymized}' → '{original}'")
                        else:
                            print(f"⚠️ Balise '{anonymized}' trouvée mais pas en tant que mot entier (probablement dans une autre balise)")
                
                text_element["text"] = restored_text
        
        # Reconstituer le PDF (même logique que l'anonymisation)
        buffer = io.BytesIO()
        from reportlab.pdfgen import canvas as rl_canvas
        
        first_page = pdf_elements[0] if pdf_elements else None
        if first_page:
            page_size = (first_page["page_size"].width, first_page["page_size"].height)
        else:
            page_size = A4
            
        c = rl_canvas.Canvas(buffer, pagesize=page_size)
        
        # Reconstituer chaque page avec le texte restauré
        for page_data in pdf_elements:
            page_size = (page_data["page_size"].width, page_data["page_size"].height)
            c.setPageSize(page_size)
            
            # Images
            for image_element in page_data["images"]:
                try:
                    img_data = image_element["data"]
                    bbox = image_element["bbox"]
                    img_reader = ImageReader(io.BytesIO(img_data))
                    c.drawImage(
                        img_reader,
                        bbox.x0, page_size[1] - bbox.y1,
                        width=bbox.width,
                        height=bbox.height
                    )
                except Exception as e:
                    print(f"⚠️ Erreur image: {e}")
            
            # Dessins
            for drawing in page_data["drawings"]:
                try:
                    if drawing.get("type") == "l":
                        points = drawing.get("items", [])
                        if len(points) >= 2:
                            c.line(points[0][1][0], page_size[1] - points[0][1][1],
                                  points[1][1][0], page_size[1] - points[1][1][1])
                    elif drawing.get("type") == "re":
                        rect = drawing.get("rect")
                        if rect:
                            c.rect(rect[0], page_size[1] - rect[3], 
                                  rect[2] - rect[0], rect[3] - rect[1])
                except Exception as e:
                    print(f"⚠️ Erreur dessin: {e}")
            
            # Texte restauré
            for text_element in page_data["text_elements"]:
                try:
                    text = text_element["text"].strip()
                    if not text:
                        continue
                        
                    bbox = text_element["bbox"]
                    font_size = text_element["size"]
                    x = bbox[0]
                    y = page_size[1] - bbox[3]
                    
                    c.setFont("Helvetica", font_size)
                    c.drawString(x, y, text)
                    
                except Exception as e:
                    print(f"⚠️ Erreur texte: {e}")
            
            c.showPage()
        
        c.save()
        pdf_bytes = buffer.getvalue()
        
        print(f"✅ PDF dé-anonymisé sécurisé généré")
        return pdf_bytes
        
    except Exception as e:
        print(f"❌ Erreur dans deanonymize_pdf_secure_with_graphics: {str(e)}")
        raise Exception(f"Erreur lors de la dé-anonymisation sécurisée: {str(e)}") 
        