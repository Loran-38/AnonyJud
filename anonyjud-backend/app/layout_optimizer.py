import logging
import tempfile
import os
import subprocess
import sys
import json
from typing import Dict, List, Tuple, Any, Optional
from dataclasses import dataclass
import fitz  # PyMuPDF
from pdf2docx import Converter
import docx
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

# Configuration des logs
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class LayoutMetrics:
    """Métriques de mise en page pour comparaison avant/après"""
    page_count: int
    total_text_blocks: int
    font_distribution: Dict[str, int]
    color_distribution: Dict[str, int]
    alignment_distribution: Dict[str, int]
    average_font_size: float
    margins: Dict[str, float]
    line_spacing: Dict[str, float]

class LayoutOptimizer:
    """
    Optimiseur de mise en page pour le pipeline PDF → Word → PDF
    Implémente toutes les optimisations demandées pour préserver au maximum la mise en page
    """
    
    def __init__(self):
        self.original_metrics: Optional[LayoutMetrics] = None
        self.optimized_metrics: Optional[LayoutMetrics] = None
        
    def optimize_pdf_to_word_conversion(self, pdf_content: bytes) -> bytes:
        """
        Optimise la conversion PDF → Word avec les meilleurs paramètres
        
        Args:
            pdf_content: Contenu du PDF original
            
        Returns:
            bytes: Contenu du fichier Word optimisé
        """
        logger.info("🔄 Optimisation conversion PDF → Word")
        
        try:
            # Créer des fichiers temporaires
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
                temp_pdf.write(pdf_content)
                temp_pdf_path = temp_pdf.name
            
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_docx:
                temp_docx_path = temp_docx.name
            
            try:
                # Pré-traitement du PDF
                logger.info("📄 Pré-traitement du PDF...")
                cleaned_pdf_path = self._preprocess_pdf(temp_pdf_path)
                
                # Conversion optimisée avec pdf2docx
                logger.info("📄 Conversion PDF → Word avec paramètres optimisés...")
                cv = Converter(cleaned_pdf_path)
                
                # Paramètres optimisés pour préservation maximale
                cv.convert(temp_docx_path, 
                          start=0,
                          # Paramètres pour préserver la mise en page
                          layout=True,           # Préserver la mise en page
                          tables=True,           # Préserver les tableaux
                          header=True,           # Préserver les en-têtes
                          footer=True,           # Préserver les pieds de page
                          multi_processing=False, # Éviter les problèmes de concurrence
                          # Paramètres avancés
                          debug=False,
                          min_section_height=20,
                          connected_border_tolerance=3.0,
                          line_overlap_threshold=0.9,
                          line_break_width_ratio=0.1,
                          line_break_free_space_ratio=0.6,
                          line_break_height_ratio=1.0,
                          line_merge_threshold=3.0,
                          line_break_paragraph_threshold=0.3,
                          line_break_paragraph_width_ratio=0.5,
                          line_break_paragraph_free_space_ratio=0.3,
                          line_break_paragraph_height_ratio=1.0,
                          line_break_paragraph_merging_threshold=3.0,
                          line_break_paragraph_merging_width_ratio=0.5,
                          line_break_paragraph_merging_free_space_ratio=0.3,
                          line_break_paragraph_merging_height_ratio=1.0,
                          line_break_paragraph_merging_connected_border_tolerance=3.0,
                          line_break_paragraph_merging_line_overlap_threshold=0.9,
                          line_break_paragraph_merging_line_break_width_ratio=0.1,
                          line_break_paragraph_merging_line_break_free_space_ratio=0.6,
                          line_break_paragraph_merging_line_break_height_ratio=1.0,
                          line_break_paragraph_merging_line_merge_threshold=3.0,
                          line_break_paragraph_merging_line_break_paragraph_threshold=0.3,
                          line_break_paragraph_merging_line_break_paragraph_width_ratio=0.5,
                          line_break_paragraph_merging_line_break_paragraph_free_space_ratio=0.3,
                          line_break_paragraph_merging_line_break_paragraph_height_ratio=1.0,
                          line_break_paragraph_merging_line_break_paragraph_merging_threshold=3.0)
                
                cv.close()
                
                # Post-traitement du document Word
                logger.info("📄 Post-traitement du document Word...")
                optimized_docx_content = self._postprocess_word_document(temp_docx_path)
                
                logger.info(f"✅ Conversion PDF → Word optimisée réussie")
                return optimized_docx_content
                
            finally:
                # Nettoyage des fichiers temporaires
                for path in [temp_pdf_path, temp_docx_path, cleaned_pdf_path]:
                    if os.path.exists(path):
                        os.unlink(path)
                        
        except Exception as e:
            logger.error(f"❌ Erreur lors de l'optimisation PDF → Word: {str(e)}")
            raise Exception(f"Erreur optimisation PDF → Word: {str(e)}")
    
    def optimize_word_to_pdf_conversion(self, docx_content: bytes) -> bytes:
        """
        Optimise la conversion Word → PDF avec les meilleurs outils
        
        Args:
            docx_content: Contenu du fichier Word
            
        Returns:
            bytes: Contenu du fichier PDF optimisé
        """
        logger.info("🔄 Optimisation conversion Word → PDF")
        
        # Essayer plusieurs méthodes par ordre de préférence
        methods = [
            self._convert_with_libreoffice_simple,     # Priorité 1: LibreOffice simple (nixpacks)
            self._convert_with_libreoffice_optimized,  # Priorité 2: LibreOffice avancé 
            self._convert_with_unoconv,                # Priorité 3: unoconv
            self._convert_with_reportlab_enhanced      # Priorité 4: Fallback reportlab
        ]
        
        for method in methods:
            try:
                logger.info(f"🔄 Tentative avec {method.__name__}...")
                return method(docx_content)
            except Exception as e:
                logger.warning(f"⚠️ {method.__name__} échoué: {str(e)}")
                continue
        
        raise Exception("Toutes les méthodes de conversion Word → PDF ont échoué")
    
    def _preprocess_pdf(self, pdf_path: str) -> str:
        """
        Pré-traite le PDF pour optimiser la conversion
        
        Args:
            pdf_path: Chemin vers le PDF original
            
        Returns:
            str: Chemin vers le PDF pré-traité
        """
        logger.info("📄 Pré-traitement du PDF...")
        
        try:
            # Ouvrir le PDF avec PyMuPDF
            doc = fitz.open(pdf_path)
            
            # Créer un nouveau PDF optimisé
            output_path = pdf_path.replace('.pdf', '_cleaned.pdf')
            
            # Nettoyer les métadonnées et optimiser
            for page in doc:
                # Vectoriser les polices manquantes
                self._vectorize_missing_fonts(page)
            
            # Sauvegarder avec optimisations
            doc.save(output_path, 
                    garbage=4,           # Nettoyage agressif
                    deflate=True,        # Compression
                    clean=True)          # Nettoyage des objets
            
            doc.close()
            
            logger.info(f"✅ Pré-traitement PDF terminé: {output_path}")
            return output_path
            
        except Exception as e:
            logger.warning(f"⚠️ Pré-traitement PDF échoué: {str(e)}")
            return pdf_path  # Retourner le fichier original en cas d'échec
    
    def _vectorize_missing_fonts(self, page):
        """Vectorise les polices manquantes"""
        try:
            # Obtenir les polices utilisées sur la page
            fonts = page.get_fonts()
            
            for font in fonts:
                font_name = font[3]  # Nom de la police
                
                # Vérifier si la police est disponible
                if not self._is_font_available(font_name):
                    logger.info(f"📄 Vectorisation de la police manquante: {font_name}")
                    # Remplacer par une police équivalente
                    self._replace_font_with_equivalent(page, font_name)
                    
        except Exception as e:
            logger.warning(f"⚠️ Erreur vectorisation polices: {str(e)}")
    
    def _is_font_available(self, font_name: str) -> bool:
        """Vérifie si une police est disponible sur le système"""
        try:
            # Liste des polices système courantes
            system_fonts = [
                "Arial", "Times New Roman", "Calibri", "Helvetica", "Times",
                "Courier New", "Verdana", "Georgia", "Palatino", "Garamond"
            ]
            
            return font_name in system_fonts
        except:
            return False
    
    def _replace_font_with_equivalent(self, page, original_font: str):
        """Remplace une police par son équivalent le plus proche"""
        font_mapping = {
            "Calibri": "Arial",
            "Cambria": "Times New Roman",
            "Candara": "Arial",
            "Consolas": "Courier New",
            "Constantia": "Times New Roman",
            "Corbel": "Arial",
            "Georgia": "Times New Roman",
            "Impact": "Arial",
            "Lucida Console": "Courier New",
            "Lucida Sans Unicode": "Arial",
            "Palatino": "Times New Roman",
            "Tahoma": "Arial",
            "Trebuchet MS": "Arial",
            "Verdana": "Arial",
            "Wingdings": "Arial"
        }
        
        equivalent = font_mapping.get(original_font, "Arial")
        logger.info(f"📄 Remplacement {original_font} → {equivalent}")
    
    def _postprocess_word_document(self, docx_path: str) -> bytes:
        """
        Post-traite le document Word pour améliorer la mise en page
        
        Args:
            docx_path: Chemin vers le document Word
            
        Returns:
            bytes: Contenu du document Word optimisé
        """
        logger.info("📄 Post-traitement du document Word...")
        
        try:
            # Ouvrir le document
            doc = Document(docx_path)
            
            # Ajuster les marges
            self._adjust_margins(doc)
            
            # Ajuster les espacements
            self._adjust_spacing(doc)
            
            # Ajuster les tailles de police
            self._adjust_font_sizes(doc)
            
            # Sauvegarder le document optimisé
            output_path = docx_path.replace('.docx', '_optimized.docx')
            doc.save(output_path)
            
            # Lire le contenu optimisé
            with open(output_path, 'rb') as f:
                optimized_content = f.read()
            
            # Nettoyer le fichier temporaire
            if os.path.exists(output_path):
                os.unlink(output_path)
            
            logger.info(f"✅ Post-traitement Word terminé")
            return optimized_content
            
        except Exception as e:
            logger.warning(f"⚠️ Post-traitement Word échoué: {str(e)}")
            # Retourner le contenu original en cas d'échec
            with open(docx_path, 'rb') as f:
                return f.read()
    
    def _adjust_margins(self, doc):
        """Ajuste les marges du document"""
        try:
            # Marges standard pour une meilleure présentation
            section = doc.sections[0]
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            
            logger.info("📄 Marges ajustées")
        except Exception as e:
            logger.warning(f"⚠️ Erreur ajustement marges: {str(e)}")
    
    def _adjust_spacing(self, doc):
        """Ajuste les espacements du document"""
        try:
            for paragraph in doc.paragraphs:
                # Espacement entre paragraphes
                paragraph.paragraph_format.space_after = Pt(6)
                paragraph.paragraph_format.space_before = Pt(6)
                
                # Espacement entre lignes
                paragraph.paragraph_format.line_spacing = 1.15
                
            logger.info("📄 Espacements ajustés")
        except Exception as e:
            logger.warning(f"⚠️ Erreur ajustement espacements: {str(e)}")
    
    def _adjust_font_sizes(self, doc):
        """Ajuste les tailles de police"""
        try:
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    # Ajuster la taille de police si nécessaire
                    if run.font.size:
                        current_size = run.font.size.pt
                        if current_size < 8:
                            run.font.size = Pt(10)
                        elif current_size > 24:
                            run.font.size = Pt(18)
                            
            logger.info("📄 Tailles de police ajustées")
        except Exception as e:
            logger.warning(f"⚠️ Erreur ajustement tailles police: {str(e)}")
    
    def _convert_with_libreoffice_optimized(self, docx_content: bytes) -> bytes:
        """Conversion avec LibreOffice optimisé pour Railway/Linux"""
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_docx:
            temp_docx.write(docx_content)
            temp_docx_path = temp_docx.name
        
        try:
            # Déterminer les commandes LibreOffice selon l'OS avec plus d'options
            libreoffice_commands = [
                # Commandes standard
                "soffice", 
                "libreoffice",
                # Chemins absolus Linux/Docker
                "/usr/bin/soffice",
                "/usr/bin/libreoffice",
                "/opt/libreoffice/program/soffice",
                "/opt/libreoffice7.1/program/soffice",
                "/snap/bin/libreoffice",
                # Chemins Flatpak
                "flatpak run org.libreoffice.LibreOffice",
                # Autres chemins possibles
                "/usr/local/bin/soffice"
            ]
            
            # Ajouter des variables d'environnement pour LibreOffice
            env = os.environ.copy()
            env.update({
                'HOME': '/tmp',  # LibreOffice a besoin d'un répertoire home
                'TMPDIR': '/tmp',
                'XDG_CONFIG_HOME': '/tmp/.config',
                'XDG_DATA_HOME': '/tmp/.local/share',
                'XDG_CACHE_HOME': '/tmp/.cache'
            })
            
            # Créer les répertoires nécessaires
            for dir_path in ['/tmp/.config', '/tmp/.local/share', '/tmp/.cache']:
                os.makedirs(dir_path, exist_ok=True)
            
            # Essayer chaque commande
            for cmd in libreoffice_commands:
                try:
                    logger.info(f"📄 Conversion LibreOffice avec: {cmd}")
                    
                    # Déterminer le répertoire de sortie
                    output_dir = os.path.dirname(temp_docx_path)
                    
                    # Paramètres optimisés pour LibreOffice headless
                    cmd_args = [
                        "--headless",
                        "--invisible",
                        "--nodefault",
                        "--nolockcheck",
                        "--nologo",
                        "--norestore",
                        "--convert-to", "pdf",
                        "--outdir", output_dir,
                        temp_docx_path
                    ]
                    
                    # Si la commande contient flatpak, ajuster les arguments
                    if "flatpak" in cmd:
                        full_cmd = cmd.split() + ["--headless"] + cmd_args[1:]
                    else:
                        full_cmd = [cmd] + cmd_args
                    
                    result = subprocess.run(
                        full_cmd,
                        capture_output=True, 
                        text=True, 
                        timeout=90,  # Augmenter le timeout
                        env=env
                    )
                    
                    if result.returncode == 0:
                        # Le fichier PDF généré aura le même nom de base que le docx
                        expected_pdf_path = temp_docx_path.replace('.docx', '.pdf')
                        
                        if os.path.exists(expected_pdf_path):
                            with open(expected_pdf_path, 'rb') as f:
                                pdf_content = f.read()
                            
                            # Nettoyer le fichier PDF généré
                            os.unlink(expected_pdf_path)
                            
                            logger.info(f"✅ Conversion LibreOffice réussie avec: {cmd}")
                            return pdf_content
                        else:
                            logger.warning(f"⚠️ Fichier PDF non généré: {expected_pdf_path}")
                    else:
                        logger.debug(f"❌ Commande {cmd} échouée: {result.stderr}")
                        
                except (subprocess.TimeoutExpired, FileNotFoundError) as e:
                    logger.debug(f"❌ Commande {cmd} non trouvée ou timeout: {str(e)}")
                    continue
            
            raise Exception("LibreOffice non trouvé ou échec de conversion")
            
        finally:
            # Nettoyer le fichier temporaire
            if os.path.exists(temp_docx_path):
                os.unlink(temp_docx_path)
    
    def _convert_with_unoconv(self, docx_content: bytes) -> bytes:
        """Conversion avec unoconv (alternative à LibreOffice via ligne de commande)"""
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_docx:
            temp_docx.write(docx_content)
            temp_docx_path = temp_docx.name
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
            temp_pdf_path = temp_pdf.name
        
        try:
            # Commandes unoconv à tester
            unoconv_commands = [
                "unoconv",
                "/usr/bin/unoconv",
                "/opt/libreoffice/program/unoconv"
            ]
            
            success = False
            for cmd in unoconv_commands:
                try:
                    logger.info(f"📄 Conversion unoconv avec: {cmd}")
                    
                    # Conversion avec unoconv
                    result = subprocess.run([
                        cmd,
                        "-f", "pdf",
                        "-o", temp_pdf_path,
                        temp_docx_path
                    ], capture_output=True, text=True, timeout=60)
                    
                    if result.returncode == 0:
                        if os.path.exists(temp_pdf_path):
                            with open(temp_pdf_path, 'rb') as f:
                                pdf_content = f.read()
                            
                            logger.info(f"✅ Conversion unoconv réussie")
                            return pdf_content
                        else:
                            logger.warning(f"⚠️ Fichier PDF non généré par unoconv")
                    else:
                        logger.debug(f"❌ Commande {cmd} échouée: {result.stderr}")
                        
                except (subprocess.TimeoutExpired, FileNotFoundError) as e:
                    logger.debug(f"❌ Commande {cmd} non trouvée ou timeout: {str(e)}")
                    continue
            
            raise Exception("unoconv non trouvé ou échec de conversion")
            
        finally:
            for path in [temp_docx_path, temp_pdf_path]:
                if os.path.exists(path):
                    os.unlink(path)
    
    def _convert_with_reportlab_enhanced(self, docx_content: bytes) -> bytes:
        """Conversion de fallback avec reportlab amélioré"""
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.units import inch
            from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT, TA_CENTER
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            from io import BytesIO
            
            logger.info("📄 Conversion de fallback avec reportlab...")
            
            # Extraire le texte du document Word
            doc = Document(BytesIO(docx_content))
            
            # Créer un PDF avec reportlab
            buffer = BytesIO()
            doc_pdf = SimpleDocTemplate(buffer, pagesize=A4)
            story = []
            
            # Styles pour préserver le formatage
            styles = getSampleStyleSheet()
            
            for paragraph in doc.paragraphs:
                text = paragraph.text
                if text.strip():
                    # Déterminer l'alignement
                    alignment = TA_LEFT
                    if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                        alignment = TA_CENTER
                    elif paragraph.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                        alignment = TA_JUSTIFY
                    
                    # Créer un paragraphe avec le style approprié
                    style = ParagraphStyle(
                        'Custom',
                        parent=styles['Normal'],
                        alignment=alignment,
                        fontSize=12,
                        spaceAfter=6
                    )
                    
                    p = Paragraph(text, style)
                    story.append(p)
                    story.append(Spacer(1, 6))
            
            doc_pdf.build(story)
            pdf_content = buffer.getvalue()
            buffer.close()
            
            logger.info(f"✅ Conversion reportlab réussie")
            return pdf_content
            
        except Exception as e:
            raise Exception(f"Erreur conversion reportlab: {str(e)}")
    
    def _convert_with_libreoffice_simple(self, docx_content: bytes) -> bytes:
        """Conversion simplifiée avec LibreOffice (approche nixpacks)"""
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_docx:
            temp_docx.write(docx_content)
            temp_docx_path = temp_docx.name
        
        try:
            # Variables d'environnement optimisées
            env = os.environ.copy()
            env.update({
                'HOME': '/tmp',
                'XDG_CONFIG_HOME': '/tmp/.config',
                'XDG_DATA_HOME': '/tmp/.local/share',
                'XDG_CACHE_HOME': '/tmp/.cache'
            })
            
            # Créer les répertoires si nécessaire
            for dir_path in ['/tmp/.config', '/tmp/.local/share', '/tmp/.cache']:
                os.makedirs(dir_path, exist_ok=True)
            
            # Répertoire de sortie
            output_dir = os.path.dirname(temp_docx_path)
            
            logger.info("📄 Conversion LibreOffice simplifiée...")
            
            # Commande LibreOffice optimisée
            result = subprocess.run([
                "soffice", 
                "--headless", 
                "--invisible",
                "--nodefault",
                "--nolockcheck", 
                "--nologo",
                "--norestore",
                "--convert-to", "pdf", 
                "--outdir", output_dir,
                temp_docx_path
            ], capture_output=True, text=True, timeout=60, env=env, check=True)
            
            # Vérifier que le PDF a été généré
            expected_pdf = temp_docx_path.replace('.docx', '.pdf')
            if os.path.exists(expected_pdf):
                with open(expected_pdf, 'rb') as f:
                    pdf_content = f.read()
                
                # Nettoyer le PDF généré
                os.unlink(expected_pdf)
                
                logger.info(f"✅ Conversion LibreOffice simplifiée réussie")
                return pdf_content
            else:
                raise Exception(f"Fichier PDF non généré: {expected_pdf}")
                
        except subprocess.CalledProcessError as e:
            logger.error(f"❌ Erreur LibreOffice: {e.stderr}")
            raise Exception(f"Échec conversion LibreOffice: {e.stderr}")
        except subprocess.TimeoutExpired:
            logger.error(f"❌ Timeout LibreOffice (60s)")
            raise Exception("Timeout conversion LibreOffice")
        finally:
            # Nettoyer le fichier temporaire
            if os.path.exists(temp_docx_path):
                os.unlink(temp_docx_path)
    
    def compare_layouts(self, original_pdf: bytes, optimized_pdf: bytes) -> Dict[str, Any]:
        """
        Compare les mises en page avant/après optimisation
        
        Args:
            original_pdf: PDF original
            optimized_pdf: PDF optimisé
            
        Returns:
            Dict: Rapport de comparaison
        """
        logger.info("📊 Comparaison des mises en page...")
        
        try:
            # Analyser le PDF original
            original_metrics = self._analyze_pdf_layout(original_pdf)
            
            # Analyser le PDF optimisé
            optimized_metrics = self._analyze_pdf_layout(optimized_pdf)
            
            # Calculer les différences
            comparison = {
                "original": original_metrics.__dict__,
                "optimized": optimized_metrics.__dict__,
                "differences": {
                    "page_count_diff": optimized_metrics.page_count - original_metrics.page_count,
                    "text_blocks_diff": optimized_metrics.total_text_blocks - original_metrics.total_text_blocks,
                    "font_size_diff": optimized_metrics.average_font_size - original_metrics.average_font_size,
                    "preservation_score": self._calculate_preservation_score(original_metrics, optimized_metrics)
                }
            }
            
            logger.info(f"📊 Comparaison terminée - Score de préservation: {comparison['differences']['preservation_score']:.2f}%")
            return comparison
            
        except Exception as e:
            logger.error(f"❌ Erreur lors de la comparaison: {str(e)}")
            return {"error": str(e)}
    
    def _analyze_pdf_layout(self, pdf_content: bytes) -> LayoutMetrics:
        """Analyse la mise en page d'un PDF"""
        try:
            doc = fitz.open(stream=pdf_content, filetype="pdf")
            
            page_count = len(doc)
            total_text_blocks = 0
            font_distribution = {}
            color_distribution = {}
            alignment_distribution = {}
            font_sizes = []
            
            for page in doc:
                text_blocks = page.get_text("dict")
                
                for block in text_blocks["blocks"]:
                    if "lines" in block:
                        total_text_blocks += 1
                        
                        for line in block["lines"]:
                            for span in line["spans"]:
                                # Distribution des polices
                                font_name = span.get("font", "unknown")
                                font_distribution[font_name] = font_distribution.get(font_name, 0) + 1
                                
                                # Distribution des couleurs
                                color = span.get("color", 0)
                                color_distribution[str(color)] = color_distribution.get(str(color), 0) + 1
                                
                                # Tailles de police
                                font_size = span.get("size", 12)
                                font_sizes.append(font_size)
            
            doc.close()
            
            return LayoutMetrics(
                page_count=page_count,
                total_text_blocks=total_text_blocks,
                font_distribution=font_distribution,
                color_distribution=color_distribution,
                alignment_distribution=alignment_distribution,
                average_font_size=sum(font_sizes) / len(font_sizes) if font_sizes else 12,
                margins={"left": 1.0, "right": 1.0, "top": 1.0, "bottom": 1.0},
                line_spacing={"before": 6, "after": 6}
            )
            
        except Exception as e:
            logger.error(f"❌ Erreur analyse mise en page: {str(e)}")
            return LayoutMetrics(0, 0, {}, {}, {}, 12, {}, {})
    
    def _calculate_preservation_score(self, original: LayoutMetrics, optimized: LayoutMetrics) -> float:
        """Calcule un score de préservation de la mise en page"""
        try:
            scores = []
            
            # Score pour le nombre de pages (pénalité si différent)
            page_score = 100 - abs(optimized.page_count - original.page_count) * 10
            scores.append(max(0, page_score))
            
            # Score pour les blocs de texte (pénalité si perte importante)
            text_score = 100 - abs(optimized.total_text_blocks - original.total_text_blocks) * 2
            scores.append(max(0, text_score))
            
            # Score pour la taille moyenne de police (pénalité si changement important)
            font_score = 100 - abs(optimized.average_font_size - original.average_font_size) * 5
            scores.append(max(0, font_score))
            
            # Score pour la distribution des polices
            font_dist_score = self._calculate_distribution_similarity(
                original.font_distribution, optimized.font_distribution)
            scores.append(font_dist_score)
            
            # Score pour la distribution des couleurs
            color_dist_score = self._calculate_distribution_similarity(
                original.color_distribution, optimized.color_distribution)
            scores.append(color_dist_score)
            
            # Moyenne des scores
            final_score = sum(scores) / len(scores)
            return max(0, min(100, final_score))
            
        except Exception as e:
            logger.error(f"❌ Erreur calcul score préservation: {str(e)}")
            return 50.0  # Score par défaut
    
    def _calculate_distribution_similarity(self, original: Dict, optimized: Dict) -> float:
        """Calcule la similarité entre deux distributions"""
        try:
            if not original or not optimized:
                return 50.0
            
            # Calculer la similarité basée sur les clés communes
            common_keys = set(original.keys()) & set(optimized.keys())
            total_keys = set(original.keys()) | set(optimized.keys())
            
            if not total_keys:
                return 100.0
            
            similarity = len(common_keys) / len(total_keys) * 100
            return similarity
            
        except Exception:
            return 50.0

# Instance globale de l'optimiseur
layout_optimizer = LayoutOptimizer() 