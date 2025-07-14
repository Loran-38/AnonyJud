import re
import json
from typing import Dict, List, Tuple, Any

# Imports pour le support PDF
try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT, TA_CENTER
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from io import BytesIO
    import logging
    PDF_SUPPORT = True
    print("‚úì Support PDF activ√© avec reportlab")
except ImportError as e:
    PDF_SUPPORT = False
    print(f"‚ö† Support PDF non disponible: {e}")
    print("Pour activer le support PDF, installez: pip install reportlab")

# Imports pour le pipeline PDF ‚Üí Word ‚Üí PDF avec pr√©servation de mise en page
try:
    from pdf2docx import Converter
    import tempfile
    import os
    import subprocess
    import sys
    PDF_ENHANCED_PIPELINE = True
    print("‚úì Pipeline PDF enhanced activ√© avec pdf2docx")
except ImportError as e:
    PDF_ENHANCED_PIPELINE = False
    print(f"‚ö† Pipeline PDF enhanced non disponible: {e}")
    print("Pour activer le pipeline enhanced, installez: pip install pdf2docx")

# Import docx2pdf comme alternative √† LibreOffice
try:
    from docx2pdf import convert as docx2pdf_convert
    DOCX2PDF_AVAILABLE = True
    print("‚úì docx2pdf disponible pour conversion Word‚ÜíPDF")
except ImportError as e:
    DOCX2PDF_AVAILABLE = False
    print(f"‚ö† docx2pdf non disponible: {e}")
    print("Pour installer docx2pdf: pip install docx2pdf")


def anonymize_text(text: str, tiers: List[Dict[str, Any]] = []) -> Tuple[str, Dict[str, str]]:
    """
    Anonymise le texte en d√©tectant les entit√©s personnelles et en les rempla√ßant par des balises.
    √âtendu : anonymise aussi les dates, num√©ros de s√©curit√© sociale, num√©ros de carte, r√©f√©rences, etc.
    
    Args:
        text: Le texte √† anonymiser
        tiers: Liste des tiers avec leurs informations personnelles (optionnel)
        
    Returns:
        Tuple contenant (texte_anonymis√©, mapping_des_remplacements)
    """
    # Initialiser le mapping des remplacements
    mapping = {}
    anonymized = text
    
    # --- ANONYMISATION DES DATES ---
    # Formats fran√ßais courants : 15/12/2023, 15-12-2023, 15 d√©cembre 2023, 2023-12-15, etc.
    date_patterns = [
        r'\b(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})\b',  # 15/12/2023 ou 15-12-2023
        r'\b(\d{4}[/-]\d{1,2}[/-]\d{1,2})\b',    # 2023-12-15
        r'\b(\d{1,2} [a-zA-Z√©√ª√Æ]+ \d{4})\b',      # 15 d√©cembre 2023
    ]
    date_count = 1
    for pattern in date_patterns:
        for match in re.finditer(pattern, anonymized):
            date = match.group(0)
            if date not in mapping.values():
                tag = f"DATE{date_count}"
                mapping[tag] = date
                anonymized = anonymized.replace(date, tag)
                date_count += 1

    # --- ANONYMISATION NUM√âROS DE S√âCURIT√â SOCIALE (NIR) ---
    nir_pattern = r'\b(\d{2} ?\d{2} ?\d{2} ?\d{3} ?\d{3} ?\d{2})\b'
    nir_count = 1
    for match in re.finditer(nir_pattern, anonymized):
        nir = match.group(0)
        if nir not in mapping.values():
            tag = f"NIR{nir_count}"
            mapping[tag] = nir
            anonymized = anonymized.replace(nir, tag)
            nir_count += 1

    # --- ANONYMISATION NUM√âROS DE CARTE BANCAIRE ---
    cb_pattern = r'\b(\d{4} ?\d{4} ?\d{4} ?\d{4})\b'
    cb_count = 1
    for match in re.finditer(cb_pattern, anonymized):
        cb = match.group(0)
        if cb not in mapping.values():
            tag = f"CB{cb_count}"
            mapping[tag] = cb
            anonymized = anonymized.replace(cb, tag)
            cb_count += 1

    # --- ANONYMISATION NUM√âROS DE DOSSIER/R√âF√âRENCE ---
    ref_pattern = r'\b([A-Z]{2,5}\d{3,10})\b'
    ref_count = 1
    for match in re.finditer(ref_pattern, anonymized):
        ref = match.group(0)
        if ref not in mapping.values():
            tag = f"REF{ref_count}"
            mapping[tag] = ref
            anonymized = anonymized.replace(ref, tag)
            ref_count += 1

    # --- ANONYMISATION NUM√âROS SIMPLES (hors t√©l√©phone d√©j√† g√©r√©) ---
    simple_num_pattern = r'\b(\d{6,})\b'
    num_count = 1
    for match in re.finditer(simple_num_pattern, anonymized):
        num = match.group(0)
        if num not in mapping.values():
            tag = f"NUM{num_count}"
            mapping[tag] = num
            anonymized = anonymized.replace(num, tag)
            num_count += 1

    # --- ANONYMISATION CLASSIQUE (tiers, emails, etc.) ---
    # Si aucun tiers n'est fourni, on utilise une d√©tection basique
    if not tiers or len(tiers) == 0:
        # Anonymisation basique (sans tiers)
        # Recherche des patterns qui ressemblent √† des noms, adresses, etc.
        
        # Exemple : num√©ros de t√©l√©phone fran√ßais
        phone_pattern = r'(?<!\d)((0|\+33|0033)[1-9](?:[ .-]?[0-9]{2}){4})(?!\d)'
        phone_matches = re.finditer(phone_pattern, anonymized)
        phone_count = 1
        
        for match in phone_matches:
            phone = match.group(0)
            if phone not in mapping.values():
                tag = f"TEL{phone_count}"
                mapping[tag] = phone
                anonymized = anonymized.replace(phone, tag)
                phone_count += 1
        
        # Exemple : adresses email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        email_matches = re.finditer(email_pattern, anonymized)
        email_count = 1
        
        for match in email_matches:
            email = match.group(0)
            if email not in mapping.values():
                tag = f"EMAIL{email_count}"
                mapping[tag] = email
                anonymized = anonymized.replace(email, tag)
                email_count += 1
    
    else:
        # Anonymisation avanc√©e avec les tiers fournis
        # Parcourir chaque tier et anonymiser ses informations
        for tier_index, tier in enumerate(tiers):
            # Utiliser le num√©ro fixe du tiers ou fallback sur l'index + 1
            tier_number = tier.get("numero", tier_index + 1)
            # Compteur local pour les champs personnalis√©s de ce tiers
            perso_count = 1
            # Traiter le nom
            if tier.get("nom"):
                nom = tier["nom"].strip()
                if nom and len(nom) > 1:  # √âviter les noms trop courts
                    tag = f"NOM{tier_number}"
                    mapping[tag] = nom
                    
                    # Remplacer toutes les occurrences (insensible √† la casse)
                    pattern = re.compile(re.escape(nom), re.IGNORECASE)
                    anonymized = pattern.sub(tag, anonymized)
                    
                    # Variantes: nom en majuscules, minuscules
                    if nom.upper() != nom:
                        anonymized = anonymized.replace(nom.upper(), tag)
                    if nom.lower() != nom:
                        anonymized = anonymized.replace(nom.lower(), tag)
            
            # Traiter le pr√©nom
            if tier.get("prenom"):
                prenom = tier["prenom"].strip()
                if prenom and len(prenom) > 1:
                    tag = f"PRENOM{tier_number}"
                    mapping[tag] = prenom
                    
                    pattern = re.compile(re.escape(prenom), re.IGNORECASE)
                    anonymized = pattern.sub(tag, anonymized)
                    
                    if prenom.upper() != prenom:
                        anonymized = anonymized.replace(prenom.upper(), tag)
                    if prenom.lower() != prenom:
                        anonymized = anonymized.replace(prenom.lower(), tag)
                        
            # Traiter les composants de l'adresse
            # Num√©ro de voie
            if tier.get("adresse_numero"):
                numero = tier["adresse_numero"].strip()
                if numero and len(numero) > 0:
                    tag = f"NUMERO{tier_number}"
                    mapping[tag] = numero
                    
                    # Remplacer toutes les occurrences
                    pattern = re.compile(re.escape(numero), re.IGNORECASE)
                    anonymized = pattern.sub(tag, anonymized)
            
            # Voie (rue, avenue, etc.)
            if tier.get("adresse_voie"):
                voie = tier["adresse_voie"].strip()
                if voie and len(voie) > 2:
                    tag = f"VOIE{tier_number}"
                    mapping[tag] = voie
                    
                    # Remplacer toutes les occurrences (insensible √† la casse)
                    pattern = re.compile(re.escape(voie), re.IGNORECASE)
                    anonymized = pattern.sub(tag, anonymized)
                    
                    # Variantes: en majuscules, minuscules
                    if voie.upper() != voie:
                        anonymized = anonymized.replace(voie.upper(), tag)
                    if voie.lower() != voie:
                        anonymized = anonymized.replace(voie.lower(), tag)
            
            # Code postal
            if tier.get("adresse_code_postal"):
                code_postal = tier["adresse_code_postal"].strip()
                if code_postal and len(code_postal) > 0:
                    tag = f"CODEPOSTAL{tier_number}"
                    mapping[tag] = code_postal
                    
                    # Remplacer toutes les occurrences
                    anonymized = anonymized.replace(code_postal, tag)
            
            # Ville
            if tier.get("adresse_ville"):
                ville = tier["adresse_ville"].strip()
                if ville and len(ville) > 1:
                    tag = f"VILLE{tier_number}"
                    mapping[tag] = ville
                    
                    # Remplacer toutes les occurrences (insensible √† la casse)
                    pattern = re.compile(re.escape(ville), re.IGNORECASE)
                    anonymized = pattern.sub(tag, anonymized)
                    
                    # Variantes: en majuscules, minuscules
                    if ville.upper() != ville:
                        anonymized = anonymized.replace(ville.upper(), tag)
                    if ville.lower() != ville:
                        anonymized = anonymized.replace(ville.lower(), tag)
            
            # Traiter l'adresse compl√®te (pour compatibilit√© avec l'ancien format)
            if tier.get("adresse"):
                adresse = tier["adresse"].strip()
                if adresse and len(adresse) > 5:  # √âviter les adresses trop courtes
                    tag = f"ADRESSE{tier_number}"
                    mapping[tag] = adresse
                    
                    # Pour l'adresse, on fait un remplacement exact car elle peut contenir des caract√®res sp√©ciaux
                    anonymized = anonymized.replace(adresse, tag)
            
            # Traiter le t√©l√©phone fixe
            if tier.get("telephone"):
                tel = tier["telephone"].strip()
                if tel and len(tel) > 5:
                    tag = f"TEL{tier_number}"
                    mapping[tag] = tel
                    
                    # Normaliser le format du t√©l√©phone pour la recherche
                    tel_clean = re.sub(r'[. -]', '', tel)
                    pattern = re.compile(r'[. -]'.join(re.escape(c) for c in tel_clean))
                    anonymized = pattern.sub(tag, anonymized)
                    
                    # Remplacer aussi le format sans espaces/points
                    anonymized = anonymized.replace(tel_clean, tag)
            
            # Traiter le portable
            if tier.get("portable"):
                portable = tier["portable"].strip()
                if portable and len(portable) > 5:
                    tag = f"PORTABLE{tier_number}"
                    mapping[tag] = portable
                    
                    # M√™me approche que pour le t√©l√©phone
                    portable_clean = re.sub(r'[. -]', '', portable)
                    pattern = re.compile(r'[. -]'.join(re.escape(c) for c in portable_clean))
                    anonymized = pattern.sub(tag, anonymized)
                    
                    anonymized = anonymized.replace(portable_clean, tag)
            
            # Traiter l'email
            if tier.get("email"):
                email = tier["email"].strip()
                if email and '@' in email:
                    tag = f"EMAIL{tier_number}"
                    mapping[tag] = email
                    
                    anonymized = anonymized.replace(email, tag)
            
            # Traiter la soci√©t√©
            if tier.get("societe"):
                societe = tier["societe"].strip()
                if societe and len(societe) > 1:
                    tag = f"SOCIETE{tier_number}"
                    mapping[tag] = societe
                    
                    pattern = re.compile(re.escape(societe), re.IGNORECASE)
                    anonymized = pattern.sub(tag, anonymized)
                    
                    if societe.upper() != societe:
                        anonymized = anonymized.replace(societe.upper(), tag)
                    if societe.lower() != societe:
                        anonymized = anonymized.replace(societe.lower(), tag)
                        
            # Traiter les champs personnalis√©s (nouveau format)
            if tier.get("customFields") and isinstance(tier["customFields"], list):
                for custom_field in tier["customFields"]:
                    if isinstance(custom_field, dict):
                        champ_value = custom_field.get("value")
                        champ_label = custom_field.get("label")
                        
                        if champ_value and isinstance(champ_value, str):
                            champ_value = champ_value.strip()
                            if champ_value and len(champ_value) > 1:
                                # Utiliser le label personnalis√© s'il existe, sinon "PERSO"
                                label_base = "PERSO"
                                if champ_label and isinstance(champ_label, str):
                                    label_perso = champ_label.strip()
                                    # Nettoyer le label pour qu'il soit utilisable comme balise
                                    if label_perso:
                                        label_base = re.sub(r'[^A-Za-z]', '', label_perso.upper())
                                        if not label_base:  # Si le label ne contient pas de lettres
                                            label_base = "PERSO"
                                
                                tag = f"{label_base}{tier_number}"
                                mapping[tag] = champ_value
                                
                                # Remplacer toutes les occurrences (insensible √† la casse)
                                pattern = re.compile(re.escape(champ_value), re.IGNORECASE)
                                anonymized = pattern.sub(tag, anonymized)
                                
                                # Variantes: en majuscules, minuscules
                                if champ_value.upper() != champ_value:
                                    anonymized = anonymized.replace(champ_value.upper(), tag)
                                if champ_value.lower() != champ_value:
                                    anonymized = anonymized.replace(champ_value.lower(), tag)
            
            # Traiter le champ personnalis√© (ancien format pour compatibilit√©)
            if tier.get("champPerso"):
                champ_perso = tier["champPerso"]
                if champ_perso and isinstance(champ_perso, str):
                    champ_perso = champ_perso.strip()
                    if champ_perso and len(champ_perso) > 1:
                        # Utiliser le label personnalis√© s'il existe, sinon "PERSO"
                        label_base = "PERSO"
                        if tier.get("labelChampPerso") and isinstance(tier["labelChampPerso"], str):
                            label_perso = tier["labelChampPerso"].strip()
                            # Nettoyer le label pour qu'il soit utilisable comme balise
                            if label_perso:
                                label_base = re.sub(r'[^A-Za-z]', '', label_perso.upper())
                                if not label_base:  # Si le label ne contient pas de lettres
                                    label_base = "PERSO"
                        
                        tag = f"{label_base}{tier_number}"
                        mapping[tag] = champ_perso
                        
                        # Remplacer toutes les occurrences (insensible √† la casse)
                        pattern = re.compile(re.escape(champ_perso), re.IGNORECASE)
                        anonymized = pattern.sub(tag, anonymized)
                        
                        # Variantes: en majuscules, minuscules
                        if champ_perso.upper() != champ_perso:
                            anonymized = anonymized.replace(champ_perso.upper(), tag)
                        if champ_perso.lower() != champ_perso:
                            anonymized = anonymized.replace(champ_perso.lower(), tag)
    
    return anonymized, mapping

def create_pdf_from_text(text: str, title: str = "Document anonymis√©") -> bytes:
    """
    Cr√©e un fichier PDF √† partir d'un texte donn√©.
    
    Args:
        text: Le texte √† convertir en PDF
        title: Le titre du document PDF
        
    Returns:
        bytes: Le contenu du fichier PDF
    """
    if not PDF_SUPPORT:
        raise ImportError("Support PDF non disponible. Installez reportlab: pip install reportlab")
    
    logging.info(f"Cr√©ation PDF - Titre: {title}, Taille texte: {len(text)} caract√®res")
    
    # Cr√©er un buffer en m√©moire pour le PDF
    buffer = BytesIO()
    
    # Cr√©er le document PDF
    doc = SimpleDocTemplate(buffer, pagesize=A4, 
                          rightMargin=72, leftMargin=72,
                          topMargin=72, bottomMargin=18)
    
    # Styles
    styles = getSampleStyleSheet()
    title_style = styles['Title']
    normal_style = styles['Normal']
    normal_style.alignment = TA_JUSTIFY
    
    # Contenu du PDF
    story = []
    
    # Titre
    if title:
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 12))
    
    # Diviser le texte en paragraphes
    paragraphs = text.split('\n\n')
    
    for paragraph in paragraphs:
        if paragraph.strip():
            # Nettoyer le paragraphe
            clean_paragraph = paragraph.strip().replace('\n', ' ')
            # √âchapper les caract√®res sp√©ciaux pour reportlab
            clean_paragraph = clean_paragraph.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            story.append(Paragraph(clean_paragraph, normal_style))
            story.append(Spacer(1, 12))
    
    # Construire le PDF
    doc.build(story)
    
    # R√©cup√©rer le contenu du buffer
    pdf_content = buffer.getvalue()
    buffer.close()
    
    logging.info(f"PDF cr√©√© avec succ√®s - Taille: {len(pdf_content)} bytes")
    return pdf_content

def anonymize_pdf_file(file_content: bytes, tiers: List[Dict[str, Any]] = []) -> Tuple[bytes, Dict[str, str]]:
    """
    Anonymise un fichier PDF en extrayant le texte, l'anonymisant et cr√©ant un nouveau PDF.
    
    Args:
        file_content: Le contenu du fichier PDF original
        tiers: Liste des tiers avec leurs informations personnelles
        
    Returns:
        Tuple contenant (contenu_pdf_anonymis√©, mapping_des_remplacements)
    """
    if not PDF_SUPPORT:
        raise ImportError("Support PDF non disponible. Installez reportlab: pip install reportlab")
    
    logging.info("D√©but anonymisation PDF")
    
    try:
        # Importer la fonction d'extraction de texte PDF
        from .utils import extract_text_from_pdf
        
        # Extraire le texte du PDF original
        extracted_text = extract_text_from_pdf(file_content)
        
        logging.info(f"Texte extrait du PDF - Taille: {len(extracted_text)} caract√®res")
        
        # Anonymiser le texte
        anonymized_text, mapping = anonymize_text(extracted_text, tiers)
        
        logging.info(f"Texte anonymis√© - Nouvelles balises: {len(mapping)}")
        
        # Cr√©er un nouveau PDF avec le texte anonymis√©
        pdf_content = create_pdf_from_text(anonymized_text, "Document PDF Anonymis√©")
        
        logging.info("Anonymisation PDF termin√©e avec succ√®s")
        return pdf_content, mapping
        
    except Exception as e:
        logging.error(f"Erreur lors de l'anonymisation PDF: {str(e)}")
        raise

def deanonymize_pdf_file(file_content: bytes, mapping: Dict[str, str]) -> bytes:
    """
    D√©-anonymise un fichier PDF en rempla√ßant les balises par les valeurs originales.
    
    Args:
        file_content: Le contenu du fichier PDF anonymis√©
        mapping: Dictionnaire de mapping des balises vers les valeurs originales
        
    Returns:
        bytes: Le contenu du fichier PDF d√©-anonymis√©
    """
    if not PDF_SUPPORT:
        raise ImportError("Support PDF non disponible. Installez reportlab: pip install reportlab")
    
    logging.info("D√©but d√©-anonymisation PDF")
    
    try:
        # Importer la fonction d'extraction de texte PDF
        from .utils import extract_text_from_pdf
        
        # Extraire le texte du PDF anonymis√©
        extracted_text = extract_text_from_pdf(file_content)
        
        logging.info(f"Texte extrait du PDF anonymis√© - Taille: {len(extracted_text)} caract√®res")
        
        # D√©-anonymiser le texte
        deanonymized_text = extracted_text
        
        if mapping:
            logging.info(f"Application du mapping - {len(mapping)} remplacements")
            
            for tag, original_value in mapping.items():
                if tag in deanonymized_text:
                    deanonymized_text = deanonymized_text.replace(tag, original_value)
                    logging.debug(f"Remplac√© {tag} par {original_value}")
        
        logging.info("D√©-anonymisation du texte termin√©e")
        
        # Cr√©er un nouveau PDF avec le texte d√©-anonymis√©
        pdf_content = create_pdf_from_text(deanonymized_text, "Document PDF D√©-anonymis√©")
        
        logging.info("D√©-anonymisation PDF termin√©e avec succ√®s")
        return pdf_content
        
    except Exception as e:
        logging.error(f"Erreur lors de la d√©-anonymisation PDF: {str(e)}")
        raise 


# ===== NOUVEAU PIPELINE PDF ‚Üí WORD ‚Üí PDF =====

def convert_pdf_to_word_enhanced(pdf_content: bytes) -> bytes:
    """
    Convertit un PDF en Word en pr√©servant strictement la mise en page avec pdf2docx.
    
    Args:
        pdf_content: Le contenu du fichier PDF
        
    Returns:
        bytes: Le contenu du fichier Word (.docx) g√©n√©r√©
    """
    if not PDF_ENHANCED_PIPELINE:
        raise ImportError("Pipeline PDF enhanced non disponible. Installez pdf2docx: pip install pdf2docx")
    
    logging.info("üîÑ D√©but conversion PDF ‚Üí Word avec pr√©servation de mise en page")
    
    try:
        # Cr√©er des fichiers temporaires pour la conversion
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
            temp_pdf.write(pdf_content)
            temp_pdf_path = temp_pdf.name
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_docx:
            temp_docx_path = temp_docx.name
        
        try:
            # Conversion PDF ‚Üí DOCX avec pdf2docx
            logging.info(f"üìÑ Conversion de {temp_pdf_path} vers {temp_docx_path}")
            
            # Utiliser pdf2docx pour convertir avec pr√©servation maximale
            # Param√®tres optimis√©s pour pr√©server la mise en page
            cv = Converter(temp_pdf_path)
            cv.convert(temp_docx_path, 
                      start=0, 
                      # Param√®tres pour pr√©server au maximum la mise en page
                      multi_processing=False)  # √âviter les probl√®mes de concurrence
            cv.close()
            
            logging.info(f"‚úÖ Conversion pdf2docx termin√©e")
            
            # Lire le fichier Word g√©n√©r√©
            with open(temp_docx_path, 'rb') as f:
                docx_content = f.read()
            
            logging.info(f"‚úÖ Conversion PDF ‚Üí Word r√©ussie. Taille DOCX: {len(docx_content)} bytes")
            return docx_content
            
        finally:
            # Nettoyer les fichiers temporaires
            if os.path.exists(temp_pdf_path):
                os.unlink(temp_pdf_path)
            if os.path.exists(temp_docx_path):
                os.unlink(temp_docx_path)
    
    except Exception as e:
        logging.error(f"‚ùå Erreur lors de la conversion PDF ‚Üí Word: {str(e)}")
        raise Exception(f"Erreur conversion PDF ‚Üí Word: {str(e)}")


def convert_word_to_pdf_enhanced(docx_content: bytes) -> bytes:
    """
    Convertit un fichier Word en PDF en pr√©servant la mise en page.
    Utilise plusieurs m√©thodes par ordre de pr√©f√©rence pour la pr√©servation.
    
    Args:
        docx_content: Le contenu du fichier Word (.docx)
        
    Returns:
        bytes: Le contenu du fichier PDF g√©n√©r√©
    """
    logging.info("üîÑ D√©but conversion Word ‚Üí PDF avec pr√©servation de mise en page")
    
    # M√©thode 1: docx2pdf (Windows, meilleure pr√©servation)
    if DOCX2PDF_AVAILABLE:
        try:
            logging.info("üîÑ Tentative conversion avec docx2pdf...")
            return _convert_word_to_pdf_docx2pdf(docx_content)
        except Exception as e:
            logging.warning(f"‚ö†Ô∏è docx2pdf √©chou√©: {str(e)}")
    
    # M√©thode 2: LibreOffice headless (cross-platform, bonne pr√©servation)
    try:
        logging.info("üîÑ Tentative conversion avec LibreOffice...")
        return _convert_word_to_pdf_libreoffice(docx_content)
    except Exception as e:
        logging.warning(f"‚ö†Ô∏è LibreOffice non disponible: {str(e)}")
        
    # M√©thode 3: Fallback am√©lior√© avec reportlab (pr√©servation partielle)
    try:
        logging.info("üîÑ Fallback: conversion am√©lior√©e avec reportlab...")
        return _convert_word_to_pdf_reportlab_enhanced(docx_content)
    except Exception as e2:
        logging.error(f"‚ùå Toutes les m√©thodes de conversion ont √©chou√©: {str(e2)}")
        raise Exception(f"Impossible de convertir Word ‚Üí PDF: {str(e2)}")


def _convert_word_to_pdf_docx2pdf(docx_content: bytes) -> bytes:
    """
    Convertit Word ‚Üí PDF avec docx2pdf (Windows uniquement, excellente pr√©servation).
    """
    if not DOCX2PDF_AVAILABLE:
        raise ImportError("docx2pdf non disponible")
    
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_docx:
        temp_docx.write(docx_content)
        temp_docx_path = temp_docx.name
    
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
        temp_pdf_path = temp_pdf.name
    
    try:
        # Utiliser docx2pdf pour la conversion
        logging.info(f"üìÑ Conversion docx2pdf: {temp_docx_path} ‚Üí {temp_pdf_path}")
        docx2pdf_convert(temp_docx_path, temp_pdf_path)
        
        # V√©rifier que le fichier PDF a √©t√© cr√©√©
        if not os.path.exists(temp_pdf_path):
            raise Exception(f"Fichier PDF non g√©n√©r√© par docx2pdf: {temp_pdf_path}")
        
        # Lire le PDF g√©n√©r√©
        with open(temp_pdf_path, 'rb') as f:
            pdf_content = f.read()
        
        logging.info(f"‚úÖ Conversion docx2pdf r√©ussie. Taille PDF: {len(pdf_content)} bytes")
        return pdf_content
        
    finally:
        # Nettoyer les fichiers temporaires
        if os.path.exists(temp_docx_path):
            os.unlink(temp_docx_path)
        if os.path.exists(temp_pdf_path):
            os.unlink(temp_pdf_path)


def _convert_word_to_pdf_libreoffice(docx_content: bytes) -> bytes:
    """
    Convertit Word ‚Üí PDF avec LibreOffice en mode headless (pr√©servation maximale).
    """
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_docx:
        temp_docx.write(docx_content)
        temp_docx_path = temp_docx.name
    
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
        temp_pdf_path = temp_pdf.name
    
    try:
        # D√©terminer la commande LibreOffice selon l'OS
        if sys.platform == "win32":
            # Windows
            libreoffice_commands = [
                "soffice",
                "libreoffice",
                r"C:\Program Files\LibreOffice\program\soffice.exe",
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"
            ]
        elif sys.platform == "darwin":
            # macOS
            libreoffice_commands = [
                "/Applications/LibreOffice.app/Contents/MacOS/soffice",
                "soffice",
                "libreoffice"
            ]
        else:
            # Linux
            libreoffice_commands = [
                "soffice",
                "libreoffice",
                "/usr/bin/soffice",
                "/opt/libreoffice7.1/program/soffice"
            ]
        
        # Essayer chaque commande jusqu'√† ce qu'une fonctionne
        success = False
        for cmd in libreoffice_commands:
            try:
                # Commande LibreOffice headless
                result = subprocess.run([
                    cmd,
                    "--headless",
                    "--convert-to", "pdf",
                    "--outdir", os.path.dirname(temp_pdf_path),
                    temp_docx_path
                ], capture_output=True, text=True, timeout=30)
                
                if result.returncode == 0:
                    success = True
                    logging.info(f"‚úÖ LibreOffice conversion r√©ussie avec: {cmd}")
                    break
                else:
                    logging.debug(f"‚ùå Commande {cmd} √©chou√©e: {result.stderr}")
            except (subprocess.TimeoutExpired, FileNotFoundError) as e:
                logging.debug(f"‚ùå Commande {cmd} non trouv√©e ou timeout: {str(e)}")
                continue
        
        if not success:
            raise Exception("LibreOffice non trouv√© ou √©chec de conversion")
        
        # Le fichier PDF g√©n√©r√© aura le m√™me nom de base que le docx
        expected_pdf_path = temp_docx_path.replace('.docx', '.pdf')
        
        if not os.path.exists(expected_pdf_path):
            raise Exception(f"Fichier PDF non g√©n√©r√©: {expected_pdf_path}")
        
        # Lire le PDF g√©n√©r√©
        with open(expected_pdf_path, 'rb') as f:
            pdf_content = f.read()
        
        logging.info(f"‚úÖ Conversion Word ‚Üí PDF r√©ussie. Taille PDF: {len(pdf_content)} bytes")
        return pdf_content
        
    finally:
        # Nettoyer les fichiers temporaires
        if os.path.exists(temp_docx_path):
            os.unlink(temp_docx_path)
        if os.path.exists(temp_pdf_path):
            os.unlink(temp_pdf_path)
        
        # Nettoyer le PDF g√©n√©r√© par LibreOffice
        expected_pdf_path = temp_docx_path.replace('.docx', '.pdf')
        if os.path.exists(expected_pdf_path):
            os.unlink(expected_pdf_path)


def _convert_word_to_pdf_reportlab(docx_content: bytes) -> bytes:
    """
    M√©thode de fallback: extrait le texte du Word et g√©n√®re un PDF avec reportlab.
    """
    if not PDF_SUPPORT:
        raise ImportError("Support PDF non disponible. Installez reportlab: pip install reportlab")
    
    from docx import Document
    
    # Extraire le texte du document Word
    doc = Document(BytesIO(docx_content))
    text = ""
    
    for para in doc.paragraphs:
        text += para.text + "\n"
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text += para.text + "\n"
    
    # G√©n√©rer un PDF avec le texte extrait
    return create_pdf_from_text(text, "Document converti depuis Word")


def _convert_word_to_pdf_reportlab_enhanced(docx_content: bytes) -> bytes:
    """
    M√©thode de fallback am√©lior√©e: extrait le texte du Word avec pr√©servation partielle de la structure.
    """
    if not PDF_SUPPORT:
        raise ImportError("Support PDF non disponible. Installez reportlab: pip install reportlab")
    
    from docx import Document
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT, TA_CENTER
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    
    # Ouvrir le document Word
    doc = Document(BytesIO(docx_content))
    
    # Cr√©er un buffer pour le PDF
    buffer = BytesIO()
    pdf_doc = SimpleDocTemplate(buffer, pagesize=A4, 
                               rightMargin=72, leftMargin=72,
                               topMargin=72, bottomMargin=72)
    
    # Styles
    styles = getSampleStyleSheet()
    normal_style = styles['Normal']
    normal_style.alignment = TA_JUSTIFY
    normal_style.fontSize = 11
    normal_style.leading = 14
    
    # Titre style
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        spaceAfter=12,
        alignment=TA_CENTER
    )
    
    # Sous-titre style
    subtitle_style = ParagraphStyle(
        'CustomSubtitle',
        parent=styles['Heading2'],
        fontSize=14,
        spaceAfter=8,
        alignment=TA_LEFT
    )
    
    story = []
    
    # Traiter chaque paragraphe
    for para in doc.paragraphs:
        if para.text.strip():
            text = para.text.strip()
            
            # √âchapper les caract√®res sp√©ciaux
            text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            
            # D√©tecter les titres (heuristique simple)
            if len(text) < 100 and (text.isupper() or text.startswith('TITRE') or text.startswith('CHAPITRE')):
                story.append(Paragraph(text, title_style))
            elif len(text) < 80 and text.endswith(':'):
                story.append(Paragraph(text, subtitle_style))
            else:
                story.append(Paragraph(text, normal_style))
            
            story.append(Spacer(1, 6))
    
    # Traiter les tableaux
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = ""
                for para in cell.paragraphs:
                    if para.text.strip():
                        cell_text += para.text.strip() + " "
                row_data.append(cell_text.strip())
            if any(row_data):  # Seulement si la ligne n'est pas vide
                table_data.append(row_data)
        
        if table_data:
            # Cr√©er le tableau PDF
            pdf_table = Table(table_data)
            pdf_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(pdf_table)
            story.append(Spacer(1, 12))
    
    # Construire le PDF
    pdf_doc.build(story)
    
    # R√©cup√©rer le contenu
    pdf_content = buffer.getvalue()
    buffer.close()
    
    logging.info(f"‚úÖ Conversion reportlab enhanced r√©ussie. Taille PDF: {len(pdf_content)} bytes")
    return pdf_content


def anonymize_pdf_enhanced_pipeline(pdf_content: bytes, tiers: List[Dict[str, Any]] = []) -> Tuple[bytes, Dict[str, str]]:
    """
    Pipeline complet : PDF ‚Üí Word ‚Üí Anonymisation ‚Üí PDF
    Pr√©serve la mise en page tout au long du processus.
    
    Args:
        pdf_content: Le contenu du fichier PDF original
        tiers: Liste des tiers avec leurs informations personnelles
        
    Returns:
        Tuple contenant (pdf_anonymis√©, mapping_des_remplacements)
    """
    logging.info("üöÄ D√©but du pipeline PDF enhanced: PDF ‚Üí Word ‚Üí Anonymisation ‚Üí PDF")
    logging.info(f"üìä Taille PDF d'entr√©e: {len(pdf_content)} bytes")
    logging.info(f"üë• Nombre de tiers: {len(tiers)}")
    
    try:
        # √âtape 1: PDF ‚Üí Word avec pr√©servation de mise en page
        logging.info("üìÑ √âtape 1/3: Conversion PDF ‚Üí Word")
        docx_content = convert_pdf_to_word_enhanced(pdf_content)
        logging.info(f"‚úÖ PDF ‚Üí Word r√©ussi. Taille DOCX: {len(docx_content)} bytes")
        
        # √âtape 2: Anonymisation du fichier Word (pr√©serve le formatage)
        logging.info("üîí √âtape 2/3: Anonymisation du document Word")
        from .main import anonymize_docx_file  # Import dynamique pour √©viter la circularit√©
        anonymized_docx_content, mapping = anonymize_docx_file(docx_content, tiers)
        logging.info(f"‚úÖ Anonymisation Word r√©ussie. Taille DOCX anonymis√©: {len(anonymized_docx_content)} bytes")
        logging.info(f"üìä Mapping g√©n√©r√©: {len(mapping)} remplacements - {list(mapping.keys())}")
        
        # √âtape 3: Word anonymis√© ‚Üí PDF final
        logging.info("üìÑ √âtape 3/3: Conversion Word anonymis√© ‚Üí PDF")
        final_pdf_content = convert_word_to_pdf_enhanced(anonymized_docx_content)
        logging.info(f"‚úÖ Word ‚Üí PDF r√©ussi. Taille PDF final: {len(final_pdf_content)} bytes")
        
        logging.info(f"üéâ Pipeline PDF enhanced termin√© avec succ√®s")
        logging.info(f"üìä R√©sum√©: PDF({len(pdf_content)}) ‚Üí DOCX({len(docx_content)}) ‚Üí DOCX_ANONYM({len(anonymized_docx_content)}) ‚Üí PDF_FINAL({len(final_pdf_content)})")
        
        return final_pdf_content, mapping
        
    except Exception as e:
        logging.error(f"‚ùå Erreur dans le pipeline PDF enhanced: {str(e)}")
        raise Exception(f"Erreur pipeline PDF enhanced: {str(e)}")


def deanonymize_pdf_enhanced_pipeline(pdf_content: bytes, mapping: Dict[str, str]) -> bytes:
    """
    Pipeline de d√©-anonymisation : PDF ‚Üí Word ‚Üí D√©-anonymisation ‚Üí PDF
    Pr√©serve la mise en page tout au long du processus.
    
    Args:
        pdf_content: Le contenu du fichier PDF anonymis√©
        mapping: Dictionnaire de mapping des balises vers les valeurs originales
        
    Returns:
        bytes: Le contenu du fichier PDF d√©-anonymis√©
    """
    logging.info("üîì D√©but du pipeline PDF d√©-anonymisation enhanced")
    
    try:
        # √âtape 1: PDF anonymis√© ‚Üí Word avec pr√©servation de mise en page
        logging.info("üìÑ √âtape 1/3: Conversion PDF anonymis√© ‚Üí Word")
        docx_content = convert_pdf_to_word_enhanced(pdf_content)
        
        # √âtape 2: D√©-anonymisation du fichier Word (pr√©serve le formatage)
        logging.info("üîì √âtape 2/3: D√©-anonymisation du document Word")
        from .main import deanonymize_docx_file  # Import dynamique pour √©viter la circularit√©
        deanonymized_docx_content = deanonymize_docx_file(docx_content, mapping)
        
        # √âtape 3: Word d√©-anonymis√© ‚Üí PDF final
        logging.info("üìÑ √âtape 3/3: Conversion Word d√©-anonymis√© ‚Üí PDF")
        final_pdf_content = convert_word_to_pdf_enhanced(deanonymized_docx_content)
        
        logging.info(f"‚úÖ Pipeline PDF d√©-anonymisation enhanced termin√© avec succ√®s")
        
        return final_pdf_content
        
    except Exception as e:
        logging.error(f"‚ùå Erreur dans le pipeline PDF d√©-anonymisation enhanced: {str(e)}")
        raise Exception(f"Erreur pipeline PDF d√©-anonymisation enhanced: {str(e)}") 


# ===== ANONYMISATION DIRECTE PDF AVEC PYMUPDF =====

def anonymize_pdf_direct(pdf_content: bytes, tiers: List[Dict[str, Any]] = []) -> Tuple[bytes, Dict[str, str]]:
    """
    Anonymise directement un PDF en rempla√ßant le texte in-place avec PyMuPDF.
    Pr√©serve parfaitement la mise en page, les polices, les couleurs et la structure.
    
    Args:
        pdf_content: Le contenu du fichier PDF original
        tiers: Liste des tiers avec leurs informations personnelles
        
    Returns:
        Tuple contenant (pdf_anonymis√©, mapping_des_remplacements)
    """
    import fitz  # PyMuPDF
    
    logging.info("üöÄ D√©but anonymisation PDF directe avec PyMuPDF")
    logging.info(f"üìä Taille PDF d'entr√©e: {len(pdf_content)} bytes")
    logging.info(f"üë• Nombre de tiers: {len(tiers)}")
    
    try:
        # G√©n√©rer le mapping d'anonymisation
        full_text = ""
        
        # Ouvrir le PDF avec gestion d'erreur MuPDF am√©lior√©e
        try:
            doc = fitz.open(stream=pdf_content, filetype="pdf")
        except Exception as e:
            if "object out of range" in str(e):
                logging.warning(f"‚ö†Ô∏è Erreur MuPDF 'object out of range': {str(e)}")
                # Essayer de r√©cup√©rer le PDF en mode tol√©rant
                try:
                    doc = fitz.open(stream=pdf_content, filetype="pdf")
                    # R√©parer le PDF si possible
                    doc.save(doc.name, garbage=4, deflate=True)
                    doc.close()
                    doc = fitz.open(stream=pdf_content, filetype="pdf")
                except:
                    logging.error(f"‚ùå Impossible de r√©parer le PDF: {str(e)}")
                    raise Exception(f"PDF corrompu ou non support√©: {str(e)}")
            else:
                raise
        
        # Extraire tout le texte pour g√©n√©rer le mapping
        for page in doc:
            try:
                full_text += page.get_text() + "\n"
            except Exception as e:
                logging.warning(f"‚ö†Ô∏è Erreur extraction texte page: {str(e)}")
                continue
        
        doc.close()
        
        # G√©n√©rer le mapping d'anonymisation
        anonymized_text, mapping = anonymize_text(full_text, tiers)
        logging.info(f"üìä Mapping g√©n√©r√©: {len(mapping)} remplacements - {list(mapping.keys())}")
        
        # Inverser le mapping pour avoir {valeur_originale: balise_anonymis√©e}
        reverse_mapping = {v: k for k, v in mapping.items()}
        logging.info(f"üìä Mapping invers√©: {reverse_mapping}")
        
        # Ouvrir le PDF pour modification avec gestion d'erreur am√©lior√©e
        try:
            doc = fitz.open(stream=pdf_content, filetype="pdf")
        except Exception as e:
            if "object out of range" in str(e):
                logging.warning(f"‚ö†Ô∏è Erreur MuPDF lors de l'ouverture pour modification: {str(e)}")
                # Continuer avec le document tel quel
                doc = fitz.open(stream=pdf_content, filetype="pdf")
            else:
                raise
        
        # Parcourir chaque page
        for page_num in range(len(doc)):
            page = doc[page_num]
            logging.info(f"üìÑ Traitement page {page_num + 1}/{len(doc)}")
            
            try:
                # Obtenir tous les blocs de texte de la page
                text_blocks = page.get_text("dict")
                
                # Parcourir chaque bloc de texte avec alignement parfait am√©lior√©
                for block in text_blocks["blocks"]:
                    if "lines" in block:  # Bloc de texte
                        _anonymize_text_block_comprehensive(page, block, reverse_mapping)
                        
            except Exception as e:
                logging.warning(f"‚ö†Ô∏è Erreur traitement page {page_num + 1}: {str(e)}")
                continue
        
        # Sauvegarder le PDF modifi√© avec gestion d'erreur am√©lior√©e
        try:
            anonymized_pdf = doc.tobytes()
        except Exception as e:
            if "object out of range" in str(e):
                logging.warning(f"‚ö†Ô∏è Erreur MuPDF lors de la sauvegarde: {str(e)}")
                # Essayer de sauvegarder avec nettoyage
                try:
                    doc.save(doc.name, garbage=4, deflate=True)
                    anonymized_pdf = doc.tobytes()
                except:
                    logging.error(f"‚ùå Impossible de sauvegarder le PDF: {str(e)}")
                    raise Exception(f"Erreur sauvegarde PDF: {str(e)}")
            else:
                raise
        
        doc.close()
        
        logging.info(f"‚úÖ Anonymisation PDF directe termin√©e avec succ√®s")
        logging.info(f"üìä Taille PDF anonymis√©: {len(anonymized_pdf)} bytes")
        
        return anonymized_pdf, mapping
        
    except Exception as e:
        logging.error(f"‚ùå Erreur lors de l'anonymisation PDF directe: {str(e)}")
        raise Exception(f"Erreur anonymisation PDF directe: {str(e)}")


def deanonymize_pdf_direct(pdf_content: bytes, mapping: Dict[str, str]) -> bytes:
    """
    D√©-anonymise directement un PDF en rempla√ßant les balises par les valeurs originales.
    Pr√©serve parfaitement la mise en page, les polices, les couleurs et la structure.
    
    Args:
        pdf_content: Le contenu du fichier PDF anonymis√©
        mapping: Dictionnaire de mapping des balises vers les valeurs originales
        
    Returns:
        bytes: Le contenu du fichier PDF d√©-anonymis√©
    """
    import fitz  # PyMuPDF
    
    logging.info("üöÄ D√©but d√©-anonymisation PDF directe avec PyMuPDF")
    logging.info(f"üìä Taille PDF d'entr√©e: {len(pdf_content)} bytes")
    logging.info(f"üìä Mapping: {len(mapping)} remplacements - {list(mapping.keys())}")
    
    try:
        # Ouvrir le PDF avec gestion d'erreur MuPDF am√©lior√©e
        try:
            doc = fitz.open(stream=pdf_content, filetype="pdf")
        except Exception as e:
            if "object out of range" in str(e):
                logging.warning(f"‚ö†Ô∏è Erreur MuPDF 'object out of range': {str(e)}")
                # Essayer de r√©cup√©rer le PDF en mode tol√©rant
                try:
                    doc = fitz.open(stream=pdf_content, filetype="pdf")
                    # R√©parer le PDF si possible
                    doc.save(doc.name, garbage=4, deflate=True)
                    doc.close()
                    doc = fitz.open(stream=pdf_content, filetype="pdf")
                except:
                    logging.error(f"‚ùå Impossible de r√©parer le PDF: {str(e)}")
                    raise Exception(f"PDF corrompu ou non support√©: {str(e)}")
            else:
                raise
        
        # Parcourir chaque page
        for page_num in range(len(doc)):
            page = doc[page_num]
            logging.info(f"üìÑ Traitement page {page_num + 1}/{len(doc)}")
            
            try:
                # Obtenir tous les blocs de texte de la page
                text_blocks = page.get_text("dict")
                
                # Parcourir chaque bloc de texte avec alignement parfait am√©lior√©
                for block in text_blocks["blocks"]:
                    if "lines" in block:  # Bloc de texte
                        _deanonymize_text_block_comprehensive(page, block, mapping)
                        
            except Exception as e:
                logging.warning(f"‚ö†Ô∏è Erreur traitement page {page_num + 1}: {str(e)}")
                continue
        
        # Sauvegarder le PDF modifi√© avec gestion d'erreur am√©lior√©e
        try:
            deanonymized_pdf = doc.tobytes()
        except Exception as e:
            if "object out of range" in str(e):
                logging.warning(f"‚ö†Ô∏è Erreur MuPDF lors de la sauvegarde: {str(e)}")
                # Essayer de sauvegarder avec nettoyage
                try:
                    doc.save(doc.name, garbage=4, deflate=True)
                    deanonymized_pdf = doc.tobytes()
                except:
                    logging.error(f"‚ùå Impossible de sauvegarder le PDF: {str(e)}")
                    raise Exception(f"Erreur sauvegarde PDF: {str(e)}")
            else:
                raise
        
        doc.close()
        
        logging.info(f"‚úÖ D√©-anonymisation PDF directe termin√©e avec succ√®s")
        logging.info(f"üìä Taille PDF d√©-anonymis√©: {len(deanonymized_pdf)} bytes")
        
        return deanonymized_pdf
        
    except Exception as e:
        logging.error(f"‚ùå Erreur lors de la d√©-anonymisation PDF directe: {str(e)}")
        raise Exception(f"Erreur d√©-anonymisation PDF directe: {str(e)}")


def _calculate_text_baseline_position(bbox, font_size):
    """
    Calcule la position optimale pour ins√©rer du texte en tenant compte de la ligne de base.
    
    Args:
        bbox: Rectangle englobant du texte original (fitz.Rect)
        font_size: Taille de la police
        
    Returns:
        tuple: (x, y) position pour insert_text
    """
    # Calculer la ligne de base approximative
    # La ligne de base est g√©n√©ralement situ√©e √† environ 20-25% de la hauteur depuis le bas
    height = bbox.height
    baseline_offset = height * 0.2  # 20% depuis le bas
    
    # Position x: coin gauche de la bo√Æte
    x = bbox.x0
    
    # Position y: bas de la bo√Æte + offset de ligne de base
    y = bbox.y1 - baseline_offset
    
    return (x, y)


def _get_text_width_estimation(text, font_size):
    """
    Estime la largeur du texte pour v√©rifier s'il rentre dans la bo√Æte englobante.
    
    Args:
        text: Texte √† mesurer
        font_size: Taille de la police
        
    Returns:
        float: Largeur estim√©e du texte
    """
    # Estimation approximative: largeur moyenne d'un caract√®re = 0.6 * font_size
    avg_char_width = font_size * 0.6
    return len(text) * avg_char_width


def _adjust_font_size_to_fit(text, bbox, original_font_size):
    """
    Ajuste la taille de la police pour que le texte rentre dans la bo√Æte englobante.
    
    Args:
        text: Texte √† ins√©rer
        bbox: Rectangle englobant disponible
        original_font_size: Taille de police originale
        
    Returns:
        float: Taille de police ajust√©e
    """
    available_width = bbox.width
    estimated_width = _get_text_width_estimation(text, original_font_size)
    
    if estimated_width <= available_width:
        return original_font_size
    
    # Calculer le facteur de r√©duction n√©cessaire
    scale_factor = available_width / estimated_width
    adjusted_size = original_font_size * scale_factor
    
    # Ne pas descendre en dessous de 6pt pour la lisibilit√©
    return max(adjusted_size, 6.0)


def _anonymize_text_block_direct(page, block, mapping: Dict[str, str]):
    """
    Anonymise un bloc de texte directement dans la page PDF avec positionnement pr√©cis.
    Le mapping contient: {valeur_originale: balise_anonymis√©e}
    """
    import fitz
    
    for line in block["lines"]:
        for span in line["spans"]:
            original_text = span["text"]
            
            if not original_text.strip():
                continue
                
            # Appliquer les remplacements du mapping
            anonymized_text = original_text
            text_changed = False
            
            # Le mapping contient {valeur_originale: balise_anonymis√©e}
            # Trier les cl√©s par longueur d√©croissante pour √©viter les remplacements partiels
            sorted_items = sorted(mapping.items(), key=lambda x: len(x[0]), reverse=True)
            
            for original_value, anonymized_tag in sorted_items:
                if original_value in anonymized_text:
                    anonymized_text = anonymized_text.replace(original_value, anonymized_tag)
                    text_changed = True
                    logging.debug(f"üîÑ Remplacement: '{original_value}' ‚Üí '{anonymized_tag}'")
            
            # Si le texte a chang√©, le remplacer dans le PDF
            if text_changed and anonymized_text != original_text:
                # Obtenir les propri√©t√©s du texte original
                font_name = span["font"]
                font_size = span["size"]
                font_flags = span["flags"]
                text_color = span["color"]
                bbox = fitz.Rect(span["bbox"])
                
                # Ajuster la taille de la police si n√©cessaire
                adjusted_font_size = _adjust_font_size_to_fit(anonymized_text, bbox, font_size)
                
                # Calculer la position optimale pour la ligne de base
                text_position = _calculate_text_baseline_position(bbox, adjusted_font_size)
                
                # Effacer l'ancien texte en le couvrant avec un rectangle blanc
                page.draw_rect(bbox, color=None, fill=fitz.pdfcolor["white"])
                
                # Ins√©rer le nouveau texte anonymis√© avec positionnement pr√©cis
                try:
                    page.insert_text(
                        text_position,  # Position calcul√©e pour la ligne de base
                        anonymized_text,
                        fontname=font_name,
                        fontsize=adjusted_font_size,
                        color=text_color
                    )
                    logging.debug(f"‚úÖ Texte remplac√© avec positionnement pr√©cis: '{original_text}' ‚Üí '{anonymized_text}' (taille: {adjusted_font_size})")
                except Exception as e:
                    logging.warning(f"‚ö†Ô∏è Erreur remplacement texte: {str(e)}")
                    # Fallback: utiliser une police par d√©faut avec positionnement pr√©cis
                    try:
                        page.insert_text(
                            text_position,
                            anonymized_text,
                            fontname="helv",  # Police par d√©faut
                            fontsize=adjusted_font_size,
                            color=text_color
                        )
                        logging.debug(f"‚úÖ Texte remplac√© avec police par d√©faut et positionnement pr√©cis")
                    except Exception as e2:
                        logging.error(f"‚ùå Impossible de remplacer le texte: {str(e2)}")


def _calculate_precise_text_position(bbox, text, fontname, fontsize, page):
    """
    Calcule la position ultra-pr√©cise pour ins√©rer du texte en utilisant les m√©triques r√©elles.
    
    Args:
        bbox: Rectangle englobant du texte original (fitz.Rect)
        text: Texte √† ins√©rer
        fontname: Nom de la police
        fontsize: Taille de la police
        page: Page PyMuPDF pour les m√©triques
        
    Returns:
        tuple: (x, y) position optimale pour insert_text
    """
    # Obtenir les m√©triques pr√©cises
    metrics = _get_precise_text_metrics(page, text, fontname, fontsize)
    
    # Position x: centr√©e horizontalement dans la bo√Æte si possible
    available_width = bbox.width
    text_width = metrics['width']
    
    if text_width <= available_width:
        # Centrer le texte horizontalement
        x = bbox.x0 + (available_width - text_width) / 2
    else:
        # Aligner √† gauche si trop large
        x = bbox.x0
    
    # Position y: calculer la ligne de base pr√©cise
    # La ligne de base est √† une distance de l'ascender depuis le haut
    bbox_height = bbox.height
    text_height = metrics['height']
    ascender = metrics['ascender']
    
    # Centrer verticalement puis ajuster pour la ligne de base
    y_center = bbox.y0 + bbox_height / 2
    y = y_center + ascender / 2
    
    # S'assurer que le texte reste dans les limites
    if y > bbox.y1:
        y = bbox.y1 - metrics['descender']
    elif y < bbox.y0 + ascender:
        y = bbox.y0 + ascender
    
    return (x, y)


def _adjust_font_size_precise(text, bbox, original_font_size, fontname, page):
    """
    Ajuste la taille de la police de mani√®re pr√©cise en utilisant les m√©triques r√©elles.
    
    Args:
        text: Texte √† ins√©rer
        bbox: Rectangle englobant disponible
        original_font_size: Taille de police originale
        fontname: Nom de la police
        page: Page PyMuPDF pour les m√©triques
        
    Returns:
        float: Taille de police ajust√©e
    """
    # Obtenir les m√©triques avec la taille originale
    metrics = _get_precise_text_metrics(page, text, fontname, original_font_size)
    
    available_width = bbox.width
    available_height = bbox.height
    
    # Calculer les facteurs d'ajustement
    width_scale = 1.0
    height_scale = 1.0
    
    if metrics['width'] > available_width:
        width_scale = available_width / metrics['width']
    
    if metrics['height'] > available_height:
        height_scale = available_height / metrics['height']
    
    # Utiliser le facteur le plus restrictif
    scale_factor = min(width_scale, height_scale)
    
    if scale_factor < 1.0:
        adjusted_size = original_font_size * scale_factor
        # Ne pas descendre en dessous de 6pt
        return max(adjusted_size, 6.0)
    
    return original_font_size


def _anonymize_text_block_ultra_precise(page, block, mapping: Dict[str, str]):
    """
    Anonymise un bloc de texte avec un positionnement ultra-pr√©cis utilisant les m√©triques r√©elles.
    Le mapping contient: {valeur_originale: balise_anonymis√©e}
    """
    import fitz
    
    for line in block["lines"]:
        for span in line["spans"]:
            original_text = span["text"]
            
            if not original_text.strip():
                continue
                
            # Appliquer les remplacements du mapping
            anonymized_text = original_text
            text_changed = False
            
            # Le mapping contient {valeur_originale: balise_anonymis√©e}
            # Trier les cl√©s par longueur d√©croissante pour √©viter les remplacements partiels
            sorted_items = sorted(mapping.items(), key=lambda x: len(x[0]), reverse=True)
            
            for original_value, anonymized_tag in sorted_items:
                if original_value in anonymized_text:
                    anonymized_text = anonymized_text.replace(original_value, anonymized_tag)
                    text_changed = True
                    logging.debug(f"üîÑ Remplacement: '{original_value}' ‚Üí '{anonymized_tag}'")
            
            # Si le texte a chang√©, le remplacer dans le PDF
            if text_changed and anonymized_text != original_text:
                # Obtenir les propri√©t√©s du texte original
                font_name = span["font"]
                font_size = span["size"]
                font_flags = span["flags"]
                text_color = span["color"]
                bbox = fitz.Rect(span["bbox"])
                
                # Ajuster la taille de la police avec pr√©cision
                adjusted_font_size = _adjust_font_size_precise(
                    anonymized_text, bbox, font_size, font_name, page
                )
                
                # Calculer la position ultra-pr√©cise
                text_position = _calculate_precise_text_position(
                    bbox, anonymized_text, font_name, adjusted_font_size, page
                )
                
                # Effacer l'ancien texte avec un rectangle l√©g√®rement plus grand
                padding = 1  # 1 point de padding
                expanded_bbox = fitz.Rect(
                    bbox.x0 - padding, bbox.y0 - padding,
                    bbox.x1 + padding, bbox.y1 + padding
                )
                page.draw_rect(expanded_bbox, color=None, fill=fitz.pdfcolor["white"])
                
                # Ins√©rer le nouveau texte anonymis√© avec positionnement ultra-pr√©cis
                try:
                    page.insert_text(
                        text_position,
                        anonymized_text,
                        fontname=font_name,
                        fontsize=adjusted_font_size,
                        color=text_color
                    )
                    logging.debug(f"‚úÖ Texte remplac√© avec positionnement ultra-pr√©cis: '{original_text}' ‚Üí '{anonymized_text}' (taille: {adjusted_font_size:.1f})")
                except Exception as e:
                    logging.warning(f"‚ö†Ô∏è Erreur remplacement texte: {str(e)}")
                    # Fallback avec police par d√©faut
                    try:
                        fallback_position = _calculate_precise_text_position(
                            bbox, anonymized_text, "helv", adjusted_font_size, page
                        )
                        page.insert_text(
                            fallback_position,
                            anonymized_text,
                            fontname="helv",
                            fontsize=adjusted_font_size,
                            color=text_color
                        )
                        logging.debug(f"‚úÖ Texte remplac√© avec police par d√©faut et positionnement ultra-pr√©cis")
                    except Exception as e2:
                        logging.error(f"‚ùå Impossible de remplacer le texte: {str(e2)}")


def _deanonymize_text_block_ultra_precise(page, block, mapping: Dict[str, str]):
    """
    D√©-anonymise un bloc de texte avec un positionnement ultra-pr√©cis utilisant les m√©triques r√©elles.
    """
    import fitz
    
    for line in block["lines"]:
        for span in line["spans"]:
            original_text = span["text"]
            
            if not original_text.strip():
                continue
                
            # Appliquer les remplacements du mapping (balise ‚Üí valeur originale)
            deanonymized_text = original_text
            text_changed = False
            
            # Trier les cl√©s par longueur d√©croissante pour √©viter les remplacements partiels
            sorted_keys = sorted(mapping.keys(), key=len, reverse=True)
            
            for anonymized_tag, original_value in mapping.items():
                if anonymized_tag in deanonymized_text:
                    deanonymized_text = deanonymized_text.replace(anonymized_tag, original_value)
                    text_changed = True
                    logging.debug(f"üîÑ D√©-anonymisation: '{anonymized_tag}' ‚Üí '{original_value}'")
            
            # Si le texte a chang√©, le remplacer dans le PDF
            if text_changed and deanonymized_text != original_text:
                # Obtenir les propri√©t√©s du texte original
                font_name = span["font"]
                font_size = span["size"]
                font_flags = span["flags"]
                text_color = span["color"]
                bbox = fitz.Rect(span["bbox"])
                
                # Ajuster la taille de la police avec pr√©cision
                adjusted_font_size = _adjust_font_size_precise(
                    deanonymized_text, bbox, font_size, font_name, page
                )
                
                # Calculer la position ultra-pr√©cise
                text_position = _calculate_precise_text_position(
                    bbox, deanonymized_text, font_name, adjusted_font_size, page
                )
                
                # Effacer l'ancien texte avec un rectangle l√©g√®rement plus grand
                padding = 1  # 1 point de padding
                expanded_bbox = fitz.Rect(
                    bbox.x0 - padding, bbox.y0 - padding,
                    bbox.x1 + padding, bbox.y1 + padding
                )
                page.draw_rect(expanded_bbox, color=None, fill=fitz.pdfcolor["white"])
                
                # Ins√©rer le nouveau texte d√©-anonymis√© avec positionnement ultra-pr√©cis
                try:
                    page.insert_text(
                        text_position,
                        deanonymized_text,
                        fontname=font_name,
                        fontsize=adjusted_font_size,
                        color=text_color
                    )
                    logging.debug(f"‚úÖ Texte d√©-anonymis√© avec positionnement ultra-pr√©cis: '{original_text}' ‚Üí '{deanonymized_text}' (taille: {adjusted_font_size:.1f})")
                except Exception as e:
                    logging.warning(f"‚ö†Ô∏è Erreur d√©-anonymisation texte: {str(e)}")
                    # Fallback avec police par d√©faut
                    try:
                        fallback_position = _calculate_precise_text_position(
                            bbox, deanonymized_text, "helv", adjusted_font_size, page
                        )
                        page.insert_text(
                            fallback_position,
                            deanonymized_text,
                            fontname="helv",
                            fontsize=adjusted_font_size,
                            color=text_color
                        )
                        logging.debug(f"‚úÖ Texte d√©-anonymis√© avec police par d√©faut et positionnement ultra-pr√©cis")
                    except Exception as e2:
                        logging.error(f"‚ùå Impossible de d√©-anonymiser le texte: {str(e2)}")


def _preserve_original_text_alignment(bbox, original_text, new_text, fontname, fontsize, page):
    """
    Calcule la position exacte pour pr√©server l'alignement original du texte.
    
    Args:
        bbox: Rectangle englobant du texte original
        original_text: Texte original
        new_text: Nouveau texte √† ins√©rer
        fontname: Nom de la police
        fontsize: Taille de la police
        page: Page PyMuPDF
        
    Returns:
        tuple: (x, y) position exacte pour pr√©server l'alignement
    """
    # Position x: utiliser exactement la position originale (pas de centrage)
    x = bbox.x0
    
    # Position y: calculer la ligne de base correctement pour √©viter les coupures
    # La ligne de base doit √™tre calcul√©e en tenant compte des descenders
    # Utiliser une approche plus pr√©cise bas√©e sur les m√©triques de la police
    
    # Estimation de la hauteur de la ligne de base (g√©n√©ralement 75-80% de la hauteur)
    baseline_ratio = 0.75
    y = bbox.y0 + (bbox.height * baseline_ratio)
    
    # Ajustement pour √©viter les coupures de descenders (g, j, p, q, y)
    # V√©rifier si le nouveau texte contient des descenders
    has_descenders = any(char in new_text for char in 'gjpqy')
    if has_descenders:
        # Ajuster l√©g√®rement vers le haut pour √©viter les coupures
        y = y - (fontsize * 0.1)  # Ajustement de 10% de la taille de police
        logging.debug(f"üîß Ajustement descenders pour '{new_text}': y = {y:.2f} (ajust√© de -{fontsize * 0.1:.2f})")
    
    logging.debug(f"üìç Position calcul√©e pour '{new_text}': x={x:.2f}, y={y:.2f} (bbox: {bbox}, ratio: {baseline_ratio})")
    
    return (x, y)


def _get_best_matching_font(original_fontname, page):
    """
    Trouve la meilleure police disponible qui correspond √† la police originale.
    
    Args:
        original_fontname: Nom de la police originale
        page: Page PyMuPDF
        
    Returns:
        str: Nom de la police √† utiliser
    """
    import fitz
    
    # Nettoyer le nom de la police original
    clean_fontname = original_fontname.lower().replace("-", "").replace(" ", "")
    
    # Mapping des polices communes
    font_mapping = {
        "helvetica": "helv",
        "arial": "helv", 
        "times": "times",
        "timesnewroman": "times",
        "courier": "cour",
        "couriernew": "cour",
        "calibri": "helv",
        "verdana": "helv",
        "georgia": "times",
        "trebuchet": "helv",
        "lucida": "helv",
        "palatino": "times",
        "garamond": "times",
        "bookman": "times",
        "avantgarde": "helv",
        "zapfdingbats": "zadb",
        "symbol": "symb"
    }
    
    # Chercher une correspondance directe
    for pattern, mapped_font in font_mapping.items():
        if pattern in clean_fontname:
            return mapped_font
    
    # Fallback intelligent bas√© sur les caract√©ristiques de la police
    if "bold" in clean_fontname or "black" in clean_fontname:
        return "helv"  # Helvetica pour le gras
    elif "italic" in clean_fontname or "oblique" in clean_fontname:
        return "helv"  # Helvetica pour l'italique
    elif "mono" in clean_fontname or "fixed" in clean_fontname:
        return "cour"  # Courier pour les polices monospace
    elif "serif" in clean_fontname or "roman" in clean_fontname:
        return "times"  # Times pour les polices serif
    else:
        return "helv"  # Helvetica par d√©faut


def _anonymize_text_block_perfect_alignment(page, block, mapping: Dict[str, str]):
    """
    Anonymise un bloc de texte en pr√©servant parfaitement l'alignement et les polices originales.
    Le mapping contient: {valeur_originale: balise_anonymis√©e}
    """
    import fitz
    
    for line in block["lines"]:
        # Traiter tous les spans d'une ligne ensemble pour maintenir l'alignement
        line_spans = []
        
        for span in line["spans"]:
            original_text = span["text"]
            
            if not original_text.strip():
                continue
                
            # Appliquer les remplacements du mapping avec remplacement s√©curis√©
            anonymized_text, text_changed = _safe_replace_in_span_text(original_text, mapping)
            
            # Ajouter ce span √† la liste des spans de la ligne
            line_spans.append({
                'span': span,
                'original_text': original_text,
                'anonymized_text': anonymized_text,
                'text_changed': text_changed
            })
        
        # Traiter tous les spans modifi√©s de la ligne
        for span_info in line_spans:
            if span_info['text_changed'] and span_info['anonymized_text'] != span_info['original_text']:
                span = span_info['span']
                original_text = span_info['original_text']
                anonymized_text = span_info['anonymized_text']
                
                # Obtenir les propri√©t√©s du texte original
                font_name = span["font"]
                font_size = span["size"]
                font_flags = span["flags"]  # Important pour le formatage
                text_color = span["color"]
                bbox = fitz.Rect(span["bbox"])
                
                # V√©rifier et ajuster les marges droites
                page_rect = page.rect
                if bbox.x1 > page_rect.x1 - 50:  # Marge droite de 50 points
                    logging.warning(f"‚ö†Ô∏è Texte trop proche du bord droit: {bbox.x1} > {page_rect.x1 - 50}")
                    # Ajuster la largeur disponible
                    available_width = page_rect.x1 - bbox.x0 - 50
                    if available_width < 100:  # Minimum 100 points
                        available_width = 100
                    bbox = fitz.Rect(bbox.x0, bbox.y0, bbox.x0 + available_width, bbox.y1)
                
                # Calculer la position exacte pour pr√©server l'alignement
                text_position = _preserve_original_text_alignment(
                    bbox, original_text, anonymized_text, font_name, font_size, page
                )
                
                # Effacer l'ancien texte avec un rectangle exact (sans padding excessif)
                try:
                    page.draw_rect(bbox, color=None, fill=fitz.pdfcolor["white"])
                except Exception as e:
                    logging.warning(f"‚ö†Ô∏è Erreur effacement rectangle: {str(e)}")
                
                # Ins√©rer le nouveau texte anonymis√© avec formatage pr√©serv√©
                success = False
                
                # Tentative 1: Utiliser la police originale exacte
                try:
                    # Normaliser la couleur pour PyMuPDF
                    normalized_color = _normalize_color(text_color)
                    
                    page.insert_text(
                        text_position,
                        anonymized_text,
                        fontname=font_name,  # Police originale exacte
                        fontsize=font_size,
                        color=normalized_color
                    )
                    success = True
                    logging.debug(f"‚úÖ Texte remplac√© avec police originale: '{original_text}' ‚Üí '{anonymized_text}' (police: {font_name})")
                except Exception as e:
                    logging.debug(f"‚ö†Ô∏è Police originale √©chou√©e: {str(e)}")
                
                # Tentative 2: Utiliser une police de base avec formatage
                if not success:
                    try:
                        # Trouver la meilleure police correspondante
                        best_font = _get_best_matching_font(font_name, page)
                        
                        # Appliquer le formatage si possible
                        formatted_font = _apply_font_formatting_safe(best_font, font_flags)
                        
                        normalized_color = _normalize_color(text_color)
                        
                        page.insert_text(
                            text_position,
                            anonymized_text,
                            fontname=formatted_font,
                            fontsize=font_size,
                            color=normalized_color
                        )
                        success = True
                        logging.debug(f"‚úÖ Texte remplac√© avec police format√©e: '{original_text}' ‚Üí '{anonymized_text}' (police: {formatted_font})")
                    except Exception as e:
                        logging.debug(f"‚ö†Ô∏è Police format√©e √©chou√©e: {str(e)}")
                
                # Tentative 3: Fallback avec police de base
                if not success:
                    try:
                        best_font = _get_best_matching_font(font_name, page)
                        normalized_color = _normalize_color(text_color)
                        
                        page.insert_text(
                            text_position,
                            anonymized_text,
                            fontname=best_font,
                            fontsize=font_size,
                            color=normalized_color
                        )
                        success = True
                        logging.debug(f"‚úÖ Texte remplac√© avec police de base: '{original_text}' ‚Üí '{anonymized_text}' (police: {best_font})")
                    except Exception as e:
                        logging.debug(f"‚ö†Ô∏è Police de base √©chou√©e: {str(e)}")
                
                # Tentative 4: Dernier fallback avec Helvetica
                if not success:
                    try:
                        normalized_color = _normalize_color(text_color)
                        
                        page.insert_text(
                            text_position,
                            anonymized_text,
                            fontname="helv",
                            fontsize=font_size,
                            color=normalized_color
                        )
                        success = True
                        logging.debug(f"‚úÖ Texte remplac√© avec Helvetica: '{original_text}' ‚Üí '{anonymized_text}'")
                    except Exception as e:
                        logging.error(f"‚ùå Impossible de remplacer le texte: {str(e)}")


def _apply_font_formatting_safe(base_font: str, font_flags: int) -> str:
    """
    Applique le formatage (gras, italique) √† une police de base de mani√®re s√©curis√©e.
    Retourne la police de base si le formatage n'est pas disponible.
    
    Args:
        base_font: Police de base (helv, times, cour)
        font_flags: Flags de formatage PyMuPDF
        
    Returns:
        str: Nom de police avec formatage appliqu√© ou police de base
    """
    import fitz
    
    # Constantes PyMuPDF pour les flags
    BOLD_FLAG = 2**4   # 16
    ITALIC_FLAG = 2**5  # 32
    
    # D√©tecter le formatage
    is_bold = bool(font_flags & BOLD_FLAG)
    is_italic = bool(font_flags & ITALIC_FLAG)
    
    logging.debug(f"üé® Formatage d√©tect√©: flags={font_flags}, bold={is_bold}, italic={is_italic}")
    
    # Si pas de formatage, retourner la police de base
    if not is_bold and not is_italic:
        return base_font
    
    # Essayer d'appliquer le formatage
    try:
        # Appliquer le formatage selon la police de base
        if base_font == "helv":
            if is_bold and is_italic:
                result = "helv-bolditalic"
            elif is_bold:
                result = "helv-bold"
            elif is_italic:
                result = "helv-italic"
            else:
                result = "helv"
        elif base_font == "times":
            if is_bold and is_italic:
                result = "times-bolditalic"
            elif is_bold:
                result = "times-bold"
            elif is_italic:
                result = "times-italic"
            else:
                result = "times"
        elif base_font == "cour":
            if is_bold and is_italic:
                result = "cour-bolditalic"
            elif is_bold:
                result = "cour-bold"
            elif is_italic:
                result = "cour-italic"
            else:
                result = "cour"
        else:
            # Pour les autres polices, utiliser Helvetica avec formatage
            if is_bold and is_italic:
                result = "helv-bolditalic"
            elif is_bold:
                result = "helv-bold"
            elif is_italic:
                result = "helv-italic"
            else:
                result = "helv"
        
        # Tester si la police format√©e est disponible
        try:
            fitz.Font(result)
            logging.debug(f"üé® Police format√©e disponible: '{base_font}' ‚Üí '{result}'")
            return result
        except:
            logging.debug(f"‚ö†Ô∏è Police format√©e non disponible: '{result}', utilisation de '{base_font}'")
            return base_font
            
    except Exception as e:
        logging.debug(f"‚ö†Ô∏è Erreur formatage police: {str(e)}, utilisation de '{base_font}'")
        return base_font


def _deanonymize_text_block_perfect_alignment(page, block, mapping: Dict[str, str]):
    """
    D√©-anonymise un bloc de texte en pr√©servant parfaitement l'alignement et les polices originales.
    """
    import fitz
    
    for line in block["lines"]:
        # Traiter tous les spans d'une ligne ensemble pour maintenir l'alignement
        line_spans = []
        
        for span in line["spans"]:
            original_text = span["text"]
            
            if not original_text.strip():
                continue
                
            # Appliquer les remplacements du mapping avec remplacement s√©curis√© (balise ‚Üí valeur originale)
            deanonymized_text, text_changed = _safe_replace_in_span_text(original_text, mapping)
            
            # Ajouter ce span √† la liste des spans de la ligne
            line_spans.append({
                'span': span,
                'original_text': original_text,
                'deanonymized_text': deanonymized_text,
                'text_changed': text_changed
            })
        
        # Traiter tous les spans modifi√©s de la ligne
        for span_info in line_spans:
            if span_info['text_changed'] and span_info['deanonymized_text'] != span_info['original_text']:
                span = span_info['span']
                original_text = span_info['original_text']
                deanonymized_text = span_info['deanonymized_text']
                
                # Obtenir les propri√©t√©s du texte original
                font_name = span["font"]
                font_size = span["size"]
                font_flags = span["flags"]  # Important pour le formatage
                text_color = span["color"]
                bbox = fitz.Rect(span["bbox"])
                
                # V√©rifier et ajuster les marges droites
                page_rect = page.rect
                if bbox.x1 > page_rect.x1 - 50:  # Marge droite de 50 points
                    logging.warning(f"‚ö†Ô∏è Texte trop proche du bord droit: {bbox.x1} > {page_rect.x1 - 50}")
                    # Ajuster la largeur disponible
                    available_width = page_rect.x1 - bbox.x0 - 50
                    if available_width < 100:  # Minimum 100 points
                        available_width = 100
                    bbox = fitz.Rect(bbox.x0, bbox.y0, bbox.x0 + available_width, bbox.y1)
                
                # Calculer la position exacte pour pr√©server l'alignement
                text_position = _preserve_original_text_alignment(
                    bbox, original_text, deanonymized_text, font_name, font_size, page
                )
                
                # Effacer l'ancien texte avec un rectangle exact (sans padding excessif)
                try:
                    page.draw_rect(bbox, color=None, fill=fitz.pdfcolor["white"])
                except Exception as e:
                    logging.warning(f"‚ö†Ô∏è Erreur effacement rectangle: {str(e)}")
                
                # Ins√©rer le nouveau texte d√©-anonymis√© avec formatage pr√©serv√©
                success = False
                
                # Tentative 1: Utiliser la police originale exacte
                try:
                    # Normaliser la couleur pour PyMuPDF
                    normalized_color = _normalize_color(text_color)
                    
                    page.insert_text(
                        text_position,
                        deanonymized_text,
                        fontname=font_name,  # Police originale exacte
                        fontsize=font_size,
                        color=normalized_color
                    )
                    success = True
                    logging.debug(f"‚úÖ Texte d√©-anonymis√© avec police originale: '{original_text}' ‚Üí '{deanonymized_text}' (police: {font_name})")
                except Exception as e:
                    logging.debug(f"‚ö†Ô∏è Police originale √©chou√©e: {str(e)}")
                
                # Tentative 2: Utiliser une police de base avec formatage
                if not success:
                    try:
                        # Trouver la meilleure police correspondante
                        best_font = _get_best_matching_font(font_name, page)
                        
                        # Appliquer le formatage si possible
                        formatted_font = _apply_font_formatting_safe(best_font, font_flags)
                        
                        normalized_color = _normalize_color(text_color)
                        
                        page.insert_text(
                            text_position,
                            deanonymized_text,
                            fontname=formatted_font,
                            fontsize=font_size,
                            color=normalized_color
                        )
                        success = True
                        logging.debug(f"‚úÖ Texte d√©-anonymis√© avec police format√©e: '{original_text}' ‚Üí '{deanonymized_text}' (police: {formatted_font})")
                    except Exception as e:
                        logging.debug(f"‚ö†Ô∏è Police format√©e √©chou√©e: {str(e)}")
                
                # Tentative 3: Fallback avec police de base
                if not success:
                    try:
                        best_font = _get_best_matching_font(font_name, page)
                        normalized_color = _normalize_color(text_color)
                        
                        page.insert_text(
                            text_position,
                            deanonymized_text,
                            fontname=best_font,
                            fontsize=font_size,
                            color=normalized_color
                        )
                        success = True
                        logging.debug(f"‚úÖ Texte d√©-anonymis√© avec police de base: '{original_text}' ‚Üí '{deanonymized_text}' (police: {best_font})")
                    except Exception as e:
                        logging.debug(f"‚ö†Ô∏è Police de base √©chou√©e: {str(e)}")
                
                # Tentative 4: Dernier fallback avec Helvetica
                if not success:
                    try:
                        normalized_color = _normalize_color(text_color)
                        
                        page.insert_text(
                            text_position,
                            deanonymized_text,
                            fontname="helv",
                            fontsize=font_size,
                            color=normalized_color
                        )
                        success = True
                        logging.debug(f"‚úÖ Texte d√©-anonymis√© avec Helvetica: '{original_text}' ‚Üí '{deanonymized_text}'")
                    except Exception as e:
                        logging.error(f"‚ùå Impossible de d√©-anonymiser le texte: {str(e)}")


# ===== FIN ANONYMISATION DIRECTE PDF =====

def _safe_replace_with_word_boundaries(text: str, old_value: str, new_value: str) -> str:
    """
    Remplace une valeur dans le texte de mani√®re agressive pour l'anonymisation.
    
    Args:
        text: Texte dans lequel effectuer le remplacement
        old_value: Valeur √† remplacer
        new_value: Nouvelle valeur
        
    Returns:
        str: Texte avec les remplacements effectu√©s
    """
    # √âchapper les caract√®res sp√©ciaux pour les expressions r√©guli√®res
    escaped_old = re.escape(old_value)
    
    # CORRECTION: Utiliser un pattern plus agressif pour l'anonymisation
    # Pattern 1: Avec limites de mots (pour les mots entiers)
    pattern1 = r'\b' + escaped_old + r'\b'
    result = re.sub(pattern1, new_value, text, flags=re.IGNORECASE)
    
    # Pattern 2: Sans limites de mots (pour capturer les cas manqu√©s)
    # Utiliser des espaces ou ponctuation autour
    pattern2 = r'(\s|^)' + escaped_old + r'(\s|$|[.,;:!?])'
    result = re.sub(pattern2, r'\1' + new_value + r'\2', result, flags=re.IGNORECASE)
    
    # Pattern 3: Remplacement direct pour les cas restants
    # Utiliser une recherche simple mais avec v√©rification de contexte
    if old_value.lower() in result.lower():
        # Remplacer toutes les occurrences restantes
        result = result.replace(old_value, new_value)
        result = result.replace(old_value.upper(), new_value.upper())
        result = result.replace(old_value.lower(), new_value.lower())
        result = result.replace(old_value.title(), new_value.title())
    
    return result


def _safe_replace_in_span_text(span_text: str, mapping: Dict[str, str]) -> Tuple[str, bool]:
    """
    Remplace les valeurs dans un span de texte en utilisant des limites de mots.
    
    Args:
        span_text: Texte du span
        mapping: Mapping des remplacements {old_value: new_value}
        
    Returns:
        Tuple[str, bool]: (texte_modifi√©, changement_effectu√©)
    """
    modified_text = span_text
    text_changed = False
    
    # Trier les cl√©s par longueur d√©croissante pour √©viter les remplacements partiels
    sorted_items = sorted(mapping.items(), key=lambda x: len(x[0]), reverse=True)
    
    for old_value, new_value in sorted_items:
        # Utiliser le remplacement s√©curis√© avec limites de mots
        before_replace = modified_text
        modified_text = _safe_replace_with_word_boundaries(modified_text, old_value, new_value)
        
        if modified_text != before_replace:
            text_changed = True
            logging.debug(f"üîÑ Remplacement s√©curis√©: '{old_value}' ‚Üí '{new_value}'")
    
    return modified_text, text_changed


def _normalize_color(color_value):
    """
    Normalise une valeur de couleur pour PyMuPDF.
    
    Args:
        color_value: Valeur de couleur (peut √™tre int, float, tuple, etc.)
        
    Returns:
        tuple: Couleur normalis√©e (R, G, B) avec des valeurs entre 0 et 1
    """
    if color_value is None:
        return (0, 0, 0)  # Noir par d√©faut
    
    # Si c'est d√©j√† un tuple/liste de 3 valeurs
    if isinstance(color_value, (tuple, list)) and len(color_value) == 3:
        return tuple(float(c) for c in color_value)
    
    # Si c'est un entier (couleur RGB encod√©e)
    if isinstance(color_value, int):
        # Convertir l'entier en RGB
        r = ((color_value >> 16) & 0xFF) / 255.0
        g = ((color_value >> 8) & 0xFF) / 255.0
        b = (color_value & 0xFF) / 255.0
        return (r, g, b)
    
    # Si c'est un float (niveau de gris)
    if isinstance(color_value, float):
        if 0 <= color_value <= 1:
            return (color_value, color_value, color_value)
        else:
            # Normaliser si > 1
            normalized = color_value / 255.0 if color_value > 1 else color_value
            return (normalized, normalized, normalized)
    
    # Fallback: noir avec log d√©taill√©
    logging.debug(f"üé® Couleur non reconnue: {color_value} (type: {type(color_value)}), utilisation du noir")
    return (0, 0, 0)


def _apply_font_formatting_safe(base_font: str, font_flags: int) -> str:
    """
    Applique le formatage (gras, italique) √† une police de base de mani√®re s√©curis√©e.
    Retourne la police de base si le formatage n'est pas disponible.
    
    Args:
        base_font: Police de base (helv, times, cour)
        font_flags: Flags de formatage PyMuPDF
        
    Returns:
        str: Nom de police avec formatage appliqu√© ou police de base
    """
    import fitz
    
    # Constantes PyMuPDF pour les flags
    BOLD_FLAG = 2**4   # 16
    ITALIC_FLAG = 2**5  # 32
    
    # D√©tecter le formatage
    is_bold = bool(font_flags & BOLD_FLAG)
    is_italic = bool(font_flags & ITALIC_FLAG)
    
    logging.debug(f"üé® Formatage d√©tect√©: flags={font_flags}, bold={is_bold}, italic={is_italic}")
    
    # Si pas de formatage, retourner la police de base
    if not is_bold and not is_italic:
        return base_font
    
    # Essayer d'appliquer le formatage
    try:
        # Appliquer le formatage selon la police de base
        if base_font == "helv":
            if is_bold and is_italic:
                result = "helv-bolditalic"
            elif is_bold:
                result = "helv-bold"
            elif is_italic:
                result = "helv-italic"
            else:
                result = "helv"
        elif base_font == "times":
            if is_bold and is_italic:
                result = "times-bolditalic"
            elif is_bold:
                result = "times-bold"
            elif is_italic:
                result = "times-italic"
            else:
                result = "times"
        elif base_font == "cour":
            if is_bold and is_italic:
                result = "cour-bolditalic"
            elif is_bold:
                result = "cour-bold"
            elif is_italic:
                result = "cour-italic"
            else:
                result = "cour"
        else:
            # Pour les autres polices, utiliser Helvetica avec formatage
            if is_bold and is_italic:
                result = "helv-bolditalic"
            elif is_bold:
                result = "helv-bold"
            elif is_italic:
                result = "helv-italic"
            else:
                result = "helv"
        
        # Tester si la police format√©e est disponible
        try:
            fitz.Font(result)
            logging.debug(f"üé® Police format√©e disponible: '{base_font}' ‚Üí '{result}'")
            return result
        except:
            logging.debug(f"‚ö†Ô∏è Police format√©e non disponible: '{result}', utilisation de '{base_font}'")
            return base_font
            
    except Exception as e:
        logging.debug(f"‚ö†Ô∏è Erreur formatage police: {str(e)}, utilisation de '{base_font}'")
        return base_font


def _anonymize_text_block_comprehensive(page, block, mapping: Dict[str, str]):
    """
    Anonymise un bloc de texte avec gestion compl√®te des probl√®mes:
    - Marge droite respect√©e avec ajustement automatique
    - Formatage gras/couleur pr√©serv√©
    - Gestion robuste des erreurs MuPDF
    Le mapping contient: {valeur_originale: balise_anonymis√©e}
    Ajout : logs d√©taill√©s pour chaque remplacement.
    """
    import fitz
    for line in block["lines"]:
        # Traiter tous les spans d'une ligne ensemble pour maintenir l'alignement
        line_spans = []
        for span in line["spans"]:
            original_text = span["text"]
            if not original_text.strip():
                continue
            # Appliquer les remplacements du mapping avec remplacement s√©curis√©
            anonymized_text, text_changed = _safe_replace_in_span_text(original_text, mapping)
            # Ajouter ce span √† la liste des spans de la ligne
            line_spans.append({
                'span': span,
                'original_text': original_text,
                'anonymized_text': anonymized_text,
                'text_changed': text_changed
            })
        # Traiter tous les spans modifi√©s de la ligne
        for span_info in line_spans:
            if span_info['text_changed'] and span_info['anonymized_text'] != span_info['original_text']:
                span = span_info['span']
                original_text = span_info['original_text']
                anonymized_text = span_info['anonymized_text']
                font_name = span["font"]
                font_size = span["size"]
                font_flags = span["flags"]
                text_color = span["color"]
                bbox = fitz.Rect(span["bbox"])
                # LOG D√âTAILL√â
                logging.info(f"üîÑ Remplacement PDF: '{original_text}' ‚Üí '{anonymized_text}' | Page: {page.number+1} | BBox: {bbox} | Police: {font_name} | Taille: {font_size} | Flags: {font_flags} | Couleur: {text_color}")
                # CORRECTION 1: V√©rifier et ajuster les marges droites de mani√®re robuste
                page_rect = page.rect
                margin_right = 50  # Marge droite de 50 points
                
                if bbox.x1 > page_rect.x1 - margin_right:
                    logging.warning(f"‚ö†Ô∏è Texte trop proche du bord droit: {bbox.x1} > {page_rect.x1 - margin_right}")
                    logging.info(f"üìè Texte original: '{original_text}' ‚Üí Texte anonymis√©: '{anonymized_text}'")
                    
                    # Ajuster la largeur disponible
                    available_width = page_rect.x1 - bbox.x0 - margin_right
                    if available_width < 100:  # Minimum 100 points
                        available_width = 100
                        logging.warning(f"‚ö†Ô∏è Largeur disponible tr√®s r√©duite: {available_width} points")
                    
                    # Cr√©er un nouveau bbox avec la largeur ajust√©e
                    original_bbox = bbox
                    bbox = fitz.Rect(bbox.x0, bbox.y0, bbox.x0 + available_width, bbox.y1)
                    logging.info(f"üìè Bbox ajust√© pour marge droite: {original_bbox} ‚Üí {bbox}")
                    
                    # V√©rifier si le texte anonymis√© est plus long que l'original
                    if len(anonymized_text) > len(original_text):
                        logging.warning(f"‚ö†Ô∏è Texte anonymis√© plus long que l'original: {len(anonymized_text)} > {len(original_text)}")
                        # Ajuster la taille de police si n√©cessaire
                        font_size = font_size * 0.9  # R√©duire de 10%
                        logging.info(f"üìè Taille de police ajust√©e: {font_size:.1f}")
                
                # Calculer la position exacte pour pr√©server l'alignement
                text_position = _preserve_original_text_alignment(
                    bbox, original_text, anonymized_text, font_name, font_size, page
                )
                
                # Effacer l'ancien texte avec gestion d'erreur robuste
                try:
                    page.draw_rect(bbox, color=None, fill=fitz.pdfcolor["white"])
                except Exception as e:
                    logging.warning(f"‚ö†Ô∏è Erreur effacement rectangle: {str(e)}")
                    # Essayer avec un rectangle l√©g√®rement plus petit
                    try:
                        smaller_bbox = fitz.Rect(bbox.x0 + 1, bbox.y0 + 1, bbox.x1 - 1, bbox.y1 - 1)
                        page.draw_rect(smaller_bbox, color=None, fill=fitz.pdfcolor["white"])
                    except Exception as e2:
                        logging.warning(f"‚ö†Ô∏è Impossible d'effacer le rectangle: {str(e2)}")
                
                # CORRECTION 2: Ins√©rer le nouveau texte avec formatage pr√©serv√© d√®s l'anonymisation
                success = _insert_text_with_preserved_formatting(
                    page, text_position, anonymized_text, font_name, font_size, font_flags, text_color, original_text
                )
                
                if not success:
                    logging.error(f"‚ùå Impossible de remplacer le texte: '{original_text}' ‚Üí '{anonymized_text}'")


def _deanonymize_text_block_comprehensive(page, block, mapping: Dict[str, str]):
    """
    D√©-anonymise un bloc de texte avec gestion compl√®te des probl√®mes:
    - Marge droite respect√©e avec ajustement automatique
    - Formatage gras/couleur pr√©serv√© et restaur√©
    - Gestion robuste des erreurs MuPDF
    Le mapping contient: {balise_anonymis√©e: valeur_originale}
    """
    import fitz
    
    for line in block["lines"]:
        # Traiter tous les spans d'une ligne ensemble pour maintenir l'alignement
        line_spans = []
        
        for span in line["spans"]:
            original_text = span["text"]
            
            if not original_text.strip():
                continue
                
            # Appliquer les remplacements du mapping avec remplacement s√©curis√©
            deanonymized_text, text_changed = _safe_replace_in_span_text(original_text, mapping)
            
            # Ajouter ce span √† la liste des spans de la ligne
            line_spans.append({
                'span': span,
                'original_text': original_text,
                'deanonymized_text': deanonymized_text,
                'text_changed': text_changed
            })
        
        # Traiter tous les spans modifi√©s de la ligne
        for span_info in line_spans:
            if span_info['text_changed'] and span_info['deanonymized_text'] != span_info['original_text']:
                span = span_info['span']
                original_text = span_info['original_text']
                deanonymized_text = span_info['deanonymized_text']
                
                # Obtenir les propri√©t√©s du texte original
                font_name = span["font"]
                font_size = span["size"]
                font_flags = span["flags"]  # Important pour le formatage
                text_color = span["color"]
                bbox = fitz.Rect(span["bbox"])
                
                # CORRECTION 1: V√©rifier et ajuster les marges droites de mani√®re robuste
                page_rect = page.rect
                margin_right = 50  # Marge droite de 50 points
                
                if bbox.x1 > page_rect.x1 - margin_right:
                    logging.warning(f"‚ö†Ô∏è Texte trop proche du bord droit: {bbox.x1} > {page_rect.x1 - margin_right}")
                    # Ajuster la largeur disponible
                    available_width = page_rect.x1 - bbox.x0 - margin_right
                    if available_width < 100:  # Minimum 100 points
                        available_width = 100
                        logging.warning(f"‚ö†Ô∏è Largeur disponible tr√®s r√©duite: {available_width} points")
                    
                    # Cr√©er un nouveau bbox avec la largeur ajust√©e
                    bbox = fitz.Rect(bbox.x0, bbox.y0, bbox.x0 + available_width, bbox.y1)
                    logging.info(f"üìè Bbox ajust√© pour marge droite: {bbox}")
                
                # Calculer la position exacte pour pr√©server l'alignement
                text_position = _preserve_original_text_alignment(
                    bbox, original_text, deanonymized_text, font_name, font_size, page
                )
                
                # Effacer l'ancien texte avec gestion d'erreur robuste
                try:
                    page.draw_rect(bbox, color=None, fill=fitz.pdfcolor["white"])
                except Exception as e:
                    logging.warning(f"‚ö†Ô∏è Erreur effacement rectangle: {str(e)}")
                    # Essayer avec un rectangle l√©g√®rement plus petit
                    try:
                        smaller_bbox = fitz.Rect(bbox.x0 + 1, bbox.y0 + 1, bbox.x1 - 1, bbox.y1 - 1)
                        page.draw_rect(smaller_bbox, color=None, fill=fitz.pdfcolor["white"])
                    except Exception as e2:
                        logging.warning(f"‚ö†Ô∏è Impossible d'effacer le rectangle: {str(e2)}")
                
                # CORRECTION 2: Ins√©rer le nouveau texte avec formatage pr√©serv√© lors de la d√©anonymisation
                success = _insert_text_with_preserved_formatting(
                    page, text_position, deanonymized_text, font_name, font_size, font_flags, text_color, original_text
                )
                
                if not success:
                    logging.error(f"‚ùå Impossible de remplacer le texte: '{original_text}' ‚Üí '{deanonymized_text}'")


def _insert_text_with_preserved_formatting(page, text_position, new_text, font_name, font_size, font_flags, text_color, original_text):
    """
    Ins√®re du texte en pr√©servant PRIORITAIREMENT la police originale, puis le formatage si possible.
    NOUVELLE STRAT√âGIE: Police originale > Formatage gras/italique.
    
    Args:
        page: Page PyMuPDF
        text_position: Position (x, y) pour ins√©rer le texte
        new_text: Nouveau texte √† ins√©rer
        font_name: Nom de la police originale
        font_size: Taille de la police
        font_flags: Flags de formatage (gras, italique)
        text_color: Couleur du texte original
        original_text: Texte original (pour le logging)
        
    Returns:
        bool: True si l'insertion a r√©ussi, False sinon
    """
    import fitz
    
    # Normaliser la couleur pour PyMuPDF
    normalized_color = _normalize_color(text_color)
    
    # D√©tecter le formatage du texte original
    BOLD_FLAG = 2**4   # 16
    ITALIC_FLAG = 2**5  # 32
    is_bold = bool(font_flags & BOLD_FLAG)
    is_italic = bool(font_flags & ITALIC_FLAG)
    
    # Log d√©taill√© du formatage d√©tect√©
    logging.info(f"üé® NOUVELLE STRAT√âGIE - Texte: '{original_text}' ‚Üí '{new_text}'")
    logging.info(f"üé® Police: {font_name}, Taille: {font_size}, Gras: {is_bold}, Italique: {is_italic}")
    logging.info(f"üé® Couleur: {text_color} ‚Üí {normalized_color}, Flags: {font_flags}")
    
    # PRIORIT√â 1: TOUJOURS essayer d'abord la police originale avec syst√®me de fallback intelligent
    if _try_original_font_with_fallback(page, text_position, new_text, font_name, font_size, font_flags, text_color, original_text):
        return True
    
    # PRIORIT√â 2: Si la police originale a totalement √©chou√©, essayer de pr√©server le formatage gras/italique
    if is_bold or is_italic:
        logging.info(f"üé® FALLBACK FORMATAGE - Police originale impossible, essai formatage gras/italique")
        
        try:
            # Utiliser l'√©quivalent PyMuPDF de la police
            equivalent_font = _get_pymupdf_font_equivalent(font_name)
            
            # FORCER l'application du formatage gras/italique
            if is_bold and is_italic:
                formatted_font = f"{equivalent_font}-boldoblique" if equivalent_font == "helv" else f"{equivalent_font}-bolditalic"
            elif is_bold:
                formatted_font = f"{equivalent_font}-bold"
            elif is_italic:
                formatted_font = f"{equivalent_font}-oblique" if equivalent_font == "helv" else f"{equivalent_font}-italic"
            else:
                formatted_font = equivalent_font
            
            logging.info(f"üé® TENTATIVE FORMATAGE √âQUIVALENT - Police: {font_name} ‚Üí {equivalent_font} ‚Üí {formatted_font}")
            
            page.insert_text(
                text_position,
                new_text,
                fontname=formatted_font,
                fontsize=font_size,
                color=normalized_color
            )
            logging.info(f"‚úÖ FORMATAGE √âQUIVALENT R√âUSSI - '{original_text}' ‚Üí '{new_text}' (police: {formatted_font})")
            return True
        except Exception as e:
            logging.warning(f"‚ö†Ô∏è Police √©quivalente format√©e √©chou√©e: {str(e)}")
        
        # PRIORIT√â 3: Forcer le gras avec les polices standards PyMuPDF
        if is_bold:
            try:
                # Essayer avec les polices standards PyMuPDF en gras
                standard_bold_fonts = ["helv-bold", "times-bold", "cour-bold"]
                
                for bold_font in standard_bold_fonts:
                    try:
                        page.insert_text(
                            text_position,
                            new_text,
                            fontname=bold_font,
                            fontsize=font_size,
                            color=normalized_color
                        )
                        logging.info(f"‚úÖ FORMATAGE GRAS STANDARD - '{original_text}' ‚Üí '{new_text}' (police: {bold_font})")
                        return True
                    except Exception as font_e:
                        logging.debug(f"‚ö†Ô∏è Police standard {bold_font} √©chou√©e: {str(font_e)}")
                        continue
                
            except Exception as e:
                logging.warning(f"‚ö†Ô∏è Toutes les polices standard gras ont √©chou√©: {str(e)}")
    
    # PRIORIT√â 4: Au minimum pr√©server la couleur avec police √©quivalente
    try:
        equivalent_font = _get_pymupdf_font_equivalent(font_name)
        
        page.insert_text(
            text_position,
            new_text,
            fontname=equivalent_font,
            fontsize=font_size,
            color=normalized_color
        )
        logging.info(f"‚úÖ COULEUR PR√âSERV√âE √âQUIVALENTE - '{original_text}' ‚Üí '{new_text}' (police: {font_name} ‚Üí {equivalent_font}, couleur: {normalized_color})")
        return True
    except Exception as e:
        logging.warning(f"‚ö†Ô∏è Police √©quivalente √©chou√©e: {str(e)}")
    
    # Dernier recours: Helvetica avec couleur
    try:
        page.insert_text(
            text_position,
            new_text,
            fontname="helv",
            fontsize=font_size,
            color=normalized_color
        )
        logging.warning(f"‚ö†Ô∏è FALLBACK HELVETICA - Couleur pr√©serv√©e: '{original_text}' ‚Üí '{new_text}' (couleur: {normalized_color})")
        return True
    except Exception as e:
        logging.error(f"‚ùå Dernier fallback √©chou√©: {str(e)}")
    
    return False


def _apply_font_formatting_comprehensive(base_font: str, font_flags: int) -> str:
    """
    Applique le formatage (gras, italique) √† une police de base de mani√®re compl√®te.
    Teste la disponibilit√© de chaque police avant de la retourner.
    
    Args:
        base_font: Police de base (helv, times, cour)
        font_flags: Flags de formatage PyMuPDF
        
    Returns:
        str: Nom de police avec formatage appliqu√© ou police de base
    """
    import fitz
    
    # Constantes PyMuPDF pour les flags
    BOLD_FLAG = 2**4   # 16
    ITALIC_FLAG = 2**5  # 32
    
    # D√©tecter le formatage
    is_bold = bool(font_flags & BOLD_FLAG)
    is_italic = bool(font_flags & ITALIC_FLAG)
    
    logging.debug(f"üé® Formatage d√©tect√©: flags={font_flags}, bold={is_bold}, italic={is_italic}")
    
    # Si pas de formatage, retourner la police de base
    if not is_bold and not is_italic:
        return base_font
    
    # Mapping des polices avec formatage
    font_mapping = {
        "helv": {
            "bold": "helv-bold",
            "italic": "helv-oblique",
            "bold_italic": "helv-boldoblique"
        },
        "times": {
            "bold": "times-bold",
            "italic": "times-italic",
            "bold_italic": "times-bolditalic"
        },
        "cour": {
            "bold": "cour-bold",
            "italic": "cour-oblique",
            "bold_italic": "cour-boldoblique"
        }
    }
    
    # D√©terminer la police format√©e
    if base_font in font_mapping:
        font_variants = font_mapping[base_font]
        
        if is_bold and is_italic:
            formatted_font = font_variants.get("bold_italic", base_font)
        elif is_bold:
            formatted_font = font_variants.get("bold", base_font)
        elif is_italic:
            formatted_font = font_variants.get("italic", base_font)
        else:
            formatted_font = base_font
        
        logging.debug(f"üé® Police format√©e: {base_font} ‚Üí {formatted_font}")
        return formatted_font
    
    # Si la police de base n'est pas reconnue, retourner telle quelle
    logging.debug(f"üé® Police non reconnue, retour de la police de base: {base_font}")
    return base_font

# Dictionnaire des polices Microsoft Word les plus courantes avec leurs √©quivalents PyMuPDF
MICROSOFT_WORD_FONTS = {
    # Polices serif
    "times new roman": "times",
    "times-roman": "times",
    "timesnewroman": "times",
    "times": "times",
    "georgia": "times",  # Fallback vers Times
    "book antiqua": "times",
    "bookantiqua": "times",
    "garamond": "times",
    "palatino": "times",
    "palatino linotype": "times",
    "palatinolinotype": "times",
    
    # Polices sans-serif
    "arial": "helv",
    "helvetica": "helv",
    "calibri": "helv",  # Fallback vers Helvetica
    "verdana": "helv",
    "tahoma": "helv",
    "trebuchet ms": "helv",
    "trebuchetms": "helv",
    "franklin gothic medium": "helv",
    "franklingingothicmedium": "helv",
    "century gothic": "helv",
    "centurygothic": "helv",
    "lucida sans unicode": "helv",
    "lucidasansunicode": "helv",
    "ms sans serif": "helv",
    "mssansserif": "helv",
    "segoe ui": "helv",
    "segoeui": "helv",
    
    # Polices monospace
    "courier new": "cour",
    "courier": "cour",
    "couriernew": "cour",
    "consolas": "cour",
    "lucida console": "cour",
    "lucidaconsole": "cour",
    "monaco": "cour",
    "menlo": "cour",
    
    # Polices sp√©ciales
    "comic sans ms": "helv",  # Fallback vers Helvetica
    "comicsansms": "helv",
    "impact": "helv",
    "papyrus": "helv",
    "brush script mt": "helv",
    "brushscriptmt": "helv",
    "chiller": "helv",
    "curlz mt": "helv",
    "curlzmt": "helv",
}

def _get_pymupdf_font_equivalent(font_name: str) -> str:
    """
    Trouve l'√©quivalent PyMuPDF d'une police Microsoft Word.
    
    Args:
        font_name: Nom de la police originale
        
    Returns:
        str: Nom de la police √©quivalente PyMuPDF (helv, times, cour)
    """
    if not font_name:
        return "helv"
    
    # Normaliser le nom de la police (minuscules, sans espaces)
    normalized_name = font_name.lower().replace(" ", "").replace("-", "")
    
    # Chercher dans le dictionnaire
    equivalent = MICROSOFT_WORD_FONTS.get(normalized_name)
    
    if equivalent:
        logging.debug(f"üîÑ Police mapp√©e: {font_name} ‚Üí {equivalent}")
        return equivalent
    
    # Si pas trouv√©, essayer de deviner par le nom
    if "times" in normalized_name or "roman" in normalized_name:
        logging.debug(f"üîÑ Police devin√©e (serif): {font_name} ‚Üí times")
        return "times"
    elif "courier" in normalized_name or "mono" in normalized_name or "console" in normalized_name:
        logging.debug(f"üîÑ Police devin√©e (monospace): {font_name} ‚Üí cour")
        return "cour"
    else:
        logging.debug(f"üîÑ Police par d√©faut (sans-serif): {font_name} ‚Üí helv")
        return "helv"

def _try_original_font_with_fallback(page, text_position, new_text, font_name, font_size, font_flags, text_color, original_text):
    """
    Essaie d'utiliser la police originale avec un syst√®me de fallback intelligent.
    
    Args:
        page: Page PyMuPDF
        text_position: Position (x, y) pour ins√©rer le texte
        new_text: Nouveau texte √† ins√©rer
        font_name: Nom de la police originale
        font_size: Taille de la police
        font_flags: Flags de formatage (gras, italique)
        text_color: Couleur du texte original
        original_text: Texte original (pour le logging)
        
    Returns:
        bool: True si l'insertion a r√©ussi, False sinon
    """
    import fitz
    
    # Normaliser la couleur pour PyMuPDF
    normalized_color = _normalize_color(text_color)
    
    # D√©tecter le formatage du texte original
    BOLD_FLAG = 2**4   # 16
    ITALIC_FLAG = 2**5  # 32
    is_bold = bool(font_flags & BOLD_FLAG)
    is_italic = bool(font_flags & ITALIC_FLAG)
    
    logging.info(f"üé® TENTATIVE POLICE ORIGINALE - Texte: '{original_text}' ‚Üí '{new_text}'")
    logging.info(f"üé® Police: {font_name}, Taille: {font_size}, Gras: {is_bold}, Italique: {is_italic}")
    
    # √âTAPE 1: Essayer la police originale exacte (priorit√© absolue)
    try:
        page.insert_text(
            text_position,
            new_text,
            fontname=font_name,
            fontsize=font_size,
            color=normalized_color
        )
        logging.info(f"‚úÖ POLICE ORIGINALE EXACTE R√âUSSIE - '{original_text}' ‚Üí '{new_text}' (police: {font_name})")
        return True
    except Exception as e:
        logging.debug(f"‚ö†Ô∏è Police originale exacte √©chou√©e: {str(e)}")
    
    # √âTAPE 2: Si police originale √©choue, essayer sans formatage (juste la police de base)
    try:
        # Retirer le formatage du nom de police (enlever -bold, -italic, etc.)
        base_font_name = font_name.replace("-bold", "").replace("-italic", "").replace("-oblique", "").replace("-bolditalic", "").replace("-boldoblique", "")
        
        page.insert_text(
            text_position,
            new_text,
            fontname=base_font_name,
            fontsize=font_size,
            color=normalized_color
        )
        logging.info(f"‚úÖ POLICE ORIGINALE BASE R√âUSSIE - '{original_text}' ‚Üí '{new_text}' (police: {base_font_name})")
        logging.warning(f"‚ö†Ô∏è FORMATAGE PERDU - Le gras/italique n'a pas pu √™tre appliqu√© pour pr√©server la police originale")
        return True
    except Exception as e:
        logging.debug(f"‚ö†Ô∏è Police originale base √©chou√©e: {str(e)}")
    
    # √âTAPE 3: Essayer l'√©quivalent PyMuPDF avec formatage pr√©serv√©
    try:
        equivalent_font = _get_pymupdf_font_equivalent(font_name)
        
        # APPLIQUER LE FORMATAGE √† la police √©quivalente
        if is_bold and is_italic:
            formatted_equivalent = f"{equivalent_font}-boldoblique" if equivalent_font == "helv" else f"{equivalent_font}-bolditalic"
        elif is_bold:
            formatted_equivalent = f"{equivalent_font}-bold"
        elif is_italic:
            formatted_equivalent = f"{equivalent_font}-oblique" if equivalent_font == "helv" else f"{equivalent_font}-italic"
        else:
            formatted_equivalent = equivalent_font
        
        logging.info(f"üé® FORMATAGE √âQUIVALENT - Police: {font_name} ‚Üí {equivalent_font} ‚Üí {formatted_equivalent}")
        
        page.insert_text(
            text_position,
            new_text,
            fontname=formatted_equivalent,
            fontsize=font_size,
            color=normalized_color
        )
        logging.info(f"‚úÖ POLICE √âQUIVALENTE + FORMATAGE R√âUSSI - '{original_text}' ‚Üí '{new_text}' (police: {font_name} ‚Üí {formatted_equivalent})")
        return True
    except Exception as e:
        logging.debug(f"‚ö†Ô∏è Police √©quivalente format√©e √©chou√©e: {str(e)}")
    
    # √âTAPE 4: Fallback police √©quivalente sans formatage
    try:
        equivalent_font = _get_pymupdf_font_equivalent(font_name)
        
        page.insert_text(
            text_position,
            new_text,
            fontname=equivalent_font,
            fontsize=font_size,
            color=normalized_color
        )
        logging.info(f"‚úÖ POLICE √âQUIVALENTE SANS FORMATAGE - '{original_text}' ‚Üí '{new_text}' (police: {font_name} ‚Üí {equivalent_font})")
        logging.warning(f"‚ö†Ô∏è FORMATAGE PERDU - Le gras/italique n'a pas pu √™tre appliqu√© avec la police √©quivalente")
        return True
    except Exception as e:
        logging.debug(f"‚ö†Ô∏è Police √©quivalente sans formatage √©chou√©e: {str(e)}")
    
    return False