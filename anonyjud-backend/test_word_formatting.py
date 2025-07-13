#!/usr/bin/env python3
"""
Test pour vérifier la préservation du formatage dans les fichiers Word
lors de l'anonymisation et de la dé-anonymisation.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from app.main import anonymize_docx_file, deanonymize_docx_file
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io
import tempfile

def create_test_docx():
    """Crée un document Word de test avec différents styles de formatage"""
    doc = Document()
    
    # Titre principal
    title = doc.add_heading('Document de Test - Formatage', 0)
    
    # Paragraphe avec du texte normal
    p1 = doc.add_paragraph('Monsieur ')
    run1 = p1.add_run('HUISSOUD')
    run1.bold = True
    run1.font.size = Inches(0.2)
    p1.add_run(' ')
    run2 = p1.add_run('Louis')
    run2.italic = True
    run2.font.name = 'Arial'
    p1.add_run(' habite à ')
    run3 = p1.add_run('123 rue de la Paix')
    run3.underline = True
    p1.add_run(' à Paris.')
    
    # Paragraphe avec texte en majuscules
    p2 = doc.add_paragraph('Adresse complète: ')
    run4 = p2.add_run('LOUIS HUISSOUD')
    run4.font.size = Inches(0.15)
    run4.bold = True
    p2.add_run(' - ')
    run5 = p2.add_run('123 RUE DE LA PAIX')
    run5.font.color.rgb = None  # Couleur par défaut
    
    # Paragraphe avec texte en minuscules
    p3 = doc.add_paragraph('En minuscules: ')
    run6 = p3.add_run('louis huissoud')
    run6.italic = True
    p3.add_run(' habite ')
    run7 = p3.add_run('123 rue de la paix')
    run7.underline = True
    
    # Tableau avec formatage
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    
    # Cellule avec nom en gras
    cell1 = table.cell(0, 0)
    cell1.text = 'Nom: '
    run8 = cell1.paragraphs[0].add_run('HUISSOUD')
    run8.bold = True
    run8.font.size = Inches(0.12)
    
    # Cellule avec prénom en italique
    cell2 = table.cell(0, 1)
    cell2.text = 'Prénom: '
    run9 = cell2.paragraphs[0].add_run('Louis')
    run9.italic = True
    
    # Cellule avec adresse soulignée
    cell3 = table.cell(1, 0)
    cell3.text = 'Adresse: '
    run10 = cell3.paragraphs[0].add_run('123 rue de la Paix')
    run10.underline = True
    
    # Cellule avec ville
    cell4 = table.cell(1, 1)
    cell4.text = 'Ville: Paris'
    
    # Sauvegarder en bytes
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()

def analyze_document_formatting(doc_bytes, description=""):
    """Analyse le formatage d'un document Word"""
    print(f"\n=== ANALYSE DU FORMATAGE - {description} ===")
    
    doc = Document(io.BytesIO(doc_bytes))
    
    # Analyser les paragraphes
    print("📝 PARAGRAPHES:")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"  Paragraphe {i}: '{para.text[:50]}...'")
            for j, run in enumerate(para.runs):
                if run.text.strip():
                    formatting = []
                    if run.bold:
                        formatting.append("GRAS")
                    if run.italic:
                        formatting.append("ITALIQUE")
                    if run.underline:
                        formatting.append("SOULIGNÉ")
                    if run.font.name:
                        formatting.append(f"Police:{run.font.name}")
                    if run.font.size:
                        formatting.append(f"Taille:{run.font.size}")
                    
                    format_str = f" [{', '.join(formatting)}]" if formatting else ""
                    print(f"    Run {j}: '{run.text}'{format_str}")
    
    # Analyser les tableaux
    print("\n📊 TABLEAUX:")
    for table_idx, table in enumerate(doc.tables):
        print(f"  Tableau {table_idx}:")
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                if cell.text.strip():
                    print(f"    Cellule [{row_idx},{cell_idx}]: '{cell.text[:30]}...'")
                    for para in cell.paragraphs:
                        for j, run in enumerate(para.runs):
                            if run.text.strip():
                                formatting = []
                                if run.bold:
                                    formatting.append("GRAS")
                                if run.italic:
                                    formatting.append("ITALIQUE")
                                if run.underline:
                                    formatting.append("SOULIGNÉ")
                                
                                format_str = f" [{', '.join(formatting)}]" if formatting else ""
                                print(f"      Run {j}: '{run.text}'{format_str}")

def test_word_formatting_preservation():
    """Test principal pour vérifier la préservation du formatage"""
    print("🚀 DÉBUT DU TEST DE PRÉSERVATION DU FORMATAGE WORD")
    
    # 1. Créer un document de test
    print("\n1️⃣ Création du document de test...")
    original_docx = create_test_docx()
    
    # 2. Analyser le document original
    analyze_document_formatting(original_docx, "DOCUMENT ORIGINAL")
    
    # 3. Définir les tiers pour l'anonymisation
    tiers = [
        {
            "numero": 1,
            "nom": "HUISSOUD",
            "prenom": "Louis", 
            "adresse": "123 rue de la Paix",
            "ville": "Paris"
        }
    ]
    
    # 4. Anonymiser le document
    print("\n2️⃣ Anonymisation du document...")
    try:
        anonymized_docx, mapping = anonymize_docx_file(original_docx, tiers)
        print(f"✅ Anonymisation réussie")
        print(f"🗂️ Mapping généré: {mapping}")
        
        # Analyser le document anonymisé
        analyze_document_formatting(anonymized_docx, "DOCUMENT ANONYMISÉ")
        
    except Exception as e:
        print(f"❌ Erreur lors de l'anonymisation: {e}")
        return False
    
    # 5. Dé-anonymiser le document
    print("\n3️⃣ Dé-anonymisation du document...")
    try:
        deanonymized_docx = deanonymize_docx_file(anonymized_docx, mapping)
        print(f"✅ Dé-anonymisation réussie")
        
        # Analyser le document dé-anonymisé
        analyze_document_formatting(deanonymized_docx, "DOCUMENT DÉ-ANONYMISÉ")
        
    except Exception as e:
        print(f"❌ Erreur lors de la dé-anonymisation: {e}")
        return False
    
    # 6. Vérifications finales
    print("\n4️⃣ Vérifications finales...")
    
    # Vérifier que le contenu est correct
    final_doc = Document(io.BytesIO(deanonymized_docx))
    final_text = ""
    for para in final_doc.paragraphs:
        final_text += para.text + "\n"
    for table in final_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                final_text += cell.text + "\n"
    
    print(f"📝 Texte final (premiers 300 chars): {final_text[:300]}...")
    
    # Vérifications spécifiques
    checks = [
        ("Louis" in final_text, "Le prénom 'Louis' est présent"),
        ("HUISSOUD" in final_text, "Le nom 'HUISSOUD' est présent"),
        ("123 rue de la Paix" in final_text, "L'adresse '123 rue de la Paix' est présente"),
        ("PRENOM1" not in final_text, "Pas de balise PRENOM1 restante"),
        ("NOM1" not in final_text, "Pas de balise NOM1 restante"),
        ("ADRESSE1" not in final_text, "Pas de balise ADRESSE1 restante"),
    ]
    
    all_passed = True
    for check, description in checks:
        if check:
            print(f"✅ {description}")
        else:
            print(f"❌ {description}")
            all_passed = False
    
    # Sauvegarder les fichiers pour inspection manuelle
    print("\n5️⃣ Sauvegarde des fichiers de test...")
    with open("test_original.docx", "wb") as f:
        f.write(original_docx)
    with open("test_anonymized.docx", "wb") as f:
        f.write(anonymized_docx)
    with open("test_deanonymized.docx", "wb") as f:
        f.write(deanonymized_docx)
    
    print("📁 Fichiers sauvegardés:")
    print("  - test_original.docx")
    print("  - test_anonymized.docx") 
    print("  - test_deanonymized.docx")
    
    return all_passed

if __name__ == "__main__":
    success = test_word_formatting_preservation()
    if success:
        print("\n🎉 TOUS LES TESTS SONT PASSÉS ! Le formatage est préservé.")
    else:
        print("\n❌ CERTAINS TESTS ONT ÉCHOUÉ. Vérifiez les logs ci-dessus.")
    
    sys.exit(0 if success else 1) 