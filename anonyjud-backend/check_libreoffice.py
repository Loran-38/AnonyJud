#!/usr/bin/env python3
"""
Script de vérification de LibreOffice pour Railway
Vérifie que LibreOffice est installé et fonctionnel pour la conversion Word → PDF
"""

import subprocess
import sys
import logging

# Configuration des logs
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def check_libreoffice_installation():
    """Vérifie que LibreOffice est installé et accessible"""
    logger.info("🔍 Vérification de l'installation LibreOffice...")
    
    # Commandes LibreOffice à tester
    libreoffice_commands = [
        "soffice",
        "libreoffice",
        "/usr/bin/soffice",
        "/opt/libreoffice7.1/program/soffice"
    ]
    
    working_command = None
    
    for cmd in libreoffice_commands:
        try:
            # Tester la commande
            result = subprocess.run([cmd, "--version"], 
                                  capture_output=True, text=True, timeout=10)
            
            if result.returncode == 0:
                working_command = cmd
                logger.info(f"✅ LibreOffice trouvé: {cmd}")
                logger.info(f"📄 Version: {result.stdout.strip()}")
                break
            else:
                logger.debug(f"❌ Commande {cmd} échouée: {result.stderr}")
                
        except (subprocess.TimeoutExpired, FileNotFoundError) as e:
            logger.debug(f"❌ Commande {cmd} non trouvée ou timeout: {str(e)}")
            continue
    
    if working_command:
        return working_command
    else:
        logger.error("❌ Aucune commande LibreOffice trouvée")
        return None

def test_libreoffice_conversion():
    """Test de conversion Word → PDF avec LibreOffice"""
    logger.info("🔄 Test de conversion Word → PDF avec LibreOffice...")
    
    try:
        # Créer un document Word de test simple
        from docx import Document
        from docx.shared import Inches
        
        # Créer un document de test
        doc = Document()
        doc.add_heading('Test LibreOffice', 0)
        doc.add_paragraph('Ceci est un test de conversion Word → PDF avec LibreOffice.')
        
        # Sauvegarder temporairement
        import tempfile
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_docx:
            temp_docx_path = temp_docx.name
            doc.save(temp_docx_path)
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
            temp_pdf_path = temp_pdf.name
        
        try:
            # Tester la conversion
            working_command = check_libreoffice_installation()
            
            if working_command:
                result = subprocess.run([
                    working_command,
                    "--headless",
                    "--convert-to", "pdf",
                    "--outdir", tempfile.gettempdir(),
                    temp_docx_path
                ], capture_output=True, text=True, timeout=30)
                
                if result.returncode == 0:
                    # Vérifier que le PDF a été créé
                    expected_pdf_path = temp_docx_path.replace('.docx', '.pdf')
                    
                    import os
                    if os.path.exists(expected_pdf_path):
                        file_size = os.path.getsize(expected_pdf_path)
                        logger.info(f"✅ Conversion LibreOffice réussie!")
                        logger.info(f"📄 PDF généré: {expected_pdf_path}")
                        logger.info(f"📊 Taille: {file_size} bytes")
                        return True
                    else:
                        logger.error(f"❌ Fichier PDF non généré: {expected_pdf_path}")
                        return False
                else:
                    logger.error(f"❌ Erreur conversion LibreOffice: {result.stderr}")
                    return False
            else:
                logger.error("❌ LibreOffice non disponible")
                return False
                
        finally:
            # Nettoyer les fichiers temporaires
            import os
            for path in [temp_docx_path, temp_pdf_path]:
                if os.path.exists(path):
                    os.unlink(path)
            
            # Nettoyer le PDF généré
            expected_pdf_path = temp_docx_path.replace('.docx', '.pdf')
            if os.path.exists(expected_pdf_path):
                os.unlink(expected_pdf_path)
                
    except Exception as e:
        logger.error(f"❌ Erreur lors du test de conversion: {str(e)}")
        return False

def check_system_requirements():
    """Vérifie les prérequis système"""
    logger.info("🔍 Vérification des prérequis système...")
    
    # Vérifier Python
    logger.info(f"🐍 Python version: {sys.version}")
    
    # Vérifier les modules Python nécessaires
    required_modules = [
        'docx',
        'fitz',  # PyMuPDF
        'reportlab',
        'pdf2docx'
    ]
    
    missing_modules = []
    for module in required_modules:
        try:
            __import__(module)
            logger.info(f"✅ Module {module} disponible")
        except ImportError:
            missing_modules.append(module)
            logger.error(f"❌ Module {module} manquant")
    
    if missing_modules:
        logger.error(f"❌ Modules manquants: {missing_modules}")
        return False
    else:
        logger.info("✅ Tous les modules Python requis sont disponibles")
        return True

def main():
    """Fonction principale de vérification"""
    logger.info("🚀 Démarrage de la vérification LibreOffice pour Railway")
    
    # Vérifier les prérequis système
    system_ok = check_system_requirements()
    
    # Vérifier l'installation LibreOffice
    libreoffice_ok = check_libreoffice_installation() is not None
    
    # Tester la conversion
    conversion_ok = test_libreoffice_conversion()
    
    # Résumé
    logger.info("📊 Résumé de la vérification:")
    logger.info(f"   - Prérequis système: {'✅' if system_ok else '❌'}")
    logger.info(f"   - LibreOffice installé: {'✅' if libreoffice_ok else '❌'}")
    logger.info(f"   - Conversion fonctionnelle: {'✅' if conversion_ok else '❌'}")
    
    if system_ok and libreoffice_ok and conversion_ok:
        logger.info("🎉 Toutes les vérifications ont réussi! LibreOffice est prêt pour Railway.")
        return True
    else:
        logger.error("❌ Certaines vérifications ont échoué. Vérifiez l'installation.")
        return False

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1) 