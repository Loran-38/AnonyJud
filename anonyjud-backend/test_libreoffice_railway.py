#!/usr/bin/env python3
"""
Test spécifique de LibreOffice pour Railway
Vérifie que LibreOffice est installé et fonctionnel pour la conversion Word → PDF
"""

import subprocess
import sys
import os
import tempfile
import logging

# Configuration des logs
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def test_libreoffice_installation():
    """Test complet de l'installation LibreOffice"""
    logger.info("🔍 Test complet de LibreOffice pour Railway")
    
    # 1. Vérifier les commandes disponibles
    commands_to_test = [
        "soffice",
        "libreoffice", 
        "/usr/bin/soffice",
        "/opt/libreoffice7.1/program/soffice"
    ]
    
    working_command = None
    
    for cmd in commands_to_test:
        try:
            logger.info(f"📄 Test de la commande: {cmd}")
            
            # Test de version
            result = subprocess.run([cmd, "--version"], 
                                  capture_output=True, text=True, timeout=10)
            
            if result.returncode == 0:
                working_command = cmd
                logger.info(f"✅ LibreOffice trouvé: {cmd}")
                logger.info(f"📄 Version: {result.stdout.strip()}")
                break
            else:
                logger.warning(f"❌ Commande {cmd} échouée: {result.stderr}")
                
        except (subprocess.TimeoutExpired, FileNotFoundError) as e:
            logger.warning(f"❌ Commande {cmd} non trouvée ou timeout: {str(e)}")
            continue
    
    if not working_command:
        logger.error("❌ Aucune commande LibreOffice trouvée")
        return False
    
    # 2. Test de conversion simple
    logger.info("🔄 Test de conversion Word → PDF...")
    
    try:
        # Créer un document Word de test
        from docx import Document
        
        doc = Document()
        doc.add_heading('Test LibreOffice Railway', 0)
        doc.add_paragraph('Ceci est un test de conversion Word → PDF avec LibreOffice sur Railway.')
        doc.add_paragraph('Si ce test réussit, LibreOffice fonctionne correctement.')
        
        # Sauvegarder temporairement
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_docx:
            temp_docx_path = temp_docx.name
            doc.save(temp_docx_path)
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
            temp_pdf_path = temp_pdf.name
        
        try:
            # Conversion avec LibreOffice
            logger.info(f"📄 Conversion avec: {working_command}")
            
            result = subprocess.run([
                working_command,
                "--headless",
                "--convert-to", "pdf:writer_pdf_Export",
                "--outdir", tempfile.gettempdir(),
                temp_docx_path
            ], capture_output=True, text=True, timeout=60)
            
            if result.returncode == 0:
                # Vérifier que le PDF a été créé
                expected_pdf_path = temp_docx_path.replace('.docx', '.pdf')
                
                if os.path.exists(expected_pdf_path):
                    file_size = os.path.getsize(expected_pdf_path)
                    logger.info(f"✅ Conversion LibreOffice réussie!")
                    logger.info(f"📄 PDF généré: {expected_pdf_path}")
                    logger.info(f"📊 Taille: {file_size} bytes")
                    
                    # Vérifier que le PDF est valide
                    if file_size > 1000:  # PDF valide doit faire plus de 1KB
                        logger.info("✅ PDF valide généré")
                        return True
                    else:
                        logger.warning("⚠️ PDF généré trop petit, possible erreur")
                        return False
                else:
                    logger.error(f"❌ Fichier PDF non généré: {expected_pdf_path}")
                    return False
            else:
                logger.error(f"❌ Erreur conversion LibreOffice: {result.stderr}")
                return False
                
        finally:
            # Nettoyer les fichiers temporaires
            for path in [temp_docx_path, temp_pdf_path]:
                if os.path.exists(path):
                    os.unlink(path)
            
            # Nettoyer le PDF généré
            expected_pdf_path = temp_docx_path.replace('.docx', '.pdf')
            if os.path.exists(expected_pdf_path):
                os.unlink(expected_pdf_path)
                
    except Exception as e:
        logger.error(f"❌ Erreur lors du test de conversion: {str(e)}")
        return False

def test_system_info():
    """Affiche les informations système"""
    logger.info("📊 Informations système:")
    
    # OS
    logger.info(f"   - OS: {sys.platform}")
    
    # Python
    logger.info(f"   - Python: {sys.version}")
    
    # Répertoires
    logger.info(f"   - PATH: {os.environ.get('PATH', 'Non défini')}")
    
    # Vérifier les répertoires LibreOffice
    libreoffice_paths = [
        "/usr/bin/soffice",
        "/opt/libreoffice7.1/program/soffice",
        "/usr/lib/libreoffice/program/soffice"
    ]
    
    for path in libreoffice_paths:
        if os.path.exists(path):
            logger.info(f"   - LibreOffice trouvé: {path}")
        else:
            logger.debug(f"   - LibreOffice non trouvé: {path}")

def main():
    """Fonction principale"""
    logger.info("🚀 Test LibreOffice pour Railway")
    
    # Informations système
    test_system_info()
    
    # Test LibreOffice
    success = test_libreoffice_installation()
    
    if success:
        logger.info("🎉 LibreOffice fonctionne correctement sur Railway!")
        return True
    else:
        logger.error("❌ LibreOffice ne fonctionne pas sur Railway")
        return False

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1) 